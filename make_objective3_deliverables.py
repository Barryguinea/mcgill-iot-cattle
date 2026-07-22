# -*- coding: utf-8 -*-
"""Prépare et empaquette les livrables de l'Objectif 3, sans PDF."""
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parent
PROJECT_PARENT = PROJECT.parent
REPORTS = PROJECT / "reports" / "objective3_demonstration"
DATA_DEMO = REPORTS / "DONNEES_DEMO"
NOTEBOOK = PROJECT / "notebooks" / "12_objectif3_notebook_demonstration.ipynb"
DOCUMENTATION = REPORTS / "documentation_notebook.md"
WORD_REPORT = REPORTS / "Objectif3_guide_utilisation.docx"
PRESENTATION = REPORTS / "Objectif3_presentation_detaillee.pptx"
PACKAGE = (
    Path.home()
    / "Desktop"
    / "Livrables_McGill_WellE"
    / "Objectif3_Notebook_demonstration"
)

O1 = PROJECT / "reports" / "objective1_pipeline_icetag"
O2 = PROJECT / "reports" / "objective2_environnement"
CORE_SOURCE = PROJECT_PARENT / "core"
THRESHOLDS_SOURCE = PROJECT_PARENT / "data" / "final_thresholds_v1.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "5F6368"
LIGHT_BLUE = "E8EEF5"
BORDER = "B7C2CE"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def _markdown(source: str) -> dict:
    source = source.strip()
    return {
        "cell_type": "markdown",
        "id": hashlib.sha1(f"markdown:{source}".encode("utf-8")).hexdigest()[:8],
        "metadata": {},
        "source": [line + "\n" for line in source.splitlines()],
    }


def _code(source: str) -> dict:
    source = source.strip()
    return {
        "cell_type": "code",
        "id": hashlib.sha1(f"code:{source}".encode("utf-8")).hexdigest()[:8],
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": [line + "\n" for line in source.splitlines()],
    }


def _build_demo_data() -> None:
    DATA_DEMO.mkdir(parents=True, exist_ok=True)
    stale_figure = REPORTS / "demo_alertes_horaire.png"
    if stale_figure.exists():
        stale_figure.unlink()

    activity = pd.read_csv(O1 / "summer_2019_pipeline_input_15min.csv")
    activity["Cow"] = activity["Cow"].astype(str)
    demo_cow = "2062"
    one_cow = activity[activity["Cow"].eq(demo_cow)].copy()
    if len(one_cow) != 8922:
        raise ValueError(f"Échantillon inattendu pour la vache {demo_cow}: {len(one_cow)}")
    one_cow.to_csv(DATA_DEMO / "summer2019_vache_2062_icetag_15min.csv", index=False)

    shutil.copy2(
        O1 / "objective1_multi_season_summary.csv",
        DATA_DEMO / "objective1_multi_season_summary.csv",
    )
    shutil.copy2(
        O1 / "summer_2019_pipeline_alerts_only.csv",
        DATA_DEMO / "summer2019_pipeline_alerts_only.csv",
    )
    shutil.copy2(
        O1
        / "renforcement_scientifique"
        / "objective1_reinforced_summary_by_confidence.csv",
        DATA_DEMO / "objective1_reinforced_summary_by_confidence.csv",
    )

    environment = pd.read_csv(O2 / "summer2019_icetag_environnement_15min.csv")
    environment["Start"] = pd.to_datetime(environment["Start"])
    herd_environment = (
        environment.groupby("Start", as_index=False)
        .agg(Steps=("Steps", "mean"), THI=("THI", "first"), n_cows=("Cow", "nunique"))
        .sort_values("Start")
    )
    if len(herd_environment) != 5795:
        raise ValueError(f"Nombre d'horodatages inattendu: {len(herd_environment)}")
    herd_environment.to_csv(
        DATA_DEMO / "summer2019_environment_herd_15min.csv", index=False
    )

    shutil.copy2(
        O2 / "summer2019_comportement_environnement.csv",
        DATA_DEMO / "summer2019_comportement_environnement.csv",
    )
    shutil.copy2(
        O2 / "mixed_model" / "objective2_thi_controle_complet.csv",
        DATA_DEMO / "objective2_thi_controle_complet.csv",
    )


def _build_notebook() -> None:
    cells = [
        _markdown(
            """
# Notebook de démonstration IoT en élevage laitier
## Objectif 3 du SOW — Tâche 3.1

Ce notebook présente un flux reproductible, de la lecture des données jusqu'à
l'interprétation scientifique. Il utilise des données McGill déjà traitées et un
échantillon IceTag compact pour que la démonstration reste rapide.

**Parcours couvert par le SOW :**
1. chargement et prétraitement des données;
2. exécution réelle du pipeline sur des données IceTag;
3. visualisations environnement–activité et environnement–comportement;
4. guide d'interprétation, contrôles et limites.

> Une alerte est un signal comportemental à vérifier, pas un diagnostic clinique de boiterie.
"""
        ),
        _code(
            """
from pathlib import Path
import json
import sys
import warnings

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from IPython.display import display
from scipy.stats import spearmanr

warnings.filterwarnings("ignore")
plt.rcParams.update({"figure.dpi": 110, "axes.grid": True, "grid.alpha": 0.22})

def locate_context():
    candidates = [Path.cwd().resolve(), *Path.cwd().resolve().parents]
    for root in candidates:
        repo_data = root / "reports" / "objective3_demonstration" / "DONNEES_DEMO"
        if repo_data.exists():
            return {
                "mode": "dépôt",
                "root": root,
                "data": repo_data,
                "output": root / "reports" / "objective3_demonstration",
                "runtime": root.parent,
                "thresholds": root.parent / "data" / "final_thresholds_v1.json",
            }
        delivery_data = root / "DONNEES_DEMO"
        if delivery_data.exists() and (root / "code" / "core").exists():
            return {
                "mode": "livraison autonome",
                "root": root,
                "data": delivery_data,
                "output": root / "RESULTATS",
                "runtime": root / "code",
                "thresholds": root / "code" / "final_thresholds_v1.json",
            }
    raise FileNotFoundError(
        "Dossier DONNEES_DEMO introuvable. Exécuter le notebook depuis le dépôt "
        "ou depuis le paquet Objectif3_Notebook_demonstration."
    )

CTX = locate_context()
DATA = CTX["data"]
OUT = CTX["output"]
OUT.mkdir(parents=True, exist_ok=True)
if str(CTX["runtime"]) not in sys.path:
    sys.path.insert(0, str(CTX["runtime"]))

from core.io import normalize_columns
from core.pipeline import run_pipeline_one_cow, summarize_one_cow

print(f"Mode d'exécution : {CTX['mode']}")
print(f"Données de démonstration : {DATA}")
print(f"Résultats : {OUT}")
"""
        ),
        _markdown(
            """
## 1. Chargement et prétraitement

L'échantillon contient toutes les observations Summer 2019 de la vache 2062.
Les colonnes hétérogènes sont normalisées vers le schéma du pipeline : `Cow`, `T`,
`Steps`, `Motion Index`, durées de couchage/debout et transitions.
"""
        ),
        _code(
            """
raw = pd.read_csv(DATA / "summer2019_vache_2062_icetag_15min.csv")
required_raw = {"Cow", "Start", "Steps", "Motion Index", "Lying Time", "Transitions"}
missing = required_raw.difference(raw.columns)
assert not missing, f"Colonnes manquantes : {sorted(missing)}"

activity = normalize_columns(raw)
assert activity["Cow"].nunique() == 1
assert activity["T"].is_monotonic_increasing
assert activity[["Steps", "Motion Index"]].notna().all().all()

print("Lignes :", f"{len(activity):,}".replace(",", " "))
print(f"Vache : {activity['Cow'].iloc[0]}")
print(f"Période : {activity['T'].min()} au {activity['T'].max()}")
print(f"Pas de temps médian : {activity['T'].diff().median()}")
display(activity[["Cow", "T", "Steps", "Motion Index", "Lying Time", "Transitions"]].head())
"""
        ),
        _markdown(
            """
### Profil journalier du signal IceTag

Le profil horaire sert d'abord à comprendre le rythme d'activité avant de chercher
des anomalies. Il ne constitue pas, à lui seul, un indicateur de boiterie.
"""
        ),
        _code(
            """
profile = activity.assign(hour=activity["T"].dt.hour).groupby("hour")["Steps"].mean()
fig, ax = plt.subplots(figsize=(9, 4.2))
ax.plot(profile.index, profile.values, marker="o", linewidth=2.2, color="#2E74B5")
ax.fill_between(profile.index, profile.values, alpha=0.16, color="#2E74B5")
ax.set(title="Profil journalier moyen — vache 2062", xlabel="Heure", ylabel="Pas moyens / 15 min")
ax.set_xticks(range(0, 24, 2))
fig.tight_layout()
fig.savefig(OUT / "demo_profil_journalier.png", dpi=150)
plt.show()
"""
        ),
        _markdown(
            """
## 2. Exécution réelle du pipeline IceTag

La cellule suivante charge les paramètres gelés, construit les caractéristiques,
ajuste l'Isolation Forest et applique les règles de persistance et de notification.
La démonstration porte sur une vache afin de rester rapide et vérifiable.
"""
        ),
        _code(
            """
with open(CTX["thresholds"], encoding="utf-8") as handle:
    params = json.load(handle)["pipeline_defaults"].copy()
params.pop("sensor_warmup_bins", None)

cow_id = activity["Cow"].iloc[0]
pipeline_output = run_pipeline_one_cow(activity, cow_id, **params)
pipeline_summary = summarize_one_cow(pipeline_output)

expected = {"n_bins": 8922, "if_anomaly_points": 582, "lameness_notifs": 17}
for key, value in expected.items():
    assert pipeline_summary[key] == value, (key, pipeline_summary[key], value)

display(pd.DataFrame([pipeline_summary]).rename(columns={
    "n_bins": "intervalles",
    "if_anomaly_points": "points atypiques IF",
    "lameness_notifs": "alertes brutes",
    "coverage_mean": "couverture moyenne (%)",
}))
print("Pipeline exécuté avec succès sur la vache 2062.")
"""
        ),
        _code(
            """
alerts = pipeline_output[pipeline_output["notif_lameness"].eq(1)].copy()
first_alert = alerts["T"].min()
window = pipeline_output[
    pipeline_output["T"].between(first_alert - pd.Timedelta(days=2), first_alert + pd.Timedelta(days=5))
].copy()

fig, ax = plt.subplots(figsize=(10, 4.3))
ax.plot(window["T"], window["Steps_sum"], color="#5F6368", linewidth=1, label="Pas / 15 min")
marked = window[window["notif_lameness"].eq(1)]
ax.scatter(marked["T"], marked["Steps_sum"], color="#C62828", s=55, zorder=3, label="Alerte")
ax.set(title="Exemple d'alertes produites par le pipeline", xlabel="Date", ylabel="Pas / 15 min")
ax.legend(frameon=False)
fig.autofmt_xdate()
fig.tight_layout()
fig.savefig(OUT / "demo_pipeline_execution.png", dpi=150)
plt.show()
"""
        ),
        _markdown(
            """
### Mise en contexte multi-saisons

Les 385 alertes initiales sont des candidats. Le renforcement de l'Objectif 1 les
classe en niveaux A, B et C selon leur caractère individuel ou collectif.
"""
        ),
        _code(
            """
multi = pd.read_csv(DATA / "objective1_multi_season_summary.csv")
confidence = pd.read_csv(DATA / "objective1_reinforced_summary_by_confidence.csv")

confidence_totals = confidence.groupby("reinforced_confidence_level")["notifications"].sum()
assert int(multi["lameness_notifs"].sum()) == 385
assert int(confidence_totals.sum()) == 385

view = multi[["season", "n_cows", "n_intervals", "lameness_notifs", "notifs_per_100_cow_days"]].copy()
view.columns = ["Saison", "Vaches", "Intervalles", "Alertes brutes", "Alertes / 100 vache-jours"]
display(view)

fig, ax = plt.subplots(figsize=(8.8, 4.2))
ax.bar(view["Saison"], view["Alertes / 100 vache-jours"], color="#4C78A8")
ax.set(title="Taux d'alertes brutes normalisé", ylabel="Alertes / 100 vache-jours", xlabel="Saison")
fig.tight_layout()
fig.savefig(OUT / "demo_alertes_par_essai.png", dpi=150)
plt.show()

labels = {
    "A_individuelle_prioritaire": "A — prioritaire",
    "B_individuelle_a_verifier": "B — à vérifier",
    "C_probable_evenement_collectif": "C — contexte collectif",
}
print("Requalification des 385 alertes :")
for level, count in confidence_totals.items():
    print(f"- {labels[level]} : {int(count)}")
"""
        ),
        _markdown(
            """
## 3. Environnement, activité et comportements

Le THI combine température et humidité. Les graphiques ci-dessous sont descriptifs :
la comparaison brute des catégories de THI ne suffit pas à établir un effet causal.
"""
        ),
        _code(
            """
environment = pd.read_csv(DATA / "summer2019_environment_herd_15min.csv", parse_dates=["Start"])

def thi_category(value):
    if value < 68:
        return "Aucun (< 68)"
    if value < 72:
        return "Léger (68–71,9)"
    if value < 80:
        return "Modéré (72–79,9)"
    return "Sévère (≥ 80)"

environment["THI_cat"] = environment["THI"].map(thi_category)
order = ["Aucun (< 68)", "Léger (68–71,9)", "Modéré (72–79,9)", "Sévère (≥ 80)"]
activity_by_thi = environment.groupby("THI_cat")["Steps"].agg(["mean", "count"]).reindex(order)

fig, ax = plt.subplots(figsize=(8.8, 4.3))
bars = ax.bar(activity_by_thi.index, activity_by_thi["mean"], color=["#4C78A8", "#72B7B2", "#F2CF5B", "#E45756"])
ax.bar_label(bars, fmt="%.1f", padding=3)
ax.set(title="Activité moyenne observée selon la catégorie de THI", xlabel="Catégorie de THI", ylabel="Pas moyens / 15 min")
fig.tight_layout()
fig.savefig(OUT / "demo_activite_THI.png", dpi=150)
plt.show()

controls = pd.read_csv(DATA / "objective2_thi_controle_complet.csv")
selected_models = [
    "Association globale ajustée pour l'heure",
    "Association globale avec tendance calendaire",
    "Effet intra-jour strict",
]
display(controls[controls["modele"].isin(selected_models)][["modele", "coefficient", "ic95_bas", "ic95_haut", "p", "n"]])
"""
        ),
        _code(
            """
behavior = pd.read_csv(DATA / "summer2019_comportement_environnement.csv")
variables = ["Pct_eating", "Pct_locomotion", "Pct_Idle", "Pct_Explo"]
daily = (
    behavior.groupby("jour")
    .agg(THI=("THI_jour", "first"), **{name: (name, "mean") for name in variables})
    .reset_index()
)
assert len(daily) == 8

results = []
for name in variables:
    valid = daily[["THI", name]].dropna()
    rho, p_value = spearmanr(valid["THI"], valid[name])
    results.append({"comportement": name.replace("Pct_", ""), "jours": len(valid), "rho": rho, "p": p_value})
behavior_results = pd.DataFrame(results)
display(behavior_results.round(3))

fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.3), sharex=True)
axes[0].scatter(daily["THI"], daily["Pct_eating"] * 100, s=64, color="#2E74B5")
for _, row in daily.iterrows():
    axes[0].annotate(pd.to_datetime(row["jour"]).strftime("%d %b"), (row["THI"], row["Pct_eating"] * 100), xytext=(4, 3), textcoords="offset points", fontsize=7)
axes[0].set(title="Alimentation", xlabel="THI journalier", ylabel="Temps observé (%)")

axes[1].scatter(daily["THI"], daily["Pct_locomotion"] * 100, s=58, label="Locomotion", color="#E45756")
axes[1].scatter(daily["THI"], daily["Pct_Explo"] * 100, s=58, label="Exploration", color="#72B7B2")
axes[1].set(title="Locomotion et exploration", xlabel="THI journalier", ylabel="Temps observé (%)")
axes[1].legend(frameon=False)
fig.suptitle("Comportements observés et THI — huit jours indépendants")
fig.tight_layout()
fig.savefig(OUT / "demo_comportement_THI.png", dpi=150)
plt.show()
"""
        ),
        _markdown(
            """
## 4. Guide d'interprétation

### Alertes IceTag
- Une alerte indique une déviation persistante de l'activité individuelle.
- Les niveaux **A** et **B** doivent être vérifiés au niveau de la vache.
- Le niveau **C** est probablement lié à un événement collectif ou de gestion.
- L'IceTag mesure la quantité de mouvement, pas l'asymétrie de la démarche : la
  boiterie légère peut donc rester invisible.

### THI et activité
- Association globale : **+0,221**, `p = 6,15e-8`.
- Après tendance calendaire linéaire : **+0,129**, `p = 0,020`.
- Comparaison intra-jour stricte : **+0,061**, `p = 0,364`, non concluante.
- Conclusion : association exploratoire positive, sans démonstration d'un effet
  thermique intra-jour indépendant.

### Comportements
L'alimentation présente un signal positif (`rho = 0,738`, `p = 0,037`), mais
seulement huit jours indépendants sont disponibles. Ce résultat est suggestif et
doit être confirmé sur davantage de journées.
"""
        ),
        _code(
            """
expected_figures = [
    "demo_profil_journalier.png",
    "demo_pipeline_execution.png",
    "demo_alertes_par_essai.png",
    "demo_activite_THI.png",
    "demo_comportement_THI.png",
]
for filename in expected_figures:
    assert (OUT / filename).exists(), filename

n_bins_fr = f'{pipeline_summary["n_bins"]:,}'.replace(',', ' ')
summary_text = f'''OBJECTIF 3 — NOTEBOOK DE DÉMONSTRATION IoT

Statut : notebook exécuté intégralement et sans erreur.
Données : échantillon IceTag de l'été 2019, résultats multi-saisons et données THI-comportement.
Pipeline : exécution réelle sur la vache 2062, {n_bins_fr} intervalles et {pipeline_summary["lameness_notifs"]} alertes brutes.
Objectif 1 : 385 alertes brutes, requalifiées en 37 A, 195 B et 153 C.
Objectif 2 : association globale THI-activité positive; effet intra-jour +0,061, p = 0,364, non concluant.
Comportement : signal alimentation-THI suggestif sur huit jours indépendants.
Interprétation : dépistage comportemental et analyse exploratoire, sans diagnostic clinique ni causalité démontrée.
'''
(OUT / "synthese_demonstration.txt").write_text(summary_text, encoding="utf-8")
print(summary_text)
print("Toutes les sorties attendues ont été générées.")
"""
        ),
    ]

    notebook = {
        "cells": cells,
        "metadata": {
            "kernelspec": {
                "display_name": "Python 3",
                "language": "python",
                "name": "python3",
            },
            "language_info": {"name": "python", "version": "3.11"},
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    NOTEBOOK.write_text(
        json.dumps(notebook, ensure_ascii=False, indent=1) + "\n", encoding="utf-8"
    )


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Largeur de table invalide: {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _format_run(run, size=11, bold=False, color=INK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _format_run(paragraph.add_run(header), size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _format_run(paragraph.add_run(str(value)), size=9.2)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_run(paragraph.add_run("Page "), size=9, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    _format_run(header.add_run("Projet McGill / WELL-E | Objectif 3"), size=9, color=MUTED)
    _add_page_field(section.footer.paragraphs[0])


def _add_label(doc, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    _format_run(paragraph.add_run(label + " "), bold=True, color=DARK_BLUE)
    _format_run(paragraph.add_run(text))


def _add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    _format_run(paragraph.add_run(text))


def _build_documentation() -> None:
    documentation = """# Documentation du notebook — Objectif 3

## Livrables SOW

- Notebook final : `code/12_objectif3_notebook_demonstration.ipynb`
- Documentation : `RAPPORTS/Objectif3_guide_utilisation.docx` et ce fichier
- Support de présentation additionnel : `RAPPORTS/Objectif3_presentation_detaillee.pptx`

## Exécution

1. Ouvrir le notebook dans Jupyter.
2. Vérifier que le noyau Python possède pandas, numpy, matplotlib, scipy et scikit-learn.
3. Choisir **Restart Kernel and Run All**.
4. Vérifier le message final « Toutes les sorties attendues ont été générées ».

Le notebook détecte automatiquement s'il est exécuté depuis le dépôt ou depuis le paquet
de livraison. Le paquet contient un échantillon IceTag, les tables nécessaires, les modules
du pipeline et les paramètres gelés.

## Résultats attendus

- 8 922 intervalles pour la vache 2062;
- 582 points atypiques Isolation Forest;
- 17 alertes brutes après règles métier;
- 385 alertes brutes dans les quatre saisons, requalifiées en 37 A, 195 B et 153 C;
- association THI-activité globale positive, mais effet intra-jour non concluant;
- huit jours indépendants pour l'analyse comportementale.

## Interprétation

Une alerte signale une anomalie comportementale à vérifier. Elle ne constitue ni un
diagnostic de boiterie ni une preuve causale d'un effet thermique. Les résultats THI et
comportement restent exploratoires.
"""
    DOCUMENTATION.write_text(documentation, encoding="utf-8")

    doc = Document()
    _configure_document(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    _format_run(title.add_run("OBJECTIF 3"), size=23, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _format_run(
        subtitle.add_run("Guide d'utilisation du notebook de démonstration IoT"),
        size=14,
        bold=True,
        color=BLUE,
    )
    _add_label(doc, "Livrables SOW :", "notebook final et documentation")
    _add_label(doc, "Supports inclus :", "données compactes, sorties et présentation détaillée")
    _add_label(doc, "Statut :", "fonctionnel, exécuté intégralement et vérifié")

    doc.add_heading("1. Correspondance avec le SOW", level=1)
    _add_table(
        doc,
        ["Exigence", "Couverture dans le notebook"],
        [
            ["Chargement et prétraitement", "Lecture, schéma canonique, contrôles de colonnes et qualité temporelle"],
            ["Exécution du pipeline IceTag", "Exécution réelle sur 8 922 intervalles de la vache 2062"],
            ["Environnement–comportement", "THI–activité et THI–comportements, avec contrôles statistiques"],
            ["Guide d'interprétation", "Lecture des alertes, limites capteur, confusion temporelle et causalité"],
        ],
        [3300, 6060],
    )

    doc.add_heading("2. Procédure d'exécution", level=1)
    for text in [
        "Ouvrir code/12_objectif3_notebook_demonstration.ipynb dans Jupyter Notebook ou JupyterLab.",
        "Utiliser un environnement Python avec pandas, numpy, matplotlib, scipy et scikit-learn.",
        "Lancer Restart Kernel and Run All afin d'exécuter les cellules dans l'ordre.",
        "Confirmer le message final indiquant que toutes les sorties attendues ont été générées.",
    ]:
        _add_bullet(doc, text)
    _add_label(
        doc,
        "Autonomie :",
        "le notebook reconnaît automatiquement le dépôt de travail ou le paquet livré. "
        "Le paquet inclut les données compactes, le cœur du pipeline et les paramètres gelés.",
    )

    doc.add_heading("3. Résultats de référence", level=1)
    _add_table(
        doc,
        ["Vérification", "Résultat attendu"],
        [
            ["Démonstration pipeline", "8 922 intervalles; 582 points atypiques IF; 17 alertes brutes"],
            ["Quatre saisons", "385 alertes brutes; taux normalisés par 100 vache-jours"],
            ["Requalification", "37 A prioritaires; 195 B à vérifier; 153 C collectives"],
            ["THI–activité", "+0,221 global; +0,129 avec tendance; +0,061 intra-jour, p = 0,364"],
            ["Comportement", "alimentation rho = 0,738, p = 0,037 sur huit jours"],
        ],
        [3000, 6360],
    )

    doc.add_heading("4. Interprétation correcte", level=1)
    for text in [
        "Une alerte est un signal comportemental à vérifier et non un diagnostic clinique.",
        "L'IceTag mesure la quantité d'activité; il ne mesure pas l'asymétrie de démarche nécessaire à la boiterie légère.",
        "L'association THI–activité est positive globalement, mais l'effet intra-jour n'est pas concluant.",
        "Le signal alimentation–THI est suggestif, car huit jours indépendants seulement sont disponibles.",
    ]:
        _add_bullet(doc, text)

    doc.add_heading("5. Contenu du paquet", level=1)
    _add_table(
        doc,
        ["Dossier", "Contenu"],
        [
            ["code/", "notebook, modules du pipeline et paramètres gelés"],
            ["DONNEES_DEMO/", "échantillon IceTag et tables compactes des Objectifs 1 et 2"],
            ["RESULTATS/", "figures et synthèse générées par le notebook"],
            ["DOCUMENTATION/", "documentation textuelle de l'exécution"],
            ["RAPPORTS/", "guide Word et présentation détaillée"],
        ],
        [2400, 6960],
    )

    doc.save(WORD_REPORT)


def _build_package() -> None:
    expected_outputs = [
        "demo_profil_journalier.png",
        "demo_pipeline_execution.png",
        "demo_alertes_par_essai.png",
        "demo_activite_THI.png",
        "demo_comportement_THI.png",
        "synthese_demonstration.txt",
    ]
    missing = [name for name in expected_outputs if not (REPORTS / name).exists()]
    if missing:
        raise FileNotFoundError(f"Sorties du notebook manquantes: {missing}")
    if not PRESENTATION.exists():
        raise FileNotFoundError(f"Présentation manquante: {PRESENTATION}")
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    for folder in ["code/core", "DONNEES_DEMO", "RESULTATS", "DOCUMENTATION", "RAPPORTS"]:
        (PACKAGE / folder).mkdir(parents=True, exist_ok=True)

    shutil.copy2(NOTEBOOK, PACKAGE / "code" / NOTEBOOK.name)
    for source in sorted(CORE_SOURCE.glob("*.py")):
        shutil.copy2(source, PACKAGE / "code" / "core" / source.name)
    shutil.copy2(THRESHOLDS_SOURCE, PACKAGE / "code" / "final_thresholds_v1.json")
    for source in sorted(DATA_DEMO.glob("*.csv")):
        shutil.copy2(source, PACKAGE / "DONNEES_DEMO" / source.name)
    for name in expected_outputs:
        shutil.copy2(REPORTS / name, PACKAGE / "RESULTATS" / name)
    shutil.copy2(DOCUMENTATION, PACKAGE / "DOCUMENTATION" / DOCUMENTATION.name)
    shutil.copy2(WORD_REPORT, PACKAGE / "RAPPORTS" / WORD_REPORT.name)
    shutil.copy2(PRESENTATION, PACKAGE / "RAPPORTS" / PRESENTATION.name)

    readme = """OBJECTIF 3 — NOTEBOOK DE DÉMONSTRATION IoT

À ouvrir en premier : RAPPORTS/Objectif3_guide_utilisation.docx
Support de présentation : RAPPORTS/Objectif3_presentation_detaillee.pptx

Livrables SOW :
- code/12_objectif3_notebook_demonstration.ipynb
- DOCUMENTATION/documentation_notebook.md

Support additionnel :
- RAPPORTS/Objectif3_presentation_detaillee.pptx (12 diapositives avec notes)

Exécution :
1. Ouvrir le notebook dans Jupyter.
2. Choisir Restart Kernel and Run All.
3. Vérifier le message final de réussite.

Le paquet est autonome : données compactes, pipeline et paramètres sont inclus.
Les alertes sont des signaux comportementaux à vérifier, pas des diagnostics cliniques.
"""
    (PACKAGE / "README_livraison_objectif3.txt").write_text(readme, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--package",
        action="store_true",
        help="Construit le paquet après l'exécution réussie du notebook.",
    )
    args = parser.parse_args()
    if args.package:
        _build_package()
        print(f"Paquet Objectif 3 construit : {PACKAGE}")
        return
    _build_demo_data()
    _build_notebook()
    _build_documentation()
    print(f"Notebook préparé : {NOTEBOOK}")
    print(f"Documentation Word : {WORD_REPORT}")


if __name__ == "__main__":
    main()
