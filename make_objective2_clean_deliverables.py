# -*- coding: utf-8 -*-
"""Construit le paquet de livraison propre de l'Objectif 2, sans PDF."""
from __future__ import annotations

from pathlib import Path
import shutil

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT = Path(__file__).resolve().parent
REPORTS = PROJECT / "reports" / "objective2_environnement"
MODELS = REPORTS / "mixed_model"
PACKAGE = (
    Path.home()
    / "Desktop"
    / "Livrables_McGill_WellE"
    / "Objectif2_Environnement_x_comportement"
)

BLUE = "000000"
DARK_BLUE = "000000"
INK = "000000"
MUTED = "595959"
LIGHT_GRAY = "F2F2F2"
CALLOUT = "F2F2F2"
BORDER = "B7B7B7"
WHITE = "FFFFFF"
HEADER_FILL = "000000"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Largeur de table invalide : {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), BORDER)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = (
                    row_index < len(table.rows) - 1
                )


def _format_run(run, size=10.5, bold=False, color=INK, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        _format_run(run, size=9.5, bold=True, color=WHITE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(value))
            _format_run(run, size=9.2)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_label_paragraph(doc, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(f"{label} ")
    _format_run(run, bold=True, color=DARK_BLUE)
    _format_run(paragraph.add_run(text))


def _add_callout(doc, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT)
    p_pr.append(shading)
    run = paragraph.add_run(f"{label} ")
    _format_run(run, bold=True, color=DARK_BLUE)
    _format_run(paragraph.add_run(text))


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_run(paragraph.add_run("Page "), size=9, color=MUTED)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _format_run(
        header.add_run("Projet McGill / WELL-E | Objectif 2"),
        size=9,
        color=MUTED,
    )
    _add_page_field(section.footer.paragraphs[0])


def _format_p(value: float) -> str:
    if value < 0.001:
        return f"{value:.2e}".replace("e-0", "e-").replace(".", ",")
    return f"{value:.3f}".replace(".", ",")


def _format_num(value: float, digits=3) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def _format_int(value: float) -> str:
    return f"{int(value):,}".replace(",", " ")


def _build_report(destination: Path) -> None:
    models = pd.read_csv(MODELS / "objective2_mixed_model_summary.csv")
    complete = pd.read_csv(MODELS / "objective2_thi_controle_complet.csv")
    nonlinear = pd.read_csv(MODELS / "objective2_non_linearite.csv")
    descriptifs = pd.read_csv(MODELS / "objective2_descriptifs.csv").iloc[0]
    trimodal = dict(
        pd.read_csv(MODELS / "objective2_trimodal_summary.csv").values.tolist()
    )
    behavior = pd.read_csv(REPORTS / "summer2019_comportement_vs_THI.csv")

    mixed = models[models["modele"].str.startswith("Modèle mixte")].iloc[0]
    herd_hour = models[
        models["modele"].eq("Troupeau-timestamp : Steps ~ THI + heure")
    ].iloc[0]
    herd_day = models[
        models["modele"].eq("Troupeau-timestamp : Steps ~ THI + heure + jour")
    ].iloc[0]
    mi_day = models[
        models["modele"].str.startswith("Troupeau-timestamp : Motion Index")
    ].iloc[0]
    trend = complete[
        complete["modele"].eq("Association globale avec tendance calendaire")
    ].iloc[0]
    daily = complete[
        complete["modele"].eq("Association entre journées (unité jour)")
    ].iloc[0]
    daily_trend = complete[
        complete["modele"].eq("Association entre jours avec tendance calendaire")
    ].iloc[0]
    eating = behavior[behavior["comportement"].eq("eating")].iloc[0]

    doc = Document()
    _configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(4)
    _format_run(title.add_run("OBJECTIF 2"), size=23, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _format_run(
        subtitle.add_run("Conditions environnementales et comportement locomoteur"),
        size=14,
        bold=True,
        color=BLUE,
    )
    _add_label_paragraph(doc, "Projet :", "McGill / WELL-E, Summer 2019")
    _add_label_paragraph(doc, "Statut :", "Rapport exploratoire corrigé et paquet de livraison")
    _add_callout(
        doc,
        "Conclusion principale.",
        "L'association globale THI-activité est positive et demeure positive après "
        "contrôle d'une tendance calendaire linéaire. L'effet strictement intra-jour "
        "est plus faible et non concluant. Les résultats décrivent donc une association "
        "exploratoire, sans preuve d'un effet thermique causal indépendant.",
    )

    doc.add_heading("1. Correspondance avec le SOW", level=1)
    _add_table(
        doc,
        ["Livrable SOW", "Fichier livré"],
        [
            [
                "2.1 Jeu intégré : accéléromètres + environnement + comportements",
                "DONNEES_SYNCHRONISEES/summer2019_multimodal_cow_day.csv",
            ],
            [
                "2.1 Documentation de la synchronisation",
                "NOTES_SOW/documentation_synchronisation.md",
            ],
            [
                "2.2 Rapport d'analyse exploratoire",
                "RAPPORTS/Objectif2_rapport_livraison.docx",
            ],
            [
                "2.2 Faisabilité d'une modélisation intégrée",
                "NOTES_SOW/notes_faisabilite_modelisation.md",
            ],
        ],
        [4100, 5260],
    )

    doc.add_heading("2. Synchronisation des données", level=1)
    _add_label_paragraph(
        doc,
        "Activité et environnement.",
        f"{_format_int(descriptifs['n_bins_vache'])} intervalles vache-15 min, "
        f"{int(descriptifs['n_vaches'])} vaches et {int(descriptifs['n_jours'])} jours. "
        "Les mesures IceTag sont appariées aux sondes HOBO extérieures, agrégées "
        "au même pas de 15 minutes.",
    )
    _add_label_paragraph(
        doc,
        "Périmètre de l'analyse.",
        "Le corpus Summer 2019 traité à l'Objectif 1 compte 139 111 intervalles et "
        "18 vaches (5 juin au 6 septembre 2019). Les sondes environnementales HOBO ne "
        "couvrent la période qu'à partir du 1er juillet 2019 : l'analyse "
        "environnement-activité porte donc sur le sous-ensemble apparié de "
        f"{_format_int(descriptifs['n_bins_vache'])} intervalles (62,9 % du corpus), "
        f"{int(descriptifs['n_vaches'])} vaches et {int(descriptifs['n_jours'])} jours "
        "(1er juillet au 6 septembre). La vache 2067 est absente de ce sous-ensemble, "
        "faute de mesure environnementale concomitante. Cet écart avec l'Objectif 1 est "
        "attendu et ne traduit pas une perte de données.",
    )
    _add_label_paragraph(
        doc,
        "Comportements.",
        f"{int(trimodal['scans_comportementaux_total'])} scans sur "
        f"{int(trimodal['jours_de_scan'])} jours, pour "
        f"{int(trimodal['vaches_observees'])} vaches, reliés aux conditions du jour.",
    )
    _add_label_paragraph(
        doc,
        "Table trimodale.",
        f"{int(trimodal['scans_trimodaux_complets'])} scans sur "
        f"{int(trimodal['scans_comportementaux_total'])} disposent simultanément "
        "du comportement, de l'environnement et de l'activité quotidienne. Deux "
        "scans de la vache 5169 restent documentés avec activité IceTag absente.",
    )
    _add_table(
        doc,
        ["Fichier", "Niveau", "Rôle"],
        [
            [
                "summer2019_icetag_environnement_15min.csv",
                "vache-15 min",
                "Activité IceTag + HOBO extérieur",
            ],
            [
                "summer2019_comportement_environnement.csv",
                "vache-jour-scan",
                "Comportement + conditions journalières",
            ],
            [
                "summer2019_multimodal_cow_day.csv",
                "vache-jour-scan",
                "Livrable trimodal avec statut d'intégration",
            ],
        ],
        [3900, 1900, 3560],
    )

    doc.add_heading("3. Activité locomotrice et THI", level=1)
    _add_label_paragraph(
        doc,
        "Plage observée.",
        f"THI de {_format_num(descriptifs['THI_min'], 1)} à "
        f"{_format_num(descriptifs['THI_max'], 1)}, moyenne "
        f"{_format_num(descriptifs['THI_moy'], 1)}. Seulement "
        f"{_format_num(descriptifs['pct_THI_ge_80'], 1)} % des intervalles "
        "atteignent THI ≥ 80.",
    )
    model_rows = []
    for row in [mixed, herd_hour, herd_day, mi_day]:
        ci = f"[{_format_num(row['ic95_bas'])} ; {_format_num(row['ic95_haut'])}]"
        model_rows.append(
            [
                row["modele"].replace("Troupeau-timestamp : ", ""),
                _format_num(row["coef_THI"]),
                ci,
                _format_p(row["p"]),
            ]
        )
    _add_table(
        doc,
        ["Analyse", "Effet THI", "IC 95 %", "p"],
        model_rows,
        [4300, 1450, 2250, 1360],
    )
    _add_callout(
        doc,
        "Interprétation.",
        f"Sans contrôle du jour, l'analyse troupeau-timestamp estime un effet de "
        f"+{_format_num(herd_hour['coef_THI'])} pas par unité THI "
        f"(p = {_format_p(herd_hour['p'])}). Avec contrôle du jour, l'effet est "
        f"+{_format_num(herd_day['coef_THI'])} "
        f"(IC 95 % [{_format_num(herd_day['ic95_bas'])} ; "
        f"{_format_num(herd_day['ic95_haut'])}], p = {_format_p(herd_day['p'])}). "
        "Le signal observé dépend donc fortement des différences entre jours.",
    )
    _add_label_paragraph(
        doc,
        "Contrôles complémentaires.",
        f"Avec une tendance calendaire linéaire, l'association globale reste positive "
        f"(+{_format_num(trend['coefficient'])}, p = {_format_p(trend['p'])}). "
        f"Au niveau des {int(daily['n'])} unités journalières, elle est estimée à "
        f"+{_format_num(daily['coefficient'])} (p = {_format_p(daily['p'])}) et reste "
        f"positive avec la tendance calendaire (+{_format_num(daily_trend['coefficient'])}, "
        f"p = {_format_p(daily_trend['p'])}). Ces analyses renforcent l'existence "
        "d'une association entre jours, sans établir une causalité thermique.",
    )
    quadratic = nonlinear[nonlinear["terme"].str.contains("carré")].iloc[0]
    _add_label_paragraph(
        doc,
        "Non-linéarité.",
        f"Après contrôle de l'heure et du jour, le terme quadratique n'est pas "
        f"concluant (p = {_format_p(quadratic['p'])}). Les données ne permettent "
        "pas d'établir un seuil thermique ou une courbe convexe indépendante du jour.",
    )

    doc.add_heading("4. Comportements observés et THI", level=1)
    _add_label_paragraph(
        doc,
        "Résultat exploratoire.",
        f"Au niveau du jour, l'alimentation présente rho = "
        f"{_format_num(eating['rho_vs_THI'])}, p = {_format_p(eating['p'])}, "
        f"sur seulement {int(eating['n_jours'])} jours. Les autres catégories "
        "analysées ne sont pas significatives.",
    )
    _add_label_paragraph(
        doc,
        "Portée.",
        "Ce résultat est suggestif, pas confirmatoire. Les 51 scans ne sont pas "
        "51 expositions thermiques indépendantes, puisque plusieurs vaches partagent "
        "le même THI journalier.",
    )

    doc.add_heading("5. Faisabilité d'un modèle intégré", level=1)
    _add_label_paragraph(
        doc,
        "Faisable maintenant.",
        "Modéliser l'activité IceTag en fonction du THI, avec heure et jour, tout en "
        "présentant séparément les associations entre jours et les effets intra-jour.",
    )
    _add_label_paragraph(
        doc,
        "Non robuste en l'état.",
        "Introduire les comportements observés comme covariables principales : huit "
        "jours de scan sont insuffisants pour une estimation stable.",
    )
    _add_label_paragraph(
        doc,
        "Renforcement futur.",
        "Ajouter davantage de jours de scan, tester des effets décalés et étendre "
        "l'analyse au froid hivernal après validation de la correspondance entre "
        "les éthogrammes Summer et Winter par McGill.",
    )

    doc.add_heading("6. Conclusion et limites", level=1)
    _add_label_paragraph(
        doc,
        "Conclusion.",
        "La synchronisation demandée par le SOW est réalisée et documentée. "
        "L'association globale THI-activité est positive et persiste après contrôle "
        "d'une tendance calendaire simple. Elle est principalement portée par les "
        "différences entre jours et n'est pas confirmée dans l'analyse intra-jour. "
        "Aucune relation causale ni baisse d'activité liée au stress thermique sévère "
        "ne peut être conclue.",
    )
    _add_label_paragraph(
        doc,
        "Limite thermique.",
        "Le stress sévère est rare dans le corpus; les conclusions restent limitées "
        "à la plage de THI observée au Québec pendant Summer 2019.",
    )
    _add_label_paragraph(
        doc,
        "Confondants.",
        "La progression saisonnière, l'accès à l'exercice, le jour d'observation et "
        "d'autres variables individuelles peuvent expliquer une partie du signal.",
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def _write_tables_readme(destination: Path) -> None:
    destination.write_text(
        """Guide de lecture des tableaux - Objectif 2

Résultat à retenir
  Les conditions de THI plus élevées sont associées à une activité locomotrice
  plus élevée dans Summer 2019. L'association globale est positive (+0,221;
  p = 6,15e-08) et reste positive après contrôle d'une tendance calendaire
  linéaire (+0,129; p = 0,020). L'effet strictement intra-jour est plus faible
  et non concluant (+0,061; p = 0,364). Ces résultats décrivent une association,
  sans démontrer un effet thermique causal indépendant.

objective2_thi_controle_complet.csv
  TABLEAU DE SYNTHÈSE. Il distingue explicitement :
  - l'association globale ajustée pour l'heure;
  - l'association entre les unités journalières;
  - les analyses avec tendance calendaire;
  - l'effet intra-jour avec contrôle complet du jour.
  Les estimations +0,221 et +0,061 répondent à deux questions différentes;
  elles ne se contredisent pas.

objective2_mixed_model_summary.csv
  Analyse de sensibilité à plusieurs niveaux : lignes vache-15 min,
  troupeau-timestamp et vache-jour. Le modèle mixte par vache sans contrôle
  du jour ne doit pas être utilisé seul pour conclure, car le THI est partagé
  par les vaches présentes au même timestamp.

objective2_descriptifs.csv
  Couverture du sous-ensemble apparié : intervalles, timestamps, vaches, jours
  et plage de THI.

objective2_non_linearite.csv
  Terme quadratique du THI après contrôle de l'heure et du jour : non concluant
  (p = 0,224). Aucun seuil thermique indépendant du jour n'est établi.

objective2_profil_par_thi.csv
  Tableau descriptif, sans contrôle du jour ni de l'heure. L'activité moyenne
  augmente avec la tranche de THI (6,3 à 12,8 pas par 15 minutes), mais ce profil
  ne constitue pas à lui seul une preuve causale.

objective2_trimodal_summary.csv
  Couverture de la table trimodale : 49 scans complets sur 51, 8 vaches et
  8 jours de scans comportementaux.

summer2019_comportement_vs_THI.csv
  Relations exploratoires entre comportements journaliers et THI. Le signal
  positif de l'alimentation repose sur seulement 8 jours et doit être confirmé.
""",
        encoding="utf-8",
    )


def _write_notes() -> tuple[Path, Path, Path]:
    sync = REPORTS / "documentation_synchronisation.md"
    feasibility = REPORTS / "notes_faisabilite_modelisation.md"
    readme = REPORTS / "README_livraison_objectif2.txt"

    sync.write_text(
        """# Documentation de la synchronisation - Objectif 2, Tâche 2.1

## Sources et résolution

- Activité locomotrice : accéléromètres IceTag, agrégés en intervalles de 15 minutes.
- Environnement : sondes HOBO **extérieures** (`Outside`), température et humidité.
- Comportements : scans d'observation datés, interprétés au niveau du jour.

Les deux sondes HOBO extérieures disponibles pour chaque période sont concaténées puis agrégées par intervalle de 15 minutes. Le THI est calculé à partir de la température et de l'humidité. L'environnement est commun au troupeau et est joint à chaque vache par timestamp.

## Tables livrées

1. `summer2019_icetag_environnement_15min.csv` : activité + environnement, au niveau vache-15 minutes.
2. `summer2019_comportement_environnement.csv` : comportements + environnement journalier.
3. `summer2019_multimodal_cow_day.csv` : table trimodale au niveau vache-jour-scan.

La table trimodale conserve les 51 scans. Quarante-neuf disposent des trois modalités. Deux scans de la vache 5169, datés du 16 et du 23 août 2019, sont conservés avec le statut `activité IceTag absente pour cette vache et ce jour`.

## Contrôles de qualité

- 87 501 intervalles vache-15 minutes, 17 vaches et 62 jours.
- 51 scans comportementaux, 8 vaches et 8 jours.
- 49 scans trimodaux complets; 2 scans incomplets documentés.
- Le comportement n'a pas une résolution de 15 minutes : aucune interpolation artificielle n'est effectuée.
- La période du 2 au 9 août ne comporte pas de fichiers HOBO `Outside`; seules les périodes réellement disponibles sont utilisées.

## Périmètre de l'analyse (écart avec l'Objectif 1)

Le corpus Summer 2019 traité à l'Objectif 1 compte **139 111 intervalles et 18 vaches** (5 juin au 6 septembre 2019).
Les sondes environnementales HOBO ne couvrent la période qu'**à partir du 1er juillet 2019**. L'analyse
environnement-activité porte donc sur le **sous-ensemble apparié** :

| | Objectif 1 (corpus complet) | Objectif 2 (apparié environnement) |
|---|---|---|
| Intervalles 15 min | 139 111 | 87 501 (62,9 %) |
| Vaches | 18 | 17 |
| Période | 5 juin - 6 sept 2019 | 1er juillet - 6 sept 2019 |

La vache 2067 est absente du sous-ensemble, faute de mesure environnementale concomitante.
Cet écart est **attendu** et ne traduit pas une perte de données : il reflète la couverture des sondes HOBO.
""",
        encoding="utf-8",
    )
    feasibility.write_text(
        """# Faisabilité d'un modèle intégré - Objectif 2, Tâche 2.2

## Verdict

Un modèle activité-environnement est faisable. Un modèle incluant les comportements observés comme covariables principales n'est pas robuste avec seulement huit jours de scan.

## Résultats de sensibilité

- Modèle mixte avec vache et heure : effet THI positif, +0,198 pas par unité THI.
- Analyse du troupeau par timestamp avec contrôle de l'heure : +0,221, p = 6,15e-08.
- Même analyse avec contrôle explicite du jour : +0,061, IC 95 % [-0,070; 0,192], p = 0,364.
- Motion Index avec heure et jour : +0,091, p = 0,752.
- Le terme quadratique du THI n'est plus concluant après contrôle du jour, p = 0,224.

## Contrôle complet de l'effet THI

Les estimations `+0,221` et `+0,061` répondent à deux questions différentes :

- **Association globale** : `+0,221` pas par unité THI, p = 6,15e-08. Elle combine les variations entre jours et les variations à l'intérieur des jours.
- **Association au niveau de 62 unités journalières** : `+0,273`, p = 9,20e-08, avec des erreurs HAC tenant compte de leur succession temporelle. Avec une tendance calendaire linéaire, l'estimation reste positive (`+0,162`, p = 0,030).
- **Effet intra-jour strict** : `+0,061`, IC 95 % [-0,070; 0,192], p = 0,364. Avec des erreurs regroupées par jour, la conclusion demeure non concluante (p = 0,291).

Les contrôles d'intégrité confirment 5 795 timestamps uniques, aucun doublon vache-timestamp et une valeur THI commune aux vaches présentes au même timestamp. L'association globale reste positive lorsque l'analyse est limitée aux jours ayant au moins 80 timestamps ou aux timestamps comptant au moins 14 vaches.

L'association positive globale est donc reproductible et ne résulte pas de la simple duplication des vaches. Elle est principalement portée par les différences entre jours. Le contrôle par jour retire toute cette composante et pose une question plus stricte : à heure et journée identiques, la variation résiduelle du THI explique-t-elle l'activité? La réponse actuelle est non concluante.

La progression saisonnière, le protocole d'exercice ou d'autres caractéristiques journalières peuvent encore expliquer une partie de l'association entre jours. Les données montrent une association positive, mais elles ne démontrent pas un effet thermique causal indépendant.

## Recommandations

1. Conserver le contrôle de l'heure et du jour dans l'analyse principale.
2. Présenter séparément les associations entre jours et les variations intra-jour.
3. Densifier les scans comportementaux avant de les utiliser comme covariables.
4. Tester les effets décalés lorsque davantage de jours indépendants seront disponibles.
5. Étendre au froid hivernal après validation McGill de la correspondance des éthogrammes.
""",
        encoding="utf-8",
    )
    readme.write_text(
        """Objectif 2 - Conditions environnementales et comportement locomoteur

Contenu à livrer :
- RAPPORTS/Objectif2_rapport_livraison.docx
- RAPPORTS/Objectif2_presentation_detaillee.pptx
- DONNEES_SYNCHRONISEES/ : trois tables synchronisées et le relevé des deux scans incomplets
- TABLEAUX_CSV/ : descriptifs, modèles de sensibilité, contrôle complet du THI, non-linéarité et concordance trimodale
- FIGURES/ : deux visualisations exploratoires
- NOTES_SOW/ : procédure de synchronisation et faisabilité de la modélisation
- code/ : notebook de synchronisation actuel et scripts finaux

Résultat principal : l'association globale THI-activité est positive et reproductible, principalement entre les jours. L'effet strictement intra-jour est positif mais non concluant. Les deux estimations répondent à des questions différentes et aucune causalité n'est revendiquée.
""",
        encoding="utf-8",
    )
    return sync, feasibility, readme


def _copy(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)


def build_package() -> None:
    if PACKAGE.name != "Objectif2_Environnement_x_comportement":
        raise RuntimeError("Chemin de livraison inattendu")
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    for folder in [
        "RAPPORTS",
        "DONNEES_SYNCHRONISEES",
        "TABLEAUX_CSV",
        "FIGURES",
        "NOTES_SOW",
        "code",
    ]:
        (PACKAGE / folder).mkdir(parents=True, exist_ok=True)

    sync, feasibility, readme = _write_notes()
    report_source = REPORTS / "Objectif2_rapport_livraison.docx"
    presentation_source = REPORTS / "Objectif2_presentation_detaillee.pptx"
    _build_report(report_source)

    _copy(report_source, PACKAGE / "RAPPORTS" / report_source.name)
    _copy(
        presentation_source,
        PACKAGE / "RAPPORTS" / presentation_source.name,
    )
    _copy(sync, PACKAGE / "NOTES_SOW" / sync.name)
    _copy(feasibility, PACKAGE / "NOTES_SOW" / feasibility.name)
    _copy(readme, PACKAGE / readme.name)

    for name in [
        "summer2019_icetag_environnement_15min.csv",
        "summer2019_comportement_environnement.csv",
        "summer2019_multimodal_cow_day.csv",
        "summer2019_multimodal_unmatched_scans.csv",
    ]:
        _copy(REPORTS / name, PACKAGE / "DONNEES_SYNCHRONISEES" / name)

    for name in [
        "objective2_descriptifs.csv",
        "objective2_mixed_model_summary.csv",
        "objective2_thi_controle_complet.csv",
        "objective2_non_linearite.csv",
        "objective2_profil_par_thi.csv",
        "objective2_trimodal_summary.csv",
    ]:
        _copy(MODELS / name, PACKAGE / "TABLEAUX_CSV" / name)
    _copy(
        REPORTS / "summer2019_comportement_vs_THI.csv",
        PACKAGE / "TABLEAUX_CSV" / "summer2019_comportement_vs_THI.csv",
    )
    _write_tables_readme(PACKAGE / "TABLEAUX_CSV" / "README_tableaux.txt")

    for name in [
        "summer2019_environnement_activite.png",
        "summer2019_env_comportement_v2.png",
    ]:
        _copy(REPORTS / name, PACKAGE / "FIGURES" / name)

    _copy(
        PROJECT / "notebooks" / "11_objectif2_synchro_environnement_comportement.ipynb",
        PACKAGE / "code" / "11_objectif2_synchro_environnement_comportement.ipynb",
    )
    for name in [
        "build_objective2_trimodal_dataset.py",
        "run_objective2_mixed_model.py",
    ]:
        _copy(PROJECT / name, PACKAGE / "code" / name)

    (PACKAGE / "code" / "README_execution.txt").write_text(
        """Exécution depuis le paquet de livraison

1. `python build_objective2_trimodal_dataset.py`
   Reconstruit la table trimodale à partir des deux tables synchronisées livrées.

2. `python run_objective2_mixed_model.py`
   Recalcule les descriptifs et les analyses de sensibilité dans TABLEAUX_CSV/.

Le notebook 11 documente l'extraction depuis les données brutes et doit être exécuté dans le projet McGill complet, où les fichiers IceTag et HOBO sont disponibles.
""",
        encoding="utf-8",
    )

    print("Paquet Objectif 2 construit :", PACKAGE)


if __name__ == "__main__":
    build_package()
