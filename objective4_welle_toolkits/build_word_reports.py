"""Génère les trois livrables Word de l'objectif 4 depuis les sources Markdown."""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
REPORTS = ROOT / 'reports'
BLUE = '005A84'
DARK_BLUE = '163A59'
LIGHT_BLUE = 'EAF3F8'
LIGHT_GRAY = 'F2F4F7'
MID_GRAY = '667085'
WHITE = 'FFFFFF'
BLACK = '111111'


def set_run_font(run, name='Calibri', size=None, color=BLACK,
                 bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=70, start=120, bottom=70, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start),
                          ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')

    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn('w:w'), str(widths_dxa[index]))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    set_run_font(run, size=9, color=MID_GRAY)
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(field)
    run = paragraph.add_run('  |  McGill - WELL-E')
    set_run_font(run, size=9, color=MID_GRAY)


def configure_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_specs = {
        'Heading 1': (16, BLUE, 16, 8),
        'Heading 2': (13, BLUE, 12, 6),
        'Heading 3': (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ('List Bullet', 'List Number'):
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(running_title)
    set_run_font(run, size=9, color=MID_GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_title(doc, text, subtitle):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run('OBJECTIF 4  |  SUPPORT AUX OUTILS WELL-E')
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=23, color=BLACK, bold=True)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=12.5, color=MID_GRAY)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def add_inline(paragraph, text, size=11, color=BLACK):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name='Courier New', size=size - 0.5,
                         color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), LIGHT_BLUE)
    p_pr.append(shading)
    add_inline(paragraph, text)
    return paragraph


def add_markdown_table(doc, rows):
    headers = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for index, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], BLUE)
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline(paragraph, value, size=9.5, color=WHITE)
        for run in paragraph.runs:
            run.bold = True

    for row_index, values in enumerate(body):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[index], LIGHT_GRAY)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value, size=9.25)

    if len(headers) == 2:
        widths = [3600, 5760]
    elif len(headers) == 3:
        widths = [2500, 3430, 3430]
    else:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_markdown(doc, path):
    lines = path.read_text(encoding='utf-8').splitlines()
    title_seen = False
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.15)
                paragraph.paragraph_format.right_indent = Inches(0.15)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(10)
                p_pr = paragraph._p.get_or_add_pPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), LIGHT_GRAY)
                p_pr.append(shading)
                run = paragraph.add_run('\n'.join(code_lines))
                set_run_font(run, name='Courier New', size=8.5, color=BLACK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped == '<!-- PAGEBREAK -->':
            doc.add_page_break()
            index += 1
            continue

        if stripped.startswith('|') and index + 1 < len(lines) and re.match(
                r'^\|(?:\s*:?-+:?\s*\|)+$', lines[index + 1].strip()):
            rows = []
            header = [cell.strip() for cell in stripped.strip('|').split('|')]
            rows.append(header)
            index += 2
            while index < len(lines) and lines[index].strip().startswith('|'):
                rows.append([
                    cell.strip()
                    for cell in lines[index].strip().strip('|').split('|')
                ])
                index += 1
            add_markdown_table(doc, rows)
            continue

        if stripped.startswith('# '):
            title_seen = True
            index += 1
            continue
        if stripped.startswith('## '):
            doc.add_paragraph(stripped[3:], style='Heading 1')
        elif stripped.startswith('### '):
            doc.add_paragraph(stripped[4:], style='Heading 2')
        elif re.match(r'^\d+\.\s+', stripped):
            paragraph = doc.add_paragraph(style='List Number')
            add_inline(paragraph, re.sub(r'^\d+\.\s+', '', stripped))
        elif stripped.startswith('- '):
            paragraph = doc.add_paragraph(style='List Bullet')
            add_inline(paragraph, stripped[2:])
        elif stripped:
            if stripped.startswith('**Statut') or stripped.startswith('**Verdict'):
                add_callout(doc, stripped)
            else:
                paragraph = doc.add_paragraph()
                add_inline(paragraph, stripped)
        index += 1
    if not title_seen:
        raise ValueError(f'Titre Markdown absent : {path}')


def build(source_name, output_name, title, subtitle, running_title):
    doc = Document()
    configure_document(doc, running_title)
    add_title(doc, title, subtitle)
    parse_markdown(doc, REPORTS / source_name)
    properties = doc.core_properties
    properties.title = title
    properties.subject = 'Objectif 4 - Support aux outils WELL-E'
    properties.author = 'Projet McGill WELL-E'
    properties.keywords = 'WELL-E, MATLAB, Python, audit, MATRID'
    output = REPORTS / output_name
    doc.save(output)
    print(output)


def main():
    build(
        'revue_annotee_scripts.md',
        'Objectif4_revue_annotee_scripts.docx',
        'Revue annotée de la conversion MATLAB vers Python',
        'Vérification, corrections et validation sur l’exemple du 27 mars',
        'Objectif 4.1 | Revue de code',
    )
    build(
        'note_application_Mira.md',
        'Objectif4_note_application_Mira.docx',
        'Application des outils WELL-E à Mira',
        'Procédure, prérequis et critères d’acceptation',
        'Objectif 4.2 | Application à Mira',
    )
    build(
        'guide_depannage_FAQ.md',
        'Objectif4_guide_depannage_FAQ.docx',
        'Guide de dépannage WELL-E',
        'Questions fréquentes pour la conversion et le calcul des métriques',
        'Objectif 4.2 | Dépannage WELL-E',
    )


if __name__ == '__main__':
    main()
