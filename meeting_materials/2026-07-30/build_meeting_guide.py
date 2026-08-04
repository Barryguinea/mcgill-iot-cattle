from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

def _find_skill_root() -> Path:
    """Localise la skill documents sans figer le numero de version du cache."""
    base = Path(
        "/Users/alioubarry/.codex/plugins/cache/openai-primary-runtime/documents"
    )
    candidates = sorted(base.glob("*/skills/documents"), reverse=True)
    for candidate in candidates:
        if (candidate / "scripts" / "table_geometry.py").is_file():
            return candidate
    raise RuntimeError(
        "table_geometry.py introuvable sous "
        f"{base}. Verifier l'installation de la skill documents."
    )


SKILL_ROOT = _find_skill_root()
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


OUT = Path(
    "/Users/alioubarry/Desktop/Livrables_McGill_WellE/"
    "Reunion_McGill_2026-07-30/Guide_oral_et_QR_McGill_30_juillet_2026.docx"
)

BLUE = "2474E5"
TEAL = "14877E"
INK = "111111"
MUTED = "5E6875"
LIGHT = "F2F4F7"
BLUE_LIGHT = "E9F0FB"
GREEN_LIGHT = "EAF4EF"
YELLOW_LIGHT = "FFF6D8"
RED_LIGHT = "FCEDED"
RED = "9B1C1C"
GREEN = "2D8A57"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D6DADF", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = tc_borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: str = INK,
    name: str = "Calibri",
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=27, bold=True, color=INK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=14, color=MUTED)


def add_kicker(doc: Document, text: str, color: str = BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run_font(r, size=9.5, bold=True, color=color)


def add_body(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    italic: bool = False,
    color: str = INK,
    after: float = 6,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)


def add_bullet(doc: Document, text: str, level: int = 0, color: str = INK) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f" {text}")
    set_run_font(r, color=color)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)


def add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = BLUE_LIGHT,
    accent: str = BLUE,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.line_spacing = 1.20
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    weights: list[float],
    *,
    header_fill: str = INK,
    font_size: float = 9.5,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color=header_fill, size=4)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        for c_idx, text in enumerate(values):
            cell = row.cells[c_idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, LIGHT if r_idx % 2 else WHITE)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            rr = p.add_run(str(text))
            set_run_font(rr, size=font_size, color=INK)
    widths = column_widths_from_weights(weights, 9360)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_question(doc: Document, question: str, answer: str, details: list[str] | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(question)
    set_run_font(r, size=11.5, bold=True, color=BLUE)
    add_callout(doc, "Réponse courte", answer, fill=LIGHT, accent=TEAL)
    if details:
        for item in details:
            add_bullet(doc, item)


def add_slide_script(doc: Document, number: int, title: str, script: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{number:02d}  ")
    set_run_font(r1, size=10.5, bold=True, color=BLUE)
    r2 = p.add_run(title)
    set_run_font(r2, size=10.5, bold=True, color=INK)
    add_body(doc, script, after=4, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style in ["List Bullet", "List Bullet 2", "List Number"]:
    style = styles[list_style]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25
styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
styles["List Number"].paragraph_format.left_indent = Inches(0.375)
styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("McGill / WELL-E  |  Réunion du 30 juillet 2026")
set_run_font(hr, size=9, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fr = fp.add_run("Objectif 1 — guide oral et questions-réponses")
set_run_font(fr, size=9, color=MUTED)
add_page_number(footer.add_paragraph())

# Cover / workshop-agenda pattern.
add_kicker(doc, "Préparation de réunion", BLUE)
add_title(
    doc,
    "Objectif 1 — Guide oral et questions-réponses",
    "Travail réalisé, résultats, dossiers, valeur scientifique et prochaines décisions",
)
add_body(doc, "Réunion McGill / WELL-E — jeudi 30 juillet 2026", bold_lead="Réunion")
add_body(doc, "Préparé pour Aliou Barry", bold_lead="Préparé")
doc.add_paragraph().paragraph_format.space_after = Pt(10)
add_table(
    doc,
    ["Durée cible", "Résultat principal", "Verdict", "Demande finale"],
    [
        [
            "12–15 min + questions",
            "4 saisons, 375 031 bins, 385 notifications",
            "SOW atteint; validation clinique limitée",
            "Accepter l’objectif 1 et choisir la cohorte de validation",
        ]
    ],
    [1.15, 2.1, 1.65, 2.0],
    header_fill=INK,
    font_size=9,
)
add_callout(
    doc,
    "Positionnement",
    "Soyez convaincant par la précision. Le projet est positif parce que le transfert est complet, reproductible et évalué honnêtement; il ne faut pas le présenter comme un diagnostic clinique déjà validé.",
    fill=BLUE_LIGHT,
    accent=BLUE,
)
doc.add_heading("Les trois messages à faire retenir", level=1)
add_numbered(doc, "Le SOW de l’objectif 1 est rempli : quatre saisons traitées et quatre livrables techniques présents.")
add_numbered(doc, "Le principal gain scientifique est la qualification des alertes par contexte troupeau : 37 A, 195 B et 153 C.")
add_numbered(doc, "La prochaine valeur à créer est une validation prospective avec davantage de scores locomoteurs synchronisés.")
doc.add_page_break()

# Quick reference.
add_kicker(doc, "Antisèche", TEAL)
add_title(doc, "Les chiffres à connaître sans hésiter")
add_table(
    doc,
    ["Indicateur", "Valeur", "Ce que cela signifie"],
    [
        ["Saisons", "4 / 4", "Winter 2019, Summer 2019, Fall 2019, Fall 2021"],
        ["Intervalles", "375 031", "Bins vache × 15 min; pas des cas indépendants de boiterie"],
        ["Notifications brutes", "385", "Épisodes après persistance et cooldown, pas diagnostics"],
        ["Priorité A", "37", "Signaux individuels à examiner en premier"],
        ["Priorité B", "195", "Signaux individuels plausibles à contextualiser"],
        ["Contexte C", "153", "Événements collectifs probables"],
        ["Taux global", "9,855 / 100 vache-jours", "Comparaison ajustée à l’exposition"],
        ["SLS baseline", "n=16; p=0,649; ρ=0,033", "Aucune association observée pour IF + règles"],
        ["SLS hybride", "n=14; 3 positifs; AUC=0,924; p=0,031", "Prometteur, mais exploratoire et instable"],
    ],
    [1.55, 1.8, 3.65],
    header_fill=INK,
    font_size=9.2,
)
doc.add_heading("Résultats par saison", level=2)
add_table(
    doc,
    ["Saison", "Profils", "Bins", "Total", "A", "B", "C", "Couverture"],
    [
        ["Winter 2019", "17", "136 929", "149", "14", "46", "89", "94,4 %"],
        ["Summer 2019", "18", "139 111", "127", "16", "94", "17", "98,3 %"],
        ["Fall 2019", "30", "93 860", "105", "7", "51", "47", "99,3 %"],
        ["Fall 2021", "10 traités / 8 complets", "5 131", "4", "0", "4", "0", "100,0 %"],
    ],
    [1.2, 1.45, 1.05, 0.65, 0.45, 0.45, 0.45, 0.9],
    header_fill=TEAL,
    font_size=8.5,
)
add_callout(
    doc,
    "Phrase à retenir",
    "« 385 est une sortie brute multi-saison. Après contexte troupeau, seulement 37 alertes sont prioritaires et 153 sont probablement liées à un événement collectif. »",
    fill=YELLOW_LIGHT,
    accent="7A5A00",
)
doc.add_page_break()

# Opening and agenda.
add_kicker(doc, "Prise de parole", BLUE)
add_title(doc, "Ouverture de 90 secondes")
add_callout(
    doc,
    "Texte proposé",
    "« L’objectif 1 consistait à appliquer et évaluer une pipeline de détection sur quatre corpus IceTag qui n’étaient pas directement compatibles entre eux. Mon travail a donc comporté cinq étapes : inventorier les sources, harmoniser les colonnes et les timestamps, exécuter la même pipeline avec des paramètres gelés, produire les sorties par intervalle et par alerte, puis évaluer leur sens grâce aux scans, aux scores SLS disponibles et au contexte du troupeau. Le SOW est rempli : les quatre saisons sont traitées, les quatre livrables sont présents et chaque résultat est traçable. La conclusion scientifique est positive mais précise : nous avons un système d’alertes comportementales à vérifier, renforcé par une priorisation A/B/C; nous n’avons pas encore une validation diagnostique clinique complète. »",
    fill=GREEN_LIGHT,
    accent=GREEN,
)
doc.add_heading("Ordre recommandé pour 12 à 15 minutes", level=1)
add_table(
    doc,
    ["Temps", "Sujet", "But"],
    [
        ["0:00–1:30", "Verdict et SOW", "Donner la conclusion avant les détails"],
        ["1:30–4:00", "Données et adaptation", "Rendre visible le travail d’ingénierie"],
        ["4:00–7:00", "Pipeline et reproductibilité", "Expliquer comment et pourquoi"],
        ["7:00–10:00", "Résultats A/B/C, scans et SLS", "Montrer la valeur et les limites"],
        ["10:00–12:00", "Dossiers, montant et suite", "Faciliter l’acceptation et demander une décision"],
        ["12:00+", "Questions", "Utiliser les diapositives d’annexe"],
    ],
    [1.15, 2.2, 3.65],
    header_fill=INK,
    font_size=9.5,
)
doc.add_heading("Posture pendant la réunion", level=2)
add_bullet(doc, "Commencer par le verdict, puis apporter les preuves.")
add_bullet(doc, "Distinguer systématiquement alerte, anomalie et diagnostic.")
add_bullet(doc, "Reconnaître les limites avant qu’on vous les reproche.")
add_bullet(doc, "Relier chaque étape à un fichier livré.")
add_bullet(doc, "Terminer par une demande précise à McGill.")
doc.add_page_break()

# SOW.
add_kicker(doc, "Contrat", BLUE)
add_title(doc, "Ce que l’objectif 1 demandait")
add_body(
    doc,
    "Le Specific Aim 1 du SOW est d’appliquer et d’évaluer la pipeline existante de détection de boiterie sur de nouvelles données IoT.",
)
add_table(
    doc,
    ["Tâche", "Travail demandé", "Livrable", "Preuve dans le paquet"],
    [
        [
            "1.1",
            "Adapter et exécuter IF + règles en bins de 15 min sur quatre jeux IceTag",
            "Données traitées avec alertes",
            "DONNEES_TRAITEES_ALERTES/",
        ],
        [
            "1.1",
            "Documenter l’exécution",
            "Note technique de reproductibilité",
            "NOTES_SOW/note_technique_reproductibilite.md",
        ],
        [
            "1.2",
            "Aligner les alertes aux scans comportementaux",
            "Table de concordance",
            "TABLEAUX_CSV/table_concordance.csv",
        ],
        [
            "1.2",
            "Évaluer et interpréter la concordance",
            "Rapport court de validation",
            "NOTES_SOW/rapport_validation_concordance.md",
        ],
    ],
    [0.55, 2.75, 1.85, 2.15],
    header_fill=INK,
    font_size=9.2,
)
add_callout(
    doc,
    "Verdict contractuel",
    "Les quatre livrables existent. Le rapport Word et le PowerPoint sont des couches de communication supplémentaires; ils ne remplacent pas les données et notes techniques.",
    fill=GREEN_LIGHT,
    accent=GREEN,
)
doc.add_heading("Ce que le montant de 3 700 $ couvre", level=2)
add_bullet(doc, "Audit d’un environnement de données hétérogène.")
add_bullet(doc, "Adaptation des entrées sans modifier la logique du modèle par saison.")
add_bullet(doc, "Exécution multi-saison, contrôles de qualité et exports volumineux.")
add_bullet(doc, "Alignement avec les scans, analyse SLS et qualification des limites.")
add_bullet(doc, "Documentation, synthèses, paquet de livraison et présentation.")
doc.add_page_break()

# Data.
add_kicker(doc, "Données", TEAL)
add_title(doc, "Quelles données ont été utilisées, et pourquoi")
add_body(
    doc,
    "L’inventaire complet recense 2 414 fichiers. Il a servi à identifier les bonnes sources, mais l’objectif 1 analytique utilise les quatre corpus IceTag définis dans le SOW.",
)
add_table(
    doc,
    ["Famille inventoriée", "Volume", "Rôle dans le projet"],
    [
        ["Accéléromètres IceTag", "2 256 fichiers", "Source principale de l’objectif 1"],
        ["Environnement HOBO", "151 fichiers", "Principalement objectif 2"],
        ["Scans comportementaux", "5 jeux + 1 agrégat", "Concordance de la tâche 1.2"],
        ["Documentation", "1 fichier principal", "Interprétation des expériences"],
    ],
    [2.2, 1.35, 3.45],
    header_fill=TEAL,
    font_size=9.5,
)
doc.add_heading("Pourquoi une conversion était nécessaire", level=2)
add_bullet(doc, "Noms de colonnes différents selon les fichiers.")
add_bullet(doc, "Timestamps, durées et types de données hétérogènes.")
add_bullet(doc, "Identifiants Cow_ID, couleurs et mappings variables selon l’expérience.")
add_bullet(doc, "Couverture irrégulière, profils partiels, doublons et valeurs impossibles.")
add_callout(
    doc,
    "Exemple concret",
    "L’erreur « None of ['T'] are in the columns » venait du fait que la pipeline attendait une colonne temporelle canonique T. La conversion a créé le schéma Cow, T, Steps, Motion Index, Lying/Standing et Transitions.",
    fill=YELLOW_LIGHT,
    accent="7A5A00",
)
doc.add_heading("Fall 2021 : pourquoi 10 et 8 apparaissent", level=2)
add_body(
    doc,
    "Dix profils ont été conservés dans la trace technique. Huit sont complets et soutiennent l’interprétation. Les deux profils partiels ne sont pas cachés : ils restent visibles pour la traçabilité.",
)
doc.add_page_break()

# Pipeline.
add_kicker(doc, "Méthode", BLUE)
add_title(doc, "Comment fonctionne la pipeline principale")
for step in [
    ("1. Harmonisation", "Normaliser les colonnes, les identités, les durées et le temps."),
    ("2. Bins de 15 minutes", "Agrèger les observations dans une grille temporelle commune."),
    ("3. Features robustes", "Calculer z-scores robustes, variations, moyennes mobiles, activité, repos et transitions."),
    ("4. Isolation Forest par vache", "Détecter les états atypiques par rapport au comportement habituel de la même vache."),
    ("5. Règles métier", "Exiger persistance, cohérence de plusieurs familles de signaux et couverture suffisante."),
    ("6. Notification", "Regrouper les épisodes et imposer un cooldown de 12 heures."),
    ("7. Contexte troupeau", "Comparer au troupeau et classer A, B ou C."),
]:
    add_body(doc, f"{step[0]} — {step[1]}", bold_lead=step[0])
doc.add_heading("Paramètres gelés", level=2)
add_table(
    doc,
    ["Paramètre", "Valeur", "Pourquoi"],
    [
        ["Intervalle", "15 min", "Conforme au SOW et aux scripts existants"],
        ["Contamination IF", "0,06", "Définition fixe des points atypiques"],
        ["Persistance", "7 h", "Écarter les changements trop courts"],
        ["Cooldown", "12 h", "Éviter les notifications répétées du même épisode"],
        ["Couverture minimale", "25 %", "Écarter les bins de mauvaise qualité"],
        ["Graine aléatoire", "42", "Reproductibilité"],
    ],
    [1.8, 1.2, 4.0],
    header_fill=INK,
    font_size=9.5,
)
add_callout(
    doc,
    "À dire",
    "« Les paramètres sont identiques entre saisons. Je n’ai pas ajusté les seuils après avoir vu les résultats SLS. »",
    fill=GREEN_LIGHT,
    accent=GREEN,
)
doc.add_page_break()

# Outputs and folders.
add_kicker(doc, "Livraison", TEAL)
add_title(doc, "À quoi servent les dossiers et fichiers")
add_table(
    doc,
    ["Élément", "Contenu", "Quand l’ouvrir"],
    [
        ["README_livraison_objectif1.txt", "Carte du paquet et correspondance SOW", "Toujours en premier"],
        ["RAPPORTS/", "Rapport Word + présentations", "Pour comprendre et présenter"],
        ["DONNEES_TRAITEES_ALERTES/", "Predictions, alerts_only, summary, fichier renforcé", "Pour auditer les résultats"],
        ["TABLEAUX_CSV/", "Synthèses saison, vache, confiance et concordance", "Pour vérifier les chiffres"],
        ["NOTES_SOW/", "Reproductibilité et rapport de validation", "Pour répondre au contrat"],
        ["ANNEXE_pipeline_actuelle_…/", "HYPO + instabilité + hybride", "Pour la direction scientifique actuelle"],
    ],
    [2.25, 2.9, 1.85],
    header_fill=TEAL,
    font_size=9.1,
)
doc.add_heading("Les trois CSV par saison", level=2)
add_body(
    doc,
    "predictions.csv — Niveau détaillé : une ligne par vache et par intervalle de 15 minutes, avec features, scores, drapeaux et couverture.",
    bold_lead="predictions.csv",
)
add_body(
    doc,
    "alerts_only.csv — Niveau opérationnel : seulement les débuts d’épisodes retenus après les règles et le cooldown.",
    bold_lead="alerts_only.csv",
)
add_body(
    doc,
    "summary.csv — Niveau contrôle : résumé par vache des bins, anomalies, épisodes et notifications.",
    bold_lead="summary.csv",
)
add_body(
    doc,
    "objective1_reinforced_alerts.csv — Ajoute le contexte troupeau et la priorité A/B/C tout en conservant la notification initiale.",
    bold_lead="objective1_reinforced_alerts.csv",
)
doc.add_heading("Pourquoi certains CSV datent du 29 mai", level=2)
add_body(
    doc,
    "Les fichiers saisonniers initiaux conservent leur date de génération originale pour préserver la traçabilité. Les sorties renforcées et les rapports finaux sont plus récents. Une date ancienne ne signifie donc pas que le fichier est obsolète.",
)
doc.add_page_break()

# Results interpretation.
add_kicker(doc, "Résultats", BLUE)
add_title(doc, "Comment interpréter les alertes")
add_table(
    doc,
    ["Niveau", "Nombre", "Sens", "Action"],
    [
        ["A", "37", "Signal individuel net", "Examiner en premier"],
        ["B", "195", "Signal individuel plausible", "Vérifier le contexte de la vache"],
        ["C", "153", "Plusieurs vaches touchées simultanément", "Chercher un événement commun"],
        ["D", "0", "Qualité ou contexte insuffisant", "Aucune dans la sortie finale"],
    ],
    [0.65, 0.75, 3.05, 2.55],
    header_fill=INK,
    font_size=9.5,
)
add_body(
    doc,
    "La requalification ne supprime pas les 385 notifications. Elle améliore leur sens en séparant les signaux individuels des événements probablement collectifs.",
)
doc.add_heading("Pourquoi Winter 2019 semble bruyant", level=2)
add_body(
    doc,
    "89 des 149 notifications Winter 2019 sont classées C, soit 59,7 %. Cela indique qu’une grande partie des changements a touché plusieurs vaches au même moment. Le bon réflexe est donc d’étudier le contexte collectif avant de soupçonner une boiterie individuelle.",
)
doc.add_heading("Concordance avec les scans", level=2)
add_table(
    doc,
    ["Saison", "Scans comparables", "Avec alerte ±1 jour", "Taux descriptif"],
    [
        ["Fall 2019", "27", "12", "44,4 %"],
        ["Winter 2019", "41", "11", "26,8 %"],
        ["Summer 2019", "55", "14", "25,5 %"],
        ["Fall 2021", "10", "1", "10,0 %"],
        ["Total", "133", "38", "28,6 %"],
    ],
    [1.7, 1.45, 2.2, 1.65],
    header_fill=TEAL,
    font_size=9.5,
)
add_body(
    doc,
    "Sur 396 scans datés, seuls 133 sont comparables : la vache doit disposer de données IceTag dans une fenêtre de ±1 jour. Ce filtre ramène Fall 2021 de 270 scans à 10, ce qui évite qu'un corpus de six jours écrase l'analyse.",
)
add_callout(
    doc,
    "Niveau attendu, le point à annoncer soi-même",
    "Sur les 127 scans couverts le jour même, on observe 38 concordances (29,9 %) contre 35,3 attendues (27,8 %) d'après la fréquence d'alerte propre à chaque vache. L'écart de +2,1 points n'est pas concluant (Poisson-binomial unilatéral, p = 0,324). La concordance documente donc une proximité temporelle, sans enrichissement démontré au-delà du niveau attendu.",
    fill=RED_LIGHT,
    accent=RED,
)
add_callout(
    doc,
    "Attention",
    "Ces pourcentages sont des chevauchements temporels, pas des précisions diagnostiques. Les scans comportementaux ne sont pas une vérité-terrain clinique de boiterie.",
    fill=RED_LIGHT,
    accent=RED,
)
doc.add_page_break()

# SLS and hybrid.
add_kicker(doc, "Validation", BLUE)
add_title(doc, "Ce que les scores SLS permettent réellement de dire")
add_table(
    doc,
    ["Analyse", "Cohorte", "Résultat", "Conclusion correcte"],
    [
        [
            "IF + règles",
            "16 vaches; 5 SLS ≥ 2",
            "p=0,649; ρ=0,033",
            "Aucune association observée",
        ],
        [
            "HYPO + instabilité + hybride",
            "14 évaluables; 3 SLS ≥ 2",
            "AUC=0,924; p=0,031",
            "Signal exploratoire prometteur",
        ],
    ],
    [1.65, 1.65, 1.45, 2.25],
    header_fill=INK,
    font_size=9.3,
)
add_body(
    doc,
    "L’approche IF + règles est la baseline contractuelle. L’approche HYPO + instabilité + hybride est l’évolution scientifique actuelle, livrée séparément en annexe.",
)
add_heading = doc.add_heading
add_heading("Pourquoi les deux résultats ne sont pas une comparaison clinique directe", level=2)
add_bullet(doc, "Les définitions de notification ne sont pas identiques.")
add_bullet(doc, "Les cohortes évaluables diffèrent.")
add_bullet(doc, "Il n’y a que trois vaches positives dans l’analyse hybride.")
add_bullet(doc, "Une AUC élevée peut être instable avec un si petit échantillon.")
add_callout(
    doc,
    "Formulation recommandée",
    "« L’hybride montre une meilleure séparation observationnelle dans la petite cohorte SLS disponible. C’est encourageant pour la suite, mais insuffisant pour annoncer une sensibilité, une spécificité ou une supériorité clinique générale. »",
    fill=YELLOW_LIGHT,
    accent="7A5A00",
)
doc.add_heading("Pourquoi l’annexe produit plus d’alertes", level=2)
add_body(
    doc,
    "La baseline produit 385 notifications; l’hybride en produit 1 179. L’approche actuelle est plus sensible analytiquement, mais sa charge opérationnelle est plus élevée. Une validation future devra donc étudier à la fois la détection et le nombre d’alertes à examiner.",
)
doc.add_page_break()

# Slide-by-slide.
add_kicker(doc, "Présentation", TEAL)
add_title(doc, "Trame diapositive par diapositive")
slide_scripts = [
    ("Objectif 1", "Présenter le problème : quatre corpus hétérogènes à transformer en sorties comparables et interprétables."),
    ("Le résultat en une minute", "Donner les trois chiffres principaux et le verdict : SOW atteint, diagnostic clinique non encore validé."),
    ("Ce que le SOW demandait", "Montrer les quatre livrables et leur emplacement. Citer le montant seulement après les preuves."),
    ("Les données mobilisées", "Distinguer l’inventaire de 2 414 fichiers des quatre corpus réellement analysés."),
    ("Pourquoi une adaptation", "Expliquer les formats, timestamps, identités et contrôles de qualité; rappeler l’erreur initiale sur T."),
    ("La chaîne d’analyse", "Décrire harmonisation, bins, features, IF, règles et qualification A/B/C."),
    ("Paramètres gelés", "Montrer que les saisons ont été comparées avec les mêmes réglages."),
    ("Trois niveaux de sortie", "Expliquer predictions, alerts_only, summary et le fichier renforcé."),
    ("385 notifications", "Lire les totaux par saison et préciser qu’ils doivent être normalisés par exposition."),
    ("Normalisation troupeau", "Faire comprendre que le nombre ne change pas, mais que l’ordre de revue devient utile."),
    ("Événements collectifs", "Montrer que Winter et Fall 2019 contiennent une grande part de signaux collectifs."),
    ("Concordance scans", "Présenter le chevauchement ±1 jour comme descriptif, jamais comme précision."),
    ("Scores SLS", "Présenter IF comme non associé et l’hybride comme prometteur mais fragile."),
    ("Réussi / non démontré", "Séparer nettement résultat contractuel et validation clinique."),
    ("Lire le dossier", "Donner l’ordre README → rapport → CSV; laisser l’annexe séparée."),
    ("Notebooks et scripts", "Montrer que chaque fichier répond à une étape de la preuve."),
    ("Ce que rémunère le travail", "Relier le montant à l’inventaire, l’harmonisation, l’exécution, l’évaluation et la livraison."),
    ("La suite", "Demander une cohorte clinique et une validation prospective pré-spécifiée."),
    ("Conclusion", "Résumer en une phrase, puis ouvrir les questions."),
]
for i, (title, script) in enumerate(slide_scripts, start=1):
    add_slide_script(doc, i, title, script)
doc.add_page_break()

# Q&A core.
add_kicker(doc, "Questions-réponses", BLUE)
add_title(doc, "Questions scientifiques et techniques probables")
add_question(
    doc,
    "1. Avez-vous détecté des boiteries ?",
    "J’ai détecté des alertes comportementales compatibles avec un changement à vérifier; je ne peux pas confirmer 385 boiteries cliniques.",
    [
        "Une confirmation exige un examen clinique ou un score locomoteur synchronisé.",
        "Le langage prudent rend le résultat plus défendable, pas moins utile.",
    ],
)
add_question(
    doc,
    "2. Qu’est-ce qu’une notification exactement ?",
    "Le début d’un épisode persistant retenu après anomalies, règles de cohérence, couverture et cooldown.",
    [
        "Un point anormal isolé ne devient pas automatiquement une notification.",
        "Le cooldown de 12 heures évite de recompter le même épisode.",
    ],
)
add_question(
    doc,
    "3. Le score de confiance est-il une probabilité de boiterie ?",
    "Non. C’est un score empirique de classement des alertes, pas une probabilité calibrée cliniquement.",
)
add_question(
    doc,
    "4. Pourquoi Isolation Forest ?",
    "Parce que le SOW demande explicitement la pipeline validée existante IF + règles et parce que les labels cliniques sont trop rares pour entraîner correctement un modèle supervisé.",
)
doc.add_page_break()
doc.add_heading("Questions scientifiques et techniques — suite", level=2)
add_question(
    doc,
    "5. Pourquoi appliquer aussi HYPO + instabilité + hybride ?",
    "Pour montrer la direction scientifique actuelle sans modifier ni remplacer silencieusement le livrable contractuel.",
    [
        "La baseline reste principale.",
        "L’approche actuelle est dans une annexe indépendante.",
    ],
)
add_question(
    doc,
    "6. Quelle pipeline est la meilleure ?",
    "IF est la référence contractuelle la plus comparable; l’hybride est la direction la plus prometteuse scientifiquement, mais elle doit être validée et calibrée.",
)
add_question(
    doc,
    "7. Pourquoi ne pas entraîner un classifieur supervisé ?",
    "La cohorte contient trop peu de cas SLS positifs et pas assez de labels répétés pour éviter le surapprentissage.",
)
add_question(
    doc,
    "8. Les paramètres ont-ils été optimisés sur les SLS ?",
    "Non. Les paramètres de la baseline ont été gelés avant l’analyse de concordance.",
)
doc.add_page_break()

add_kicker(doc, "Questions-réponses", BLUE)
add_title(doc, "Questions sur les résultats et les limites")
add_question(
    doc,
    "9. 385 alertes, n’est-ce pas beaucoup ?",
    "C’est un total multi-saison avant priorisation. Le taux global est 9,855 par 100 vache-jours; 37 seulement sont A et 153 sont collectives probables.",
)
add_question(
    doc,
    "10. Pourquoi Winter 2019 a-t-il 149 alertes ?",
    "La saison est longue et 89 de ces alertes sont probablement collectives. Le contexte commun explique une grande partie du volume.",
)
add_question(
    doc,
    "11. Pourquoi Fall 2021 a-t-il seulement quatre alertes ?",
    "La fenêtre IceTag analysée ne couvre qu’environ six jours; le total n’est pas comparable aux saisons de plusieurs mois.",
)
add_question(
    doc,
    "12. Fall 2021 a-t-il 10 ou 8 vaches ?",
    "Dix profils sont traités et tracés; huit sont complets pour l’interprétation. Les deux profils partiels sont conservés pour la transparence.",
)
add_question(
    doc,
    "13. Les taux de concordance 44,4 %, 26,8 % et 25,5 % sont-ils des précisions ?",
    "Non. Ce sont des proportions de scans comparables avec une alerte dans une fenêtre ±1 jour, soit 28,6 % au total sur 133 scans.",
)
add_question(
    doc,
    "14. Pourquoi Fall 2021 a-t-il une concordance de 10,0 % ?",
    "Le recouvrement IceTag est très court : sur 270 scans datés, seuls 10 disposent de données capteur à ±1 jour, et le corpus ne couvre que six jours. Le chiffre reflète surtout ce faible recouvrement.",
)
add_question(
    doc,
    "14b. Ce taux de 28,6 % dépasse-t-il ce que donnerait le hasard ?",
    "Non, pas de façon démontrée. Les alertes se déclenchent environ 9 fois par 100 vache-jours et la fenêtre couvre trois journées, donc le niveau attendu est déjà de 27,8 %. On observe 29,9 %, soit +2,1 points, avec p = 0,324 : non concluant. La concordance documente une proximité temporelle, pas un enrichissement.",
)
add_question(
    doc,
    "15. L’AUC 0,924 prouve-t-elle que l’hybride fonctionne ?",
    "Elle montre une séparation exploratoire encourageante, mais seulement trois vaches sont SLS ≥ 2. Elle ne suffit pas pour une conclusion clinique générale.",
)
add_question(
    doc,
    "16. Pourquoi ne donnez-vous pas sensibilité et spécificité ?",
    "Parce qu’il manque une vérité-terrain clinique synchronisée, complète et assez grande pour une matrice de confusion fiable.",
)
doc.add_page_break()

add_kicker(doc, "Questions-réponses", TEAL)
add_title(doc, "Questions sur les fichiers, la reproductibilité et la valeur")
add_question(
    doc,
    "17. Pourquoi y a-t-il autant de fichiers ?",
    "Parce que les sorties détaillées, les notifications, les synthèses, les preuves SOW et l’annexe ont des responsabilités différentes.",
)
add_question(
    doc,
    "18. Pourquoi certains fichiers datent-ils du 29 mai ?",
    "Ils conservent leur date de génération originale pour la traçabilité; les sorties renforcées et rapports finaux sont plus récents.",
)
add_question(
    doc,
    "19. Le projet McGill dépend-il du dépôt du mémoire ?",
    "Le paquet McGill est autonome. L’annexe lit les fruits de l’approche actuelle, mais aucun fichier du dépôt du mémoire n’est modifié.",
)
add_question(
    doc,
    "20. Peut-on reproduire les résultats ?",
    "Oui. Les paramètres sont gelés, la conversion est documentée, les notebooks suivent l’ordre audit → conversion → pipeline → concordance → renforcement, et les sorties intermédiaires sont conservées.",
)
add_question(
    doc,
    "21. La pipeline peut-elle fonctionner en temps réel ?",
    "La logique peut être adaptée à un flux continu, mais le travail livré est rétrospectif et l’usage temps réel doit encore être validé opérationnellement.",
)
add_question(
    doc,
    "22. Qu’est-ce qui a échoué ?",
    "La baseline IF n’a pas montré de concordance SLS dans la petite cohorte. C’est une limite scientifique identifiée, pas un échec d’exécution.",
)
add_question(
    doc,
    "23. En quoi l’objectif 1 est-il positif ?",
    "Il transforme quatre expériences hétérogènes en une chaîne cohérente, reproductible et interprétable, et réduit le risque de fausses interprétations collectives.",
)
add_question(
    doc,
    "24. Le montant de 3 700 $ est-il justifié ?",
    "Oui au regard du SOW : quatre semaines de transfert, adaptation, exécution, validation exploratoire, documentation et quatre livrables vérifiables.",
)

# Phrases.
add_kicker(doc, "Communication", BLUE)
add_title(doc, "Formulations à utiliser et à éviter")
add_table(
    doc,
    ["À éviter", "À dire"],
    [
        ["« 385 vaches boiteuses »", "« 385 notifications brutes sur quatre saisons »"],
        ["« Le modèle diagnostique la boiterie »", "« Le système signale un changement comportemental persistant à vérifier »"],
        ["« 44,4 % de précision »", "« 44,4 % des scans Fall 2019 ont une alerte dans ±1 jour »"],
        ["« L’AUC 0,924 valide l’hybride »", "« L’AUC est prometteuse dans une petite cohorte exploratoire »"],
        ["« IF est mauvais »", "« IF est la baseline contractuelle; sa concordance SLS n’est pas démontrée ici »"],
        ["« HYPO est définitivement meilleur »", "« HYPO + instabilité + hybride est la direction actuelle à valider »"],
        ["« Les fichiers anciens sont obsolètes »", "« Les dates initiales sont conservées pour la traçabilité »"],
        ["« Le résultat est négatif »", "« Le transfert est réussi; la portée clinique est limitée par les labels »"],
    ],
    [2.95, 4.05],
    header_fill=INK,
    font_size=9.2,
)
doc.add_heading("Réponses-ponts utiles", level=2)
add_bullet(doc, "« La réponse courte est…, et je peux vous montrer le fichier qui la documente. »")
add_bullet(doc, "« Il faut distinguer le résultat contractuel de la validation clinique. »")
add_bullet(doc, "« Ce chiffre mesure une concordance temporelle, pas une précision diagnostique. »")
add_bullet(doc, "« La limite vient surtout de la taille et de la synchronisation de la vérité-terrain. »")
add_bullet(doc, "« C’est précisément pour cela que je propose une validation prospective. »")
doc.add_page_break()

# Next steps and asks.
add_kicker(doc, "Décision", TEAL)
add_title(doc, "Ce qu’il faut demander à McGill")
add_numbered(doc, "Confirmer que les quatre livrables de l’objectif 1 sont acceptés.")
add_numbered(doc, "Identifier la meilleure cohorte avec IceTag et scores locomoteurs/SLS datés.")
add_numbered(doc, "Clarifier les mappings couleur ↔ Cow_ID restants pour Fall 2019 et Fall 2021.")
add_numbered(doc, "Définir la fenêtre clinique d’intérêt : même jour, 3 jours ou 7 jours avant le score.")
add_numbered(doc, "Choisir les critères de succès de la validation prospective : sensibilité, spécificité, délai et charge d’alertes.")
add_numbered(doc, "Décider si l’annexe hybride devient la pipeline candidate pour l’étude suivante.")
doc.add_heading("Protocole recommandé", level=2)
add_table(
    doc,
    ["Étape", "Action", "Sortie attendue"],
    [
        ["1", "Geler la version et les seuils", "Pipeline candidate versionnée"],
        ["2", "Collecter des scores répétés et indépendants", "Vérité-terrain synchronisée"],
        ["3", "Appliquer sans réajustement", "Prédictions prospectives"],
        ["4", "Calculer métriques et intervalles de confiance", "Performance avec incertitude"],
        ["5", "Évaluer la charge de revue A/B/C", "Faisabilité opérationnelle"],
    ],
    [0.65, 3.35, 3.0],
    header_fill=TEAL,
    font_size=9.5,
)
add_callout(
    doc,
    "Question finale",
    "« Pouvez-vous confirmer l’acceptation de l’objectif 1 et nous indiquer quelle cohorte fournirait le meilleur recouvrement entre capteurs et scores locomoteurs pour la validation suivante ? »",
    fill=GREEN_LIGHT,
    accent=GREEN,
)
doc.add_page_break()

# Checklist.
add_kicker(doc, "Checklist", BLUE)
add_title(doc, "La veille et le jour de la réunion")
doc.add_heading("La veille", level=2)
for item in [
    "Ouvrir le PowerPoint et vérifier les notes d’orateur.",
    "Garder le dossier Objectif1_Pipeline_detection_boiterie accessible.",
    "Ouvrir README_livraison_objectif1.txt et les trois CSV de synthèse.",
    "Répéter l’ouverture de 90 secondes et la conclusion de 20 secondes.",
    "Préparer une réponse calme sur les limites SLS.",
]:
    add_bullet(doc, item)
doc.add_heading("Pendant la réunion", level=2)
for item in [
    "Ne pas lire les diapositives; annoncer le message, puis montrer la preuve.",
    "Utiliser les annexes seulement si une question les exige.",
    "Noter les demandes de modification et les décisions d’acceptation.",
    "Demander qui détient les labels cliniques et les mappings d’identité.",
    "Conclure avec une prochaine étape, un responsable et une date.",
]:
    add_bullet(doc, item)
doc.add_heading("Conclusion de 20 secondes", level=2)
add_callout(
    doc,
    "Texte proposé",
    "« L’objectif 1 est terminé au sens du SOW : les quatre saisons sont traitées, les sorties sont reproductibles et les livrables sont présents. L’analyse montre une utilité comme système d’alerte comportementale, surtout après qualification du contexte troupeau. La prochaine étape est une validation prospective avec davantage de labels cliniques synchronisés. »",
    fill=BLUE_LIGHT,
    accent=BLUE,
)

# Sources.
doc.add_page_break()
add_kicker(doc, "Références", TEAL)
add_title(doc, "Sources locales de vérification")
add_body(
    doc,
    "Dossier livré : /Users/alioubarry/Desktop/Livrables_McGill_WellE/Objectif1_Pipeline_detection_boiterie",
    color=MUTED,
)
for source in [
    "README_livraison_objectif1.txt",
    "RAPPORTS/Objectif1_rapport_livraison.docx",
    "TABLEAUX_CSV/objective1_multi_season_summary.csv",
    "TABLEAUX_CSV/objective1_reinforced_summary_by_season.csv",
    "TABLEAUX_CSV/concordance_par_experience.csv",
    "Annexe SLS : pipeline_actuelle_validation_sls_synthese.csv",
]:
    add_bullet(doc, source, color=MUTED)
add_body(doc, "Code : /Users/alioubarry/PROJECT/core et /Users/alioubarry/PROJECT/mcgill_iot_cattle", color=MUTED)
for source in [
    "core/config.py",
    "core/pipeline.py",
    "mcgill_iot_cattle/run_objective1_reinforcement.py",
    "mcgill_iot_cattle/SOW Alliou - Complété (avec montants).docx",
]:
    add_bullet(doc, source, color=MUTED)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
