from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(
    "/Users/alioubarry/Desktop/Livrables_McGill_WellE/"
    "Reunion_McGill_2026-07-30/"
    "Fiche_questions_reponses_Objectif1_McGill_30_juillet_2026.docx"
)

BLUE = "2474E5"
DARK_BLUE = "1F4D78"
INK = "111111"
MUTED = "5E6875"
LIGHT_BLUE = "E8F0FB"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF5D6"
GOLD = "7A5A00"
RED = "9B1C1C"
GREEN = "2D8A57"
WHITE = "FFFFFF"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)


def shade_paragraph(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("McGill / WELL-E  |  Objectif 1  |  Préparation Q&R")
    set_run(run, size=9, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Réunion du 30 juillet 2026  |  Page ")
    set_run(run, size=9, color=MUTED)
    add_page_field(fp)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("FICHE DE PRÉPARATION")
    set_run(run, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Questions et réponses probables")
    set_run(run, size=26, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Objectif 1 — Pipeline de détection sur les données IceTag")
    set_run(run, size=15, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(
        "Usage : document personnel de répétition pour la rencontre McGill/WELL-E. "
        "Commencer par la réponse courte; développer seulement si la question se poursuit."
    )
    set_run(run, size=11, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    shade_paragraph(p, LIGHT_BLUE, BLUE)
    lead = p.add_run("Verdict à retenir : ")
    set_run(lead, size=12, bold=True, color=BLUE)
    body = p.add_run(
        "l’objectif 1 est atteint techniquement et contractuellement. La baseline IF est "
        "reproductible, mais non validée cliniquement. L’approche actuelle hybride montre "
        "un alignement SLS exploratoire encourageant qui doit être confirmé."
    )
    set_run(body, size=12, color=INK)


def add_section(doc, title, intro=None, page_break=False):
    if page_break:
        doc.add_page_break()
    doc.add_heading(title, level=1)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(9)
        for run in p.runs:
            set_run(run, color=MUTED, italic=True)


def add_qa(doc, number, question, short_answer, detail=None, caution=None, single_run=False):
    q = doc.add_heading(level=3)
    q.paragraph_format.page_break_before = False
    q.paragraph_format.keep_with_next = True
    run = q.add_run(f"{number}. {question}")
    set_run(run, size=12, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.03)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = bool(detail or caution)
    shade_paragraph(p, LIGHT_BLUE, BLUE)
    if single_run:
        body = p.add_run(f"Réponse : {short_answer}")
        set_run(body, color=INK)
    else:
        label = p.add_run("Réponse : ")
        set_run(label, bold=True, color=BLUE)
        body = p.add_run(short_answer)
        set_run(body, color=INK)

    if detail:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = bool(caution)
        label = p.add_run("Précision : ")
        set_run(label, bold=True, color=DARK_BLUE)
        body = p.add_run(detail)
        set_run(body, color=INK)

    if caution:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.keep_together = True
        shade_paragraph(p, LIGHT_GOLD, GOLD)
        label = p.add_run("À éviter : ")
        set_run(label, bold=True, color=GOLD)
        body = p.add_run(caution)
        set_run(body, color=INK)


def add_key_messages(doc):
    doc.add_page_break()
    doc.add_heading("Les huit réponses à maîtriser", level=1)
    messages = [
        (
            "Le mandat",
            "Adapter et exécuter une pipeline existante sur quatre corpus IceTag, puis "
            "évaluer la concordance et documenter les résultats.",
        ),
        (
            "Le volume traité",
            "Quatre saisons, 375 031 intervalles vache × 15 minutes et 385 débuts "
            "d’épisodes retenus par la baseline.",
        ),
        (
            "Le sens des 385",
            "Ce sont des notifications comportementales à vérifier, pas 385 diagnostics "
            "ni 385 vaches boiteuses.",
        ),
        (
            "A / B / C",
            "A = signal individuel prioritaire; B = signal individuel à vérifier; "
            "C = événement collectif probable. Ce ne sont pas des stades cliniques.",
        ),
        (
            "La concordance",
            "Les taux avec les scans mesurent une cooccurrence temporelle dans ± 1 jour, "
            "pas une précision diagnostique.",
        ),
        (
            "Les SLS et la baseline",
            "La baseline IF n’est pas alignée aux SLS disponibles : p = 0,649 et "
            "rho = 0,033.",
        ),
        (
            "Les SLS et l’approche actuelle",
            "L’hybride montre un signal exploratoire positif : AUC = 0,924 et p = 0,031, "
            "mais seulement 3 cas positifs et un confondant Exercise.",
        ),
        (
            "La conclusion",
            "Succès technique et contractuel; outil de recherche et de revue ciblée; "
            "validation clinique prospective encore requise.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2016, 7344])
    hdr = table.rows[0].cells
    hdr[0].text = "Repère"
    hdr[1].text = "Formulation à retenir"
    for cell in hdr:
        set_cell_margins(cell)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), DARK_BLUE)
        tc_pr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run(run, bold=True, color=WHITE)
    for row_index, (label, value) in enumerate(messages, start=1):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = 1
            if row_index % 2 == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_GRAY)
                tc_pr.append(shd)
        for run in cells[0].paragraphs[0].runs:
            set_run(run, bold=True, color=DARK_BLUE)
        for run in cells[1].paragraphs[0].runs:
            set_run(run, color=INK)
    set_table_geometry(table, [2016, 7344])


def add_questions_to_ask(doc):
    add_section(
        doc,
        "Questions à poser à McGill/WELL-E",
        "Ces questions permettent de conclure la rencontre par des décisions concrètes.",
        page_break=False,
    )
    questions = [
        "Pouvez-vous confirmer que les quatre composantes du livrable de l’objectif 1 répondent au SOW?",
        "Quelle est la définition exacte du score SLS utilisé, et le score a-t-il été attribué par le même évaluateur?",
        "Existe-t-il d’autres dates de scores locomoteurs synchronisées aux périodes IceTag?",
        "Disposez-vous d’une cohorte comportant davantage de vaches SLS ≥ 2 et, idéalement, des SLS ≥ 3?",
        "Pouvez-vous documenter le groupe Exercise et les autres traitements expérimentaux pour dissocier leur effet de celui de la locomotion?",
        "Un événement de gestion, de météo ou de protocole explique-t-il le regroupement collectif observé au début de février 2019?",
        "Pour une validation suivante, quelle charge quotidienne de notifications A et B serait acceptable pour l’équipe terrain?",
    ]
    for idx, question in enumerate(questions, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(question)
        set_run(run, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, LIGHT_BLUE, BLUE)
    lead = p.add_run("Clôture proposée — ")
    set_run(lead, bold=True, color=BLUE)
    body = p.add_run(
        "« Je propose que nous confirmions aujourd’hui la livraison de l’objectif 1, "
        "puis que nous choisissions ensemble la cohorte clinique et le protocole de "
        "validation pour l’étape suivante. »"
    )
    set_run(body, color=INK)


def add_sources(doc):
    doc.add_heading("Sources internes utilisées", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    run = p.add_run(
        "SOW Alliou complété; README_livraison_objectif1.txt; "
        "Objectif1_rapport_livraison.docx; objective1_multi_season_summary.csv; "
        "objective1_reinforced_summary_by_season.csv; concordance_par_experience.csv; "
        "pipeline_actuelle_validation_sls_synthese.csv; scripts de configuration, "
        "pipeline, moteur d’alertes et renforcement troupeau."
    )
    set_run(run, size=8.5, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)
    add_key_messages(doc)

    add_section(
        doc,
        "A. Mandat, périmètre et état d’avancement",
        "Questions probables pour vérifier si le travail correspond réellement au SOW.",
        page_break=True,
    )
    add_qa(
        doc,
        1,
        "Quel était exactement l’objectif 1?",
        "Appliquer et évaluer une pipeline existante sur quatre corpus IceTag, puis "
        "documenter les données traitées, les alertes et la concordance.",
        "La tâche 1.1 couvre l’adaptation des entrées et l’exécution. La tâche 1.2 couvre "
        "l’alignement avec les scans et informations disponibles, la table de concordance "
        "et le rapport de validation.",
    )
    add_qa(
        doc,
        2,
        "L’objectif 1 est-il terminé?",
        "Oui au sens technique et contractuel : quatre saisons sur quatre ont été traitées "
        "et les quatre composantes du SOW sont présentes.",
        "La validation clinique complète n’était pas atteignable avec les labels disponibles. "
        "Elle constitue l’étape scientifique suivante, pas un livrable manquant de l’exécution.",
        "Ne pas répondre simplement « oui, le détecteur de boiterie est validé ». Le bon verdict "
        "est : objectif du SOW atteint, validation clinique encore exploratoire.",
    )
    add_qa(
        doc,
        3,
        "Avez-vous modifié la pipeline pour obtenir de meilleurs résultats?",
        "Non pour la baseline du SOW. Les mêmes paramètres gelés ont été appliqués aux quatre saisons.",
        "L’évolution HYPO + instabilité + hybride est évaluée séparément en annexe. Cette séparation "
        "évite de confondre l’exécution contractuelle et le développement scientifique plus récent.",
    )
    add_qa(
        doc,
        4,
        "Pourquoi présenter deux approches?",
        "Parce qu’elles répondent à deux questions différentes : la baseline montre que le SOW a été "
        "exécuté de façon reproductible; l’approche actuelle teste une évolution scientifique.",
        "La comparaison aux SLS montre justement pourquoi cette distinction est utile : la baseline "
        "n’est pas alignée aux SLS disponibles, tandis que l’hybride présente un signal exploratoire positif.",
    )
    add_qa(
        doc,
        5,
        "Qu’est-ce qui a été réellement réalisé, au-delà du lancement d’un algorithme?",
        "Inventaire des sources, harmonisation des formats, contrôles de qualité, exécution commune, "
        "analyse des notifications, concordance, validation SLS exploratoire et paquet de livraison.",
        "L’inventaire a couvert 2 414 fichiers. Les sorties sont traçables à trois niveaux : intervalles, "
        "notifications et synthèses par vache ou saison.",
    )

    add_section(
        doc,
        "B. Données, qualité et couverture",
        "Répondre ici avec des définitions concrètes; ne pas confondre couverture temporelle et qualité clinique.",
    )
    add_qa(
        doc,
        6,
        "Quelles données ont été utilisées?",
        "Quatre corpus IceTag : Winter 2019, Summer 2019, Fall 2019 et Fall 2021.",
        "Ils représentent 75 profils tracés au total et 375 031 intervalles vache × 15 minutes. Les durées "
        "et les effectifs diffèrent entre les expériences.",
    )
    add_qa(
        doc,
        7,
        "Pourquoi les saisons ne sont-elles pas directement comparables?",
        "Parce qu’elles n’ont pas la même durée, le même nombre de vaches ni le même protocole expérimental.",
        "Fall 2021 ne couvre qu’environ une semaine. Les totaux bruts doivent donc être accompagnés du temps "
        "d’observation et du contexte de chaque expérience.",
    )
    add_qa(
        doc,
        8,
        "Que signifie la couverture?",
        "La couverture mesure la complétude temporelle d’un intervalle : échantillons bruts présents divisés "
        "par échantillons attendus.",
        "Elle sert à empêcher un manque de données d’être interprété comme un changement de comportement. "
        "Un intervalle sous 25 % de couverture ne peut pas générer de notification.",
        "Ne pas dire que 94,4 % de couverture signifie 94,4 % de vaches saines ou 94,4 % de précision.",
    )
    add_qa(
        doc,
        9,
        "Pourquoi Fall 2021 affiche-t-il 100 % de couverture mais seulement quatre notifications?",
        "Le 100 % décrit la complétude des intervalles conservés, pas la longueur du suivi. La période exploitable "
        "est très courte.",
        "Une bonne couverture sur une semaine ne produit pas le même volume qu’une bonne couverture sur plusieurs mois.",
    )
    add_qa(
        doc,
        10,
        "Comment les valeurs manquantes sont-elles traitées?",
        "Elles sont suivies par la couverture et ne sont pas automatiquement converties en immobilité.",
        "Les contrôles portent aussi sur les doublons, l’ordre chronologique, les identifiants de vache, les colonnes "
        "requises et les profils partiels.",
    )

    add_section(
        doc,
        "C. Fonctionnement de la pipeline",
        "La réponse centrale : Isolation Forest détecte l’atypique; les règles temporelles transforment ce signal en épisode.",
        page_break=True,
    )
    add_qa(
        doc,
        11,
        "Comment une notification est-elle produite?",
        "Agrégation à 15 minutes, création de variables robustes, Isolation Forest par vache, règles de persistance "
        "et de cohérence, contrôle de couverture, puis classement A/B/C.",
        "Une notification est inscrite au début d’un épisode retenu. Le délai de répétition de 12 heures évite de "
        "notifier continuellement le même épisode.",
    )
    add_qa(
        doc,
        12,
        "Pourquoi utiliser Isolation Forest?",
        "Parce qu’il peut construire un profil individuel sans exiger de nombreux labels cliniques.",
        "Il détecte des états atypiques par rapport à la même vache. Sa limite est précisément qu’un état atypique "
        "n’est pas nécessairement une boiterie.",
    )
    add_qa(
        doc,
        13,
        "Pourquoi entraîner le modèle par vache?",
        "Les niveaux d’activité diffèrent fortement d’une vache à l’autre. Le profil individuel réduit le risque "
        "de considérer une vache naturellement moins active comme anormale.",
        "Le contexte troupeau est ajouté ensuite pour vérifier si le changement est propre à la vache ou partagé.",
    )
    add_qa(
        doc,
        14,
        "Quels paramètres ont été utilisés?",
        "Intervalles de 15 minutes, contamination IF de 6 %, persistance sur 7 heures, taux d’anomalies d’au moins "
        "24 %, couverture minimale de 25 % et cooldown de 12 heures.",
        "Une portion initiale de 60 % sert de ligne de base. Les paramètres restent identiques entre saisons pour "
        "préserver la reproductibilité et éviter un ajustement opportuniste.",
    )
    add_qa(
        doc,
        15,
        "Quelle est la différence entre anomalie, épisode et notification?",
        "Une anomalie est un intervalle atypique; un épisode est une séquence persistante qui satisfait les règles; "
        "une notification est le début enregistré de cet épisode.",
        "Cette distinction explique pourquoi le nombre d’intervalles anormaux est beaucoup plus élevé que les 385 notifications.",
    )
    add_qa(
        doc,
        16,
        "Est-ce que la pipeline détecte directement la boiterie?",
        "Non. Elle détecte des changements comportementaux persistants compatibles avec un problème à vérifier.",
        "Une observation clinique ou un score locomoteur reste nécessaire pour confirmer la boiterie et mesurer la "
        "sensibilité ou la spécificité.",
        "Ne pas appeler les sorties « diagnostics » ou « cas de boiterie confirmés ».",
    )

    add_section(
        doc,
        "D. Résultats, volume de notifications et priorités A/B/C",
        "Présenter les nombres avec leur signification opérationnelle, pas comme une prévalence clinique.",
    )
    add_qa(
        doc,
        17,
        "Combien de notifications ont été produites?",
        "385 au total : 149 en Winter 2019, 127 en Summer 2019, 105 en Fall 2019 et 4 en Fall 2021.",
        "Ce sont des débuts d’épisodes retenus. Une même vache peut recevoir plusieurs notifications au cours du suivi.",
    )
    add_qa(
        doc,
        18,
        "385 notifications, est-ce trop?",
        "Le total brut seul ne permet pas de répondre. Il faut tenir compte de la durée, du nombre de vaches et du "
        "fait qu’une partie importante des signaux est collective.",
        "Le classement A/B/C réduit la charge de revue : 37 A à examiner en premier, 195 B à vérifier et 153 C à "
        "interpréter d’abord comme événements de contexte.",
    )
    add_qa(
        doc,
        19,
        "Que signifient A, B et C?",
        "A : prioritaire; B : à vérifier; C : événement collectif probable.",
        "A exige un signal non collectif, distinct du troupeau et un score interne d’au moins 45. C est déclenché si "
        "au moins 30 % des vaches sont alertées dans ± 1 jour ou 50 % dans ± 3 jours.",
        "Ne pas présenter A, B et C comme des niveaux léger, modéré et sévère de boiterie.",
    )
    add_qa(
        doc,
        20,
        "Le score lame_confidence est-il une probabilité de boiterie?",
        "Non. C’est un indice empirique de classement relatif utilisé pour prioriser les notifications.",
        "Il combine plusieurs familles de signaux, mais il n’a pas été calibré comme une probabilité clinique.",
    )
    add_qa(
        doc,
        21,
        "Pourquoi autant de notifications collectives en Winter 2019?",
        "89 des 149 notifications de Winter, soit 59,7 %, surviennent dans un contexte collectif probable.",
        "Le filtre n’en détermine pas la cause. Il invite à vérifier la gestion, la météo, l’exercice, l’alimentation, "
        "les déplacements ou un problème partagé de capteurs. Un regroupement important apparaît au début de février 2019.",
    )
    add_qa(
        doc,
        22,
        "Une notification C est-elle un faux positif?",
        "Pas nécessairement. Elle signifie que le signal est partagé par plusieurs vaches et qu’une cause commune "
        "doit être recherchée avant une interprétation individuelle.",
        "Une modification réelle de routine ou de conditions environnementales peut produire un signal collectif valide.",
    )
    add_qa(
        doc,
        23,
        "Pourquoi n’y a-t-il aucune catégorie A en Fall 2021?",
        "Aucune des quatre notifications n’a réuni simultanément les critères de non-collectivité, de contraste avec "
        "le troupeau et de score interne élevé.",
        "La courte fenêtre de Fall 2021 limite aussi la quantité d’information disponible pour l’interprétation.",
    )

    add_section(
        doc,
        "E. Concordance, SLS et approche actuelle",
        "C’est la section la plus sensible : distinguer cooccurrence, association exploratoire et validation clinique.",
        page_break=False,
    )
    add_qa(
        doc,
        24,
        "Qu’est-ce que le taux de concordance avec les scans?",
        "La proportion de scans comparables ayant au moins une notification de la même vache dans une fenêtre de ± 1 jour.",
        "Sur 396 scans datés, 133 sont comparables (la vache doit avoir des données IceTag à ± 1 jour) : 44,4 % en "
        "Fall 2019, 26,8 % en Winter 2019, 25,5 % en Summer 2019 et 10,0 % en Fall 2021, soit 28,6 % au total.",
        "Le filtre de couverture ramène Fall 2021 de 270 scans à 10, ce qui évite qu’un corpus de six jours écrase "
        "l’analyse.",
    )
    add_qa(
        doc,
        25,
        "Le taux de 44,4 % signifie-t-il que le modèle est précis à 44,4 %?",
        "Non. Il signifie seulement que 12 des 27 scans de Fall 2019 ont une notification temporellement proche.",
        "Les scans comportementaux ne sont pas une vérité-terrain clinique; on ne peut pas en déduire sensibilité, "
        "spécificité ou exactitude.",
        "Ne pas utiliser le mot « précision » pour ces taux. Dire « concordance temporelle » ou « cooccurrence ».",
    )
    add_qa(
        doc,
        "25b",
        "Ce niveau de concordance dépasse-t-il ce que donnerait le hasard?",
        "Non, pas de façon démontrée, et c’est un point à annoncer soi-même.",
        "Les alertes se déclenchent environ 9 fois par 100 vache-jours et la fenêtre couvre trois journées : le niveau "
        "attendu est donc déjà de 27,8 %. Sur les 127 scans couverts le jour même, on observe 38 concordances (29,9 %) "
        "contre 35,3 attendues, soit +2,1 points, non concluant (Poisson-binomial unilatéral, p = 0,324).",
        "Dire « la concordance documente une proximité temporelle, sans enrichissement démontré ». Ne jamais "
        "présenter 44,4 % comme une preuve que le système fonctionne.",
    )
    add_qa(
        doc,
        26,
        "Avons-nous des labels SLS?",
        "Oui, pour Winter 2019, avec un score daté du 12 mars permettant une validation exploratoire.",
        "Le nombre de cas reste limité et la sévérité faible : aucun SLS ≥ 3 dans la cohorte analysée.",
    )
    add_qa(
        doc,
        27,
        "La baseline IF est-elle alignée aux SLS?",
        "Non dans les données disponibles : 16 vaches évaluables, p = 0,649 et rho = 0,033.",
        "Ce résultat ne signifie pas que tout le travail est négatif. Il montre honnêtement que la baseline répond au "
        "mandat technique, mais qu’elle ne peut pas être revendiquée comme détecteur clinique validé.",
    )
    add_qa(
        doc,
        28,
        "Quel est le résultat de l’approche HYPO + instabilité + hybride?",
        "Un alignement exploratoire positif : les 3 vaches SLS ≥ 2 ont en moyenne 6,67 notifications contre 4,45 "
        "chez les 11 vaches SLS < 2; AUC = 0,924 et p = 0,031.",
        "La corrélation continue est positive, rho = 0,504, mais non significative au seuil de 5 %, p = 0,066.",
    )
    add_qa(
        doc,
        29,
        "Peut-on annoncer 92,4 % de précision?",
        "Non. L’AUC de 0,924 décrit la séparation observée entre deux groupes dans une très petite cohorte.",
        "Avec seulement 14 vaches et 3 cas positifs, l’estimation peut être instable. Ce n’est ni une précision, ni une "
        "sensibilité, ni une performance clinique généralisable.",
    )
    add_qa(
        doc,
        30,
        "Pourquoi 14 vaches pour l’approche actuelle et 16 pour la baseline?",
        "Les cohortes évaluables et les définitions de notification ne sont pas strictement identiques entre les deux analyses.",
        "L’approche actuelle exige une fenêtre principale de sept jours strictement antérieurs au SLS et les données "
        "nécessaires à ses trois composantes. La comparaison est donc exploratoire, pas un tête-à-tête clinique parfaitement apparié.",
    )
    add_qa(
        doc,
        31,
        "Le groupe Exercise peut-il expliquer le résultat SLS?",
        "Oui, c’est un confondant majeur : les trois vaches SLS ≥ 2 appartiennent au groupe Exercise.",
        "Une partie de la séparation peut refléter le protocole expérimental. Il faut davantage de cas répartis entre "
        "les traitements ou un modèle permettant d’ajuster correctement cet effet.",
    )
    add_qa(
        doc,
        32,
        "Pourquoi l’approche actuelle semble-t-elle meilleure que la baseline?",
        "Elle cible l’hypoactivité persistante et l’instabilité, puis combine ces signaux hiérarchiquement.",
        "Le résultat SLS est encourageant, mais l’échantillon est trop petit pour affirmer une supériorité clinique. "
        "Le terme correct reste « meilleur alignement exploratoire ».",
        single_run=True,
    )

    add_section(
        doc,
        "F. Livrables, valeur du travail et prochaine étape",
        "Conclure avec ce qui est livré, ce que le travail rend possible et les décisions attendues.",
    )
    add_qa(
        doc,
        33,
        "Quels dossiers faut-il réellement livrer?",
        "Le dossier Objectif1_Pipeline_detection_boiterie complet, avec README, RAPPORTS, DONNEES_TRAITEES_ALERTES, "
        "TABLEAUX_CSV, NOTES_SOW et l’annexe scientifique séparée.",
        "Le README est le point d’entrée. Le rapport Word sert à la lecture; les CSV permettent de vérifier les chiffres "
        "et d’auditer les sorties détaillées.",
    )
    add_qa(
        doc,
        34,
        "Comment vérifier qu’un chiffre de la présentation est réel?",
        "Chaque chiffre renvoie à un tableau CSV ou à une sortie détaillée conservée dans le paquet.",
        "predictions.csv porte les intervalles; alerts_only.csv porte les débuts d’épisodes; summary.csv synthétise par "
        "vache; les tableaux renforcés portent A/B/C et la concordance.",
    )
    add_qa(
        doc,
        35,
        "Qu’est-ce qui justifie le travail et sa facturation?",
        "La valeur ne se limite pas à exécuter un modèle : elle comprend l’inventaire de 2 414 fichiers, l’harmonisation "
        "de quatre expériences, le débogage, la reproductibilité, l’analyse scientifique, la traçabilité et les livrables.",
        "Le travail a aussi évité une conclusion trompeuse : il distingue clairement succès technique, limites de la baseline "
        "et signal exploratoire de l’approche actuelle.",
    )
    add_qa(
        doc,
        36,
        "Le résultat est-il positif ou négatif?",
        "Positif techniquement et contractuellement; mitigé pour la baseline clinique; encourageant mais exploratoire "
        "pour l’approche actuelle.",
        "Cette formulation est plus convaincante qu’un « succès total », car elle montre que les conclusions sont "
        "proportionnées aux preuves disponibles.",
    )
    add_qa(
        doc,
        37,
        "Que manque-t-il pour une validation clinique?",
        "Davantage de scores locomoteurs datés et répétés, davantage de cas SLS ≥ 2 et ≥ 3, des groupes moins confondus "
        "et un protocole préspécifié.",
        "Il faudra geler la version, les seuils et les fenêtres avant l’analyse, puis calculer sensibilité, spécificité, "
        "valeurs prédictives, délai d’alerte et intervalles de confiance.",
    )
    add_qa(
        doc,
        38,
        "Quelle est la prochaine décision demandée à McGill/WELL-E?",
        "Confirmer la conformité du livrable de l’objectif 1 et identifier la meilleure cohorte clinique pour la validation suivante.",
        "Il faut aussi clarifier le protocole SLS, documenter les traitements expérimentaux et vérifier les causes possibles "
        "des épisodes collectifs.",
    )
    add_qa(
        doc,
        39,
        "Si vous deviez résumer le travail en une phrase?",
        "Une chaîne commune a traité quatre saisons IceTag et produit des notifications comportementales traçables et "
        "priorisées; l’approche actuelle montre un signal SLS prometteur qui doit maintenant être validé prospectivement.",
    )

    add_questions_to_ask(doc)

    core = doc.core_properties
    core.title = "Questions et réponses probables — Objectif 1 McGill/WELL-E"
    core.subject = "Fiche de répétition pour la réunion du 30 juillet 2026"
    core.author = "Aliou Barry"
    core.keywords = "McGill, WELL-E, Objectif 1, IceTag, alertes, SLS, questions réponses"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
