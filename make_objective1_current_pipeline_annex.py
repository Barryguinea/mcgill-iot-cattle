#!/usr/bin/env python3
"""Build a clean Objective 1 annex for the current HYPO/instability/hybrid pipeline."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PROJECT = Path("/Users/alioubarry/PROJECT")
SOURCE = (
    PROJECT
    / "mcgill_iot_cattle"
    / "reports"
    / "objective1_pipeline_icetag"
    / "memoirev3_comparison"
)
DELIVERY = (
    Path("/Users/alioubarry/Desktop/Livrables_McGill_WellE")
    / "Objectif1_Pipeline_detection_boiterie"
)
ANNEX = DELIVERY / "ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride"
ANNEX_REPORTS = ANNEX / "RAPPORTS"
ANNEX_TABLES = ANNEX / "TABLEAUX_CSV"
ANNEX_ALERTS = ANNEX / "ALERTES_HYBRIDES"
ANNEX_SUMMARIES = ANNEX / "RESUMES_PAR_VACHE"
SLS_SOURCE = PROJECT / "memoirev3" / "data" / "validation" / "mcgill_sls"

BLUE = RGBColor(0, 0, 0)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)

SEASONS = ["fall_2019", "summer_2019", "winter_2019", "fall_2021"]
SEASON_LABELS = {
    "fall_2019": "Fall 2019",
    "summer_2019": "Summer 2019",
    "winter_2019": "Winter 2019",
    "fall_2021": "Fall 2021",
}

COMMON_RENAMES = {
    "old_lameness_notifs": "pipeline_initiale_notifications",
    "old_n_bins": "pipeline_initiale_intervalles",
    "old_coverage_mean": "pipeline_initiale_couverture_moyenne",
    "v3_n_bins": "pipeline_actuelle_intervalles",
    "v3_coverage_mean": "pipeline_actuelle_couverture_moyenne",
    "v3_legacy_if_lameness_notifs": "comparateur_if_notifications",
    "v3_behavioral_hypo_notifs": "hypo_notifications",
    "behavioral_warning_notifs": "hypo_notifications",
    "v3_instability_notifs": "instabilite_notifications",
    "instability_warning_notifs": "instabilite_notifications",
    "v3_hybrid_notifs": "hybride_notifications",
    "hybrid_warning_notifs": "hybride_notifications",
    "old_rate_per_100_cow_days": "pipeline_initiale_taux_100_vache_jours",
    "v3_hybrid_rate_per_100_cow_days": "hybride_taux_100_vache_jours",
    "old_vs_v3_legacy_if_exact_overlap": "pipeline_initiale_vs_comparateur_if_chevauchement_exact",
    "old_vs_hybrid_exact_overlap": "pipeline_initiale_vs_hybride_chevauchement_exact",
    "old_only_vs_hybrid": "pipeline_initiale_seule_vs_hybride",
    "hybrid_only_vs_old": "hybride_seul_vs_pipeline_initiale",
    "hypo_exact_notifs": "hypo_notifications_exactes",
    "instability_exact_notifs": "instabilite_notifications_exactes",
    "old_total": "pipeline_initiale_total",
    "hybrid_total": "hybride_total",
    "pearson_old_vs_hybrid_by_cow": "pearson_initiale_vs_hybride_par_vache",
    "spearman_old_vs_hybrid_by_cow": "spearman_initiale_vs_hybride_par_vache",
    "old_alerts": "pipeline_initiale_alertes",
    "hybrid_alerts": "hybride_alertes",
    "old_alerts_with_hybrid_nearby": "alertes_initiales_avec_hybride_proche",
    "old_alerts_with_hybrid_nearby_pct": "alertes_initiales_avec_hybride_proche_pct",
    "hybrid_alerts_with_old_nearby": "alertes_hybrides_avec_initiale_proche",
    "hybrid_alerts_with_old_nearby_pct": "alertes_hybrides_avec_initiale_proche_pct",
    "delta_hybrid_minus_old": "delta_hybride_moins_initiale",
}

ALERT_RENAMES = {
    "notif_lameness": "comparateur_if_notification",
    "lame_confidence": "comparateur_if_confiance",
    "behavioral_warning_score": "hypo_score",
    "behavioral_warning_cusum": "hypo_cusum",
    "behavioral_warning_families": "hypo_familles",
    "behavioral_warning_episode": "hypo_episode",
    "behavioral_warning_start": "hypo_debut",
    "behavioral_warning_notification": "hypo_notification",
    "instability_warning_score": "instabilite_score",
    "instability_warning_cusum": "instabilite_cusum",
    "instability_warning_families": "instabilite_familles",
    "instability_warning_episode": "instabilite_episode",
    "instability_warning_start": "instabilite_debut",
    "instability_warning_notification": "instabilite_notification",
    "hybrid_warning_score": "hybride_score",
    "hybrid_warning_episode": "hybride_episode",
    "hybrid_warning_surveillance": "hybride_surveillance",
    "hybrid_warning_sequence_start": "hybride_debut_sequence",
    "hybrid_warning_start": "hybride_debut",
    "hybrid_warning_notification": "hybride_notification",
    "hybrid_warning_type": "hybride_type",
    "hybrid_warning_priority": "hybride_priorite",
    "hybrid_warning_fusion_mode": "hybride_mode_fusion",
}

SUMMARY_RENAMES = {
    "if_anomaly_points": "comparateur_if_points_anomalie",
    "problem_points": "comparateur_if_points_probleme",
    "lameness_points": "comparateur_if_points_boiterie",
    "problem_starts": "comparateur_if_debuts_probleme",
    "lameness_starts": "comparateur_if_debuts_boiterie",
    "lameness_notifs": "comparateur_if_notifications",
    "critique_points": "points_critiques",
    "behavioral_warning_points": "hypo_points",
    "behavioral_warning_notifs": "hypo_notifications",
    "instability_warning_notifs": "instabilite_notifications",
    "hybrid_warning_notifs": "hybride_notifications",
    "coverage_mean": "couverture_moyenne",
    "coverage_min": "couverture_min",
}


def _ensure_dirs() -> None:
    for path in [ANNEX, ANNEX_REPORTS, ANNEX_TABLES, ANNEX_ALERTS, ANNEX_SUMMARIES]:
        path.mkdir(parents=True, exist_ok=True)


def _clean_frame(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    drop_cols = [c for c in out.columns if c.endswith("_path") or c == "v3_full_predictions_path"]
    if drop_cols:
        out = out.drop(columns=drop_cols)
    out = out.rename(columns=COMMON_RENAMES)
    if "hybrid_type_counts" in out.columns:
        out = out.rename(columns={"hybrid_type_counts": "types_hybrides"})
    return out


def _write_clean_csvs() -> tuple[pd.DataFrame, pd.DataFrame]:
    by_season = _clean_frame(pd.read_csv(SOURCE / "comparison_summary_by_season.csv"))
    by_season["season_label"] = by_season["season"].map(SEASON_LABELS).fillna(by_season["season"])
    cols = ["season_label"] + [c for c in by_season.columns if c != "season_label"]
    by_season = by_season[cols]
    by_season.to_csv(ANNEX_TABLES / "pipeline_actuelle_comparaison_par_saison.csv", index=False)

    by_cow = _clean_frame(pd.read_csv(SOURCE / "comparison_by_cow_all_seasons.csv"))
    by_cow["season_label"] = by_cow["season"].map(SEASON_LABELS).fillna(by_cow["season"])
    cols = ["season_label"] + [c for c in by_cow.columns if c != "season_label"]
    by_cow[cols].to_csv(ANNEX_TABLES / "pipeline_actuelle_comparaison_par_vache.csv", index=False)

    correlations = _clean_frame(pd.read_csv(SOURCE / "comparison_by_cow_correlations.csv"))
    correlations["season_label"] = correlations["season"].map(SEASON_LABELS).fillna(correlations["season"])
    correlations.to_csv(ANNEX_TABLES / "pipeline_actuelle_correlations_par_vache.csv", index=False)

    overlap = _clean_frame(pd.read_csv(SOURCE / "comparison_temporal_overlap_tolerance.csv"))
    overlap["season_label"] = overlap["season"].map(SEASON_LABELS).fillna(overlap["season"])
    overlap.to_csv(ANNEX_TABLES / "pipeline_actuelle_chevauchement_temporel.csv", index=False)

    top_cows = _clean_frame(pd.read_csv(SOURCE / "top_hybrid_cows_by_season.csv"))
    top_cows["season_label"] = top_cows["season"].map(SEASON_LABELS).fillna(top_cows["season"])
    top_cows.to_csv(ANNEX_TABLES / "pipeline_actuelle_top_vaches_hybride.csv", index=False)

    for season in SEASONS:
        season_dir = SOURCE / season
        label = SEASON_LABELS[season].lower().replace(" ", "_")

        alerts = pd.read_csv(season_dir / f"{season}_memoirev3_alerts_only.csv")
        alerts = alerts.rename(columns=ALERT_RENAMES)
        alerts["hybride_notification"] = pd.to_numeric(
            alerts["hybride_notification"], errors="coerce"
        ).fillna(0).astype(int)
        hybrid_alerts = alerts.loc[alerts["hybride_notification"].eq(1)].copy()
        keep = [
            "T",
            "Cow",
            "coverage_pct",
            "dataset_split",
            "comparateur_if_notification",
            "comparateur_if_confiance",
            "hypo_score",
            "hypo_notification",
            "instabilite_score",
            "instabilite_notification",
            "hybride_score",
            "hybride_notification",
            "hybride_type",
            "hybride_priorite",
            "hybride_mode_fusion",
        ]
        hybrid_alerts[[c for c in keep if c in hybrid_alerts.columns]].to_csv(
            ANNEX_ALERTS / f"{label}_pipeline_actuelle_alertes_hybrides.csv",
            index=False,
        )

        summary = pd.read_csv(season_dir / f"{season}_memoirev3_summary.csv")
        summary.rename(columns=SUMMARY_RENAMES).to_csv(
            ANNEX_SUMMARIES / f"{label}_pipeline_actuelle_resume_par_vache.csv",
            index=False,
        )

    total = pd.DataFrame(
        [
            {
                "pipeline_initiale_notifications": int(by_season["pipeline_initiale_notifications"].sum()),
                "comparateur_if_notifications": int(by_season["comparateur_if_notifications"].sum()),
                "hypo_notifications": int(by_season["hypo_notifications"].sum()),
                "instabilite_notifications": int(by_season["instabilite_notifications"].sum()),
                "hybride_notifications": int(by_season["hybride_notifications"].sum()),
                "pipeline_initiale_taux_100_vache_jours": round(
                    float(by_season["pipeline_initiale_notifications"].sum())
                    / float(by_season["n_bins"].sum() / 96.0)
                    * 100.0,
                    3,
                ),
                "hybride_taux_100_vache_jours": round(
                    float(by_season["hybride_notifications"].sum())
                    / float(by_season["n_bins"].sum() / 96.0)
                    * 100.0,
                    3,
                ),
            }
        ]
    )
    total.to_csv(ANNEX_TABLES / "pipeline_actuelle_synthese_totale.csv", index=False)

    sls_summary = json.loads((SLS_SOURCE / "mcgill_summary.json").read_text(encoding="utf-8"))
    primary_metric = next(
        item for item in sls_summary["primary_metrics"] if item["metric"] == "pre7_hybrid_notifs"
    )
    pd.DataFrame(
        [
            {
                "date_score_sls": sls_summary["protocol"]["endpoint"],
                "fenetre_capteur_jours": sls_summary["protocol"]["primary_window_days"],
                "unite_statistique": "vache",
                "n_vaches_evaluables": sls_summary["cohort"]["n_evaluable"],
                "n_sls_ge_2": sls_summary["cohort"]["n_sls_ge_2"],
                "n_sls_lt_2": sls_summary["cohort"]["n_sls_lt_2"],
                "auc_notifications_hybrides_pre7": primary_metric["auc"],
                "mann_whitney_p": primary_metric["mann_whitney_p"],
                "spearman_rho": primary_metric["spearman_rho"],
                "spearman_p": primary_metric["spearman_p"],
                "conclusion": "concordance observationnelle exploratoire; aucune validation diagnostique",
            }
        ]
    ).to_csv(ANNEX_TABLES / "pipeline_actuelle_validation_sls_synthese.csv", index=False)

    metrics = pd.read_csv(SLS_SOURCE / "mcgill_metrics.csv").rename(
        columns={
            "metric": "metrique",
            "variant": "mode_fusion",
            "mean_sls_ge_2": "moyenne_sls_ge_2",
            "mean_sls_lt_2": "moyenne_sls_lt_2",
            "auc": "auc",
            "mann_whitney_p": "mann_whitney_p",
            "spearman_rho": "spearman_rho",
            "spearman_p": "spearman_p",
        }
    )
    metrics.to_csv(ANNEX_TABLES / "pipeline_actuelle_validation_sls_metriques.csv", index=False)

    cohort = pd.read_csv(SLS_SOURCE / "mcgill_cohort_all_variants.csv")
    cohort = cohort.loc[cohort["variant"].eq("hierarchical")].copy()
    cohort = cohort.rename(
        columns={
            "variant": "mode_fusion",
            "cow": "vache",
            "sls_mar": "sls_12_mars",
            "sls_ge_2": "sls_ge_2",
            "treatment": "traitement",
            "pre7_hybrid_notifs": "notifications_hybrides_pre7",
            "pre7_hybrid_frac_time": "fraction_temps_hybride_pre7",
            "pre7_hybrid_score_max": "score_hybride_max_pre7",
        }
    )
    cohort.to_csv(ANNEX_TABLES / "pipeline_actuelle_validation_sls_cohorte.csv", index=False)

    exclusions = pd.read_csv(SLS_SOURCE / "mcgill_exclusions.csv")
    exclusions = exclusions.loc[exclusions["variant"].eq("hierarchical")].copy()
    exclusions = exclusions.rename(
        columns={
            "variant": "mode_fusion",
            "cow": "vache",
            "exclusion_reason": "motif_exclusion",
        }
    )
    exclusions.to_csv(ANNEX_TABLES / "pipeline_actuelle_validation_sls_exclusions.csv", index=False)
    return by_season, total


def _load_sls_result() -> dict[str, float | int]:
    summary = json.loads((SLS_SOURCE / "mcgill_summary.json").read_text(encoding="utf-8"))
    metric = next(item for item in summary["primary_metrics"] if item["metric"] == "pre7_hybrid_notifs")
    return {
        "n_evaluable": int(summary["cohort"]["n_evaluable"]),
        "n_sls_ge_2": int(summary["cohort"]["n_sls_ge_2"]),
        "n_sls_lt_2": int(summary["cohort"]["n_sls_lt_2"]),
        "auc": float(metric["auc"]),
        "mann_whitney_p": float(metric["mann_whitney_p"]),
        "spearman_rho": float(metric["spearman_rho"]),
        "spearman_p": float(metric["spearman_p"]),
    }


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        _set_cell_shading(hdr[i], "000000")
        _set_cell_text(hdr[i], header, bold=True, color=RGBColor(255, 255, 255))
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, attr, Inches(1))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, RGBColor(0, 0, 0)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Annexe - pipeline actuelle HYPO, instabilité et hybride")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Projet McGill / WELL-E - Objectif 1 - Analyse complémentaire")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Date: 13 juillet 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED


def _build_docx(by_season: pd.DataFrame, total: pd.DataFrame) -> Path:
    path = ANNEX_REPORTS / "Annexe_pipeline_actuelle_HYPO_instabilite_hybride.docx"
    doc = Document()
    _style_doc(doc)
    _add_title(doc)
    sls = _load_sls_result()

    doc.add_heading("Message principal", level=1)
    doc.add_paragraph(
        "La pipeline actuelle HYPO + instabilité + hybride a été appliquée aux mêmes quatre corpus IceTag que la pipeline initiale. "
        "Elle produit une détection plus large des signaux comportementaux compatibles avec une perturbation locomotrice. "
        "Des scores SLS synchronisés existent pour Winter 2019 et permettent une concordance exploratoire limitée."
    )

    doc.add_heading("Résultat global", level=1)
    row = total.iloc[0]
    _add_table(
        doc,
        ["Sortie", "Notifications", "Taux / 100 vache-jours"],
        [
            [
                "Pipeline initiale IF + règles",
                int(row["pipeline_initiale_notifications"]),
                f'{row["pipeline_initiale_taux_100_vache_jours"]:.2f}',
            ],
            ["HYPO", int(row["hypo_notifications"]), "-"],
            ["Instabilité", int(row["instabilite_notifications"]), "-"],
            [
                "Hybride final",
                int(row["hybride_notifications"]),
                f'{row["hybride_taux_100_vache_jours"]:.2f}',
            ],
        ],
    )
    doc.add_paragraph(
        "Lecture: le comparateur IF reproduit les 385 notifications initiales, tandis que la sortie hybride finale produit 1179 notifications. "
        "La hausse reflète une définition plus large de l’alerte comportementale, pas une preuve automatique de meilleure détection clinique."
    )

    doc.add_heading("Comparaison par saison", level=1)
    rows = []
    for _, r in by_season.iterrows():
        rows.append(
            [
                r["season_label"],
                int(r["pipeline_initiale_notifications"]),
                int(r["hypo_notifications"]),
                int(r["instabilite_notifications"]),
                int(r["hybride_notifications"]),
                f'{float(r["hybride_taux_100_vache_jours"]):.2f}',
            ]
        )
    _add_table(
        doc,
        ["Saison", "Initiale", "HYPO", "Instabilité", "Hybride", "Hybride / 100 v-j"],
        rows,
    )

    doc.add_heading("Concordance exploratoire avec les scores SLS", level=1)
    _add_table(
        doc,
        ["Cohorte", "Mesure principale", "Résultat", "Portée"],
        [
            [
                f'{sls["n_evaluable"]} vaches; {sls["n_sls_ge_2"]} avec SLS >= 2',
                "Notifications hybrides dans les 7 jours avant le score",
                f'AUC = {sls["auc"]:.3f}; p = {sls["mann_whitney_p"]:.3f}',
                "Concordance observationnelle exploratoire",
            ]
        ],
    )
    doc.add_paragraph(
        "Le SLS n'a servi ni à entraîner la détection ni à fixer les seuils. Le résultat est encourageant, mais la cohorte ne "
        "contient que trois vaches SLS >= 2 et le traitement Exercise est confondu avec le statut SLS. Il ne s'agit donc pas "
        "d'une validation diagnostique."
    )

    doc.add_heading("Interprétation", level=1)
    for text in [
        "HYPO cible les baisses comportementales persistantes.",
        "Instabilité cible les changements irréguliers ou abrupts de comportement.",
        "Hybride fusionne les deux familles de signaux pour produire la sortie principale de la pipeline actuelle.",
        "Sur McGill, la pipeline actuelle est utile comme annexe comparative, mais elle doit être recalibrée avant d’être substituée au livrable principal.",
    ]:
        doc.add_paragraph(text, style=None)

    doc.add_heading("Limite scientifique", level=1)
    doc.add_paragraph(
        "Le corpus McGill fournit des scores SLS synchronisés sur une sous-cohorte, mais pas assez de cas positifs pour calculer une sensibilité ou une spécificité robuste. "
        "Les alertes doivent donc être décrites comme des signaux comportementaux compatibles avec une perturbation locomotrice et à vérifier."
    )

    doc.add_heading("Fichiers inclus dans l’annexe", level=1)
    for text in [
        "TABLEAUX_CSV/pipeline_actuelle_comparaison_par_saison.csv",
        "TABLEAUX_CSV/pipeline_actuelle_comparaison_par_vache.csv",
        "TABLEAUX_CSV/pipeline_actuelle_chevauchement_temporel.csv",
        "TABLEAUX_CSV/pipeline_actuelle_validation_sls_synthese.csv",
        "TABLEAUX_CSV/pipeline_actuelle_validation_sls_metriques.csv",
        "TABLEAUX_CSV/pipeline_actuelle_validation_sls_cohorte.csv",
        "TABLEAUX_CSV/pipeline_actuelle_validation_sls_exclusions.csv",
        "ALERTES_HYBRIDES/*_pipeline_actuelle_alertes_hybrides.csv",
        "RESUMES_PAR_VACHE/*_pipeline_actuelle_resume_par_vache.csv",
    ]:
        doc.add_paragraph(text)

    doc.save(path)
    return path


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;"), style)


def _build_pdf(by_season: pd.DataFrame, total: pd.DataFrame) -> Path:
    path = ANNEX_REPORTS / "Annexe_pipeline_actuelle_HYPO_instabilite_hybride.pdf"
    sls = _load_sls_result()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#000000"),
        spaceAfter=8,
        alignment=0,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        textColor=colors.HexColor("#000000"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#000000"),
        spaceAfter=6,
    )

    story = [
        _p("Annexe - pipeline actuelle HYPO, instabilité et hybride", title),
        _p("Projet McGill / WELL-E - Objectif 1 - Analyse complémentaire<br/>Date: 13 juillet 2026", body),
        Spacer(1, 8),
        _p("Message principal", h1),
        _p(
            "La pipeline actuelle HYPO + instabilité + hybride a été appliquée aux mêmes quatre corpus IceTag que la pipeline initiale. "
            "Elle produit une détection plus large des signaux comportementaux compatibles avec une perturbation locomotrice. "
            "Des scores SLS synchronisés existent pour Winter 2019 et permettent une concordance exploratoire limitée.",
            body,
        ),
        _p("Résultat global", h1),
    ]

    row = total.iloc[0]
    table_data = [
        ["Sortie", "Notifications", "Taux / 100 vache-jours"],
        ["Pipeline initiale IF + règles", int(row["pipeline_initiale_notifications"]), f'{row["pipeline_initiale_taux_100_vache_jours"]:.2f}'],
        ["HYPO", int(row["hypo_notifications"]), "-"],
        ["Instabilité", int(row["instabilite_notifications"]), "-"],
        ["Hybride final", int(row["hybride_notifications"]), f'{row["hybride_taux_100_vache_jours"]:.2f}'],
    ]
    story.append(_pdf_table(table_data, [3.0 * inch, 1.3 * inch, 1.7 * inch]))
    story.append(
        _p(
            "Lecture: le comparateur IF reproduit les 385 notifications initiales, tandis que la sortie hybride finale produit 1179 notifications. "
            "La hausse reflète une définition plus large de l’alerte comportementale, pas une preuve automatique de meilleure détection clinique.",
            body,
        )
    )
    story.append(_p("Comparaison par saison", h1))
    season_data = [["Saison", "Initiale", "HYPO", "Instabilité", "Hybride", "Hybride / 100 v-j"]]
    for _, r in by_season.iterrows():
        season_data.append(
            [
                r["season_label"],
                int(r["pipeline_initiale_notifications"]),
                int(r["hypo_notifications"]),
                int(r["instabilite_notifications"]),
                int(r["hybride_notifications"]),
                f'{float(r["hybride_taux_100_vache_jours"]):.2f}',
            ]
        )
    story.append(_pdf_table(season_data, [1.25 * inch, 0.75 * inch, 0.75 * inch, 0.9 * inch, 0.85 * inch, 1.1 * inch]))
    story.append(_p("Concordance exploratoire avec les scores SLS", h1))
    story.append(
        _pdf_table(
            [
                ["Cohorte", "Fenêtre", "Résultat", "Portée"],
                [
                    f'{sls["n_evaluable"]} évaluables; {sls["n_sls_ge_2"]} avec SLS >= 2',
                    "7 jours avant le score",
                    f'AUC = {sls["auc"]:.3f}; p = {sls["mann_whitney_p"]:.3f}',
                    "Concordance exploratoire",
                ],
            ],
            [2.05 * inch, 1.25 * inch, 1.45 * inch, 1.55 * inch],
        )
    )
    story.append(
        _p(
            "Le SLS n'a servi ni à entraîner la détection ni à fixer les seuils. Le résultat est encourageant, mais la cohorte "
            "ne contient que trois vaches SLS >= 2 et le traitement Exercise est confondu avec le statut SLS. Il ne s'agit pas "
            "d'une validation diagnostique.",
            body,
        )
    )
    story.append(
        KeepTogether(
            [
            _p("Interprétation", h1),
            _p("HYPO cible les baisses comportementales persistantes. Instabilité cible les changements irréguliers ou abrupts. Hybride fusionne les deux familles de signaux pour produire la sortie principale de la pipeline actuelle.", body),
            _p("Sur McGill, cette pipeline est utile comme annexe comparative, mais elle doit être recalibrée avant d’être substituée au livrable principal.", body),
            ]
        )
    )
    story.extend(
        [
            _p("Limite scientifique", h1),
            _p("Le corpus McGill fournit des scores SLS synchronisés sur une sous-cohorte, mais pas assez de cas positifs pour calculer une sensibilité ou une spécificité robuste. Les alertes doivent donc être décrites comme des signaux comportementaux compatibles avec une perturbation locomotrice et à vérifier.", body),
            _p("Fichiers inclus", h1),
            _p("TABLEAUX_CSV/ contient les comparaisons, la validation SLS exploratoire et les exclusions; ALERTES_HYBRIDES/ et RESUMES_PAR_VACHE/ contiennent les sorties par saison.", body),
        ]
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)
    return path


def _pdf_table(data: list[list[object]], col_widths: list[float]) -> Table:
    table = Table(data, colWidths=col_widths, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#000000")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _write_readme() -> None:
    text = """Annexe - pipeline actuelle HYPO, instabilité et hybride

Objet:
Cette annexe applique la pipeline actuelle HYPO + instabilité + hybride aux quatre corpus IceTag McGill déjà traités pour l'Objectif 1.

À lire en premier:
- RAPPORTS/Annexe_pipeline_actuelle_HYPO_instabilite_hybride.docx

Tableaux principaux:
- TABLEAUX_CSV/pipeline_actuelle_synthese_totale.csv
- TABLEAUX_CSV/pipeline_actuelle_comparaison_par_saison.csv
- TABLEAUX_CSV/pipeline_actuelle_comparaison_par_vache.csv
- TABLEAUX_CSV/pipeline_actuelle_correlations_par_vache.csv
- TABLEAUX_CSV/pipeline_actuelle_chevauchement_temporel.csv
- TABLEAUX_CSV/pipeline_actuelle_top_vaches_hybride.csv
- TABLEAUX_CSV/pipeline_actuelle_validation_sls_synthese.csv
- TABLEAUX_CSV/pipeline_actuelle_validation_sls_metriques.csv
- TABLEAUX_CSV/pipeline_actuelle_validation_sls_cohorte.csv
- TABLEAUX_CSV/pipeline_actuelle_validation_sls_exclusions.csv

Sorties par saison:
- ALERTES_HYBRIDES/: alertes hybrides finales par saison.
- RESUMES_PAR_VACHE/: résumés par vache et par saison.

Interprétation:
La pipeline actuelle produit une détection plus large des signaux comportementaux compatibles avec une perturbation locomotrice. Les scores SLS synchronisés de Winter 2019 montrent une concordance exploratoire encourageante sur 14 vaches évaluables, dont seulement 3 avec SLS >= 2. Cette annexe reste comparative et ne constitue pas une validation diagnostique.
"""
    (ANNEX / "README_annexe_pipeline_actuelle.txt").write_text(text, encoding="utf-8")


def _update_root_readme() -> None:
    path = DELIVERY / "README_livraison_objectif1.txt"
    text = path.read_text(encoding="utf-8")
    if "- ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride/" not in text:
        text = text.replace(
            "- NOTES_SOW/\n",
            "- NOTES_SOW/\n- ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride/ (annexe complémentaire)\n",
        )
    note = (
        "\nAnnexe complémentaire:\n"
        "L'annexe HYPO + instabilité + hybride applique la pipeline actuelle aux mêmes corpus IceTag. "
        "Elle inclut une concordance SLS exploratoire limitée; le livrable principal reste la pipeline IF + règles renforcée par le contexte troupeau.\n"
    )
    if "Annexe complémentaire:" not in text:
        text += note
    path.write_text(text, encoding="utf-8")


def main() -> None:
    _ensure_dirs()
    by_season, total = _write_clean_csvs()
    docx_path = _build_docx(by_season, total)
    _write_readme()
    _update_root_readme()
    print(f"Annexe créée: {ANNEX}")
    print(f"Word: {docx_path}")


if __name__ == "__main__":
    main()
