from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPORT = Path(
    "/Users/alioubarry/Desktop/Livrables_McGill_WellE/"
    "Objectif1_Pipeline_detection_boiterie/RAPPORTS/"
    "Objectif1_rapport_livraison.docx"
)


def set_cell_text(cell, value):
    paragraph = cell.paragraphs[0]
    paragraph.text = value
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


document = Document(REPORT)

# Table 4 is the compact concordance table in section 4.
table = document.tables[4]
rows = [
    ["Expérience", "Scans couverts IceTag (±1 j)", "Avec alerte (±1 j)", "Taux descriptif"],
    ["Winter 2019", "41", "11", "26.8%"],
    ["Summer 2019", "55", "14", "25.5%"],
    ["Fall 2019", "27", "12", "44.4%"],
    ["Fall 2021", "10", "1", "10.0%"],
]

for row, values in zip(table.rows, rows):
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)

interpretation = document.paragraphs[15]
interpretation.text = (
    "Les fichiers contiennent 396 scans datés, mais seulement 133 sont réellement "
    "comparables parce que la même vache dispose de données IceTag dans une fenêtre "
    "de ±1 jour. Parmi eux, 38 coïncident avec une alerte, soit 28.6%. Dans le contrôle "
    "de sensibilité limité aux 127 scans couverts le jour même, 38 concordances sont "
    "observées (29.9%), contre 35.31 attendues (27.8%) d'après la fréquence d'alerte "
    "propre à chaque vache. L'écart de +2.1 points n'est pas statistiquement concluant "
    "(test unilatéral de Poisson-binomial, p = 0.324). Ces taux décrivent une proximité "
    "temporelle; ils ne constituent ni une sensibilité, ni une spécificité, ni une "
    "validation clinique. Le faible dénominateur de Fall 2021 (10 scans couverts) "
    "impose une prudence particulière."
)

document.save(REPORT)
