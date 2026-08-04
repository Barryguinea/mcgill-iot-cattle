from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PROJECT = Path("/Users/alioubarry/PROJECT/mcgill_iot_cattle")
SOURCE = PROJECT / "reports" / "objective1_pipeline_icetag"
DESKTOP = Path("/Users/alioubarry/Desktop/Livrables_McGill_WellE")
DEST = DESKTOP / "Objectif1_Pipeline_detection_boiterie"

READ_ME = DEST / "NOTES_SOW"
PDF_DIR = DEST / "RAPPORTS"
WORD_DIR = DEST / "RAPPORTS"
CSV_DIR = DEST / "TABLEAUX_CSV"
DATA_DIR = DEST / "DONNEES_TRAITEES_ALERTES"
ARCHIVE = DESKTOP / "Archive_interne_Objectif1_ne_pas_envoyer"
ROOT_README = DEST / "README_livraison_objectif1.txt"
SLS_CURRENT_DIR = PROJECT.parent / "memoirev3" / "data" / "validation" / "mcgill_sls"

SLS_INITIAL = {
    "n_evaluable": 16,
    "n_sls_ge_2": 5,
    "mann_whitney_p": 0.649,
    "spearman_rho": 0.033,
}


def _archive_path(name: str) -> Path:
    target = ARCHIVE / name
    if not target.exists():
        return target
    i = 2
    while (ARCHIVE / f"{name}_{i}").exists():
        i += 1
    return ARCHIVE / f"{name}_{i}"


def _archive_existing_visible_noise() -> None:
    ARCHIVE.mkdir(parents=True, exist_ok=True)

    for folder, archive_name in [
        (DEST / "RAPPORTS", "anciens_rapports"),
        (DEST / "RAPPORTS_PDF", "anciens_pdfs"),
        (DEST / "RAPPORTS_WORD", "anciens_word"),
        (DEST / "A_LIRE", "anciennes_notes_detaillees"),
        (DEST / "NOTES_SOW", "anciennes_notes_sow"),
        (DEST / "ARCHIVE_INTERNE", "ancienne_archive_dans_livrable"),
    ]:
        if folder.exists():
            shutil.move(str(folder), str(_archive_path(archive_name)))

    for generated_dir in [CSV_DIR, DATA_DIR]:
        if generated_dir.exists():
            shutil.rmtree(generated_dir)

    internal_comp = DEST / "analyses_et_resultats" / "memoirev3_comparison"
    if internal_comp.exists():
        target = _archive_path("comparaison_pipeline_alternative_interne_ne_pas_envoyer")
        shutil.move(str(internal_comp), str(target))

    detailed = DEST / "analyses_et_resultats"
    if detailed.exists():
        target = _archive_path("analyses_detaillees_internes_ne_pas_envoyer")
        shutil.move(str(detailed), str(target))

    code = DEST / "code"
    if code.exists():
        target = _archive_path("code_reproductible_interne_ne_pas_envoyer")
        shutil.move(str(code), str(target))

    for ds_store in DEST.rglob(".DS_Store"):
        ds_store.unlink(missing_ok=True)


def _copy_csv_outputs() -> None:
    CSV_DIR.mkdir(parents=True, exist_ok=True)
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    reinforced = SOURCE / "renforcement_scientifique"
    files = [
        SOURCE / "objective1_multi_season_summary.csv",
        reinforced / "objective1_reinforced_summary_by_season.csv",
        reinforced / "objective1_reinforced_summary_by_confidence.csv",
        reinforced / "objective1_reinforced_summary_by_cow.csv",
        reinforced / "objective1_collective_days.csv",
        SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv",
        SOURCE / "tache1_2_concordance" / "table_concordance.csv",
    ]
    for src in files:
        if src.exists():
            shutil.copy2(src, CSV_DIR / src.name)

    processed_files = [
        SOURCE / "fall_2019_pipeline_predictions.csv",
        SOURCE / "fall_2019_pipeline_alerts_only.csv",
        SOURCE / "fall_2019_pipeline_summary.csv",
        SOURCE / "summer_2019_pipeline_predictions.csv",
        SOURCE / "summer_2019_pipeline_alerts_only.csv",
        SOURCE / "summer_2019_pipeline_summary.csv",
        SOURCE / "winter_2019_pipeline_predictions.csv",
        SOURCE / "winter_2019_pipeline_alerts_only.csv",
        SOURCE / "winter_2019_pipeline_summary.csv",
        SOURCE / "fall_2021_pipeline_predictions.csv",
        SOURCE / "fall_2021_pipeline_alerts_only.csv",
        SOURCE / "fall_2021_pipeline_summary.csv",
        reinforced / "objective1_reinforced_alerts.csv",
    ]
    for src in processed_files:
        if src.exists():
            shutil.copy2(src, DATA_DIR / src.name)

    _normalize_visible_concordance_csvs()

    READ_ME.mkdir(parents=True, exist_ok=True)


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def _fmt_int(value) -> str:
    return f"{int(value):,}".replace(",", " ")


def _fmt_pct(value) -> str:
    return f"{float(value) * 100:.1f}%"


SEASON_LABELS = {
    "winter_2019": "Winter 2019",
    "summer_2019": "Summer 2019",
    "fall_2019": "Fall 2019",
    "fall_2021": "Fall 2021",
}

SEASON_ORDER = ["winter_2019", "summer_2019", "fall_2019", "fall_2021"]

SEASON_PERIODS = {
    "winter_2019": "2019-01-16 to 2019-04-17",
    "summer_2019": "2019-06-05 to 2019-09-06",
    "fall_2019": "2019-11-11 to 2019-12-14",
    "fall_2021": "2021-11-30 to 2021-12-06",
}

SEASON_PERIODS_SHORT = {
    "winter_2019": "Jan-Apr 2019",
    "summer_2019": "Jun-Sep 2019",
    "fall_2019": "Nov-Dec 2019",
    "fall_2021": "Nov 30-Dec 6, 2021",
}

CONFIDENCE_ORDER = {
    "A_individuelle_prioritaire": 0,
    "B_individuelle_a_verifier": 1,
    "C_probable_evenement_collectif": 2,
    "D_qualite_ou_contexte_insuffisant": 3,
}

CONCORDANCE_LABELS = {
    "Winter2019": "Winter 2019",
    "Summer2019": "Summer 2019",
    "Fall2019": "Fall 2019",
    "Fall 2021": "Fall 2021",
}


def _short_table(data: list[list[str]], widths: list[float] | None = None) -> Table:
    header_style = ParagraphStyle(
        "TableHeader",
        fontName="Helvetica-Bold",
        fontSize=8.2,
        leading=9.8,
        textColor=colors.white,
    )
    body_style = ParagraphStyle(
        "TableBody",
        fontName="Helvetica",
        fontSize=8.2,
        leading=9.8,
        textColor=colors.black,
    )
    wrapped = []
    for row_index, row in enumerate(data):
        style = header_style if row_index == 0 else body_style
        wrapped.append([Paragraph(str(cell), style) for cell in row])

    table = Table(wrapped, colWidths=widths, hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#000000")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7B7B7")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F6FA")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _normalize_visible_concordance_csvs() -> None:
    for name in ["concordance_par_experience.csv", "table_concordance.csv"]:
        path = CSV_DIR / name
        if not path.exists():
            continue
        df = pd.read_csv(path)
        if "Experiment" in df.columns:
            df["Experiment"] = df["Experiment"].astype(str).replace(CONCORDANCE_LABELS)
            season_order = {SEASON_LABELS[s]: i for i, s in enumerate(SEASON_ORDER)}
            df["_season_order"] = df["Experiment"].map(season_order).fillna(99)
            df = df.sort_values(["_season_order", "Experiment"]).drop(columns=["_season_order"])
        df.to_csv(path, index=False)


def _callout(text: str, style: ParagraphStyle, background: str = "#F2F2F2") -> Table:
    table = Table([[Paragraph(text, style)]], colWidths=[7.0 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(background)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B7B7B7")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def _footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(0.55 * inch, 0.35 * inch, "Projet McGill / WELL-E - Objectif 1")
    canvas.drawRightString(7.95 * inch, 0.35 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _set_docx_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_docx_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _set_docx_table_borders(table, color="B7B7B7", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_docx_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def _set_docx_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_docx_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def _set_docx_text(cell, text: str, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255) if white else RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_docx_cell_margins(cell)


def _add_docx_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_docx_table_width(table)
    _set_docx_table_borders(table)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            is_header = row_idx == 0
            _set_docx_text(cell, value, bold=is_header, white=is_header)
            if is_header:
                _set_docx_cell_shading(cell, "000000")
            elif row_idx % 2 == 0:
                _set_docx_cell_shading(cell, "F2F2F2")
        _set_docx_row_cant_split(table.rows[row_idx])
    _set_docx_repeat_header(table.rows[0])
    doc.add_paragraph()


def _configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(80, 80, 80)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "Projet McGill / WELL-E - Objectif 1"
    footer.style = doc.styles["Footer"]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_docx_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_docx_table_width(table)
    _set_docx_table_borders(table, color="B7B7B7")
    cell = table.cell(0, 0)
    _set_docx_cell_shading(cell, "F2F2F2")
    _set_docx_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    doc.add_paragraph()


_DECIMAL_RE = re.compile(r"(?<=\d)\.(?=\d)")


def _frenchify_runs(paragraphs) -> None:
    """Remplace le point decimal par une virgule dans des paragraphes donnes."""
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if "." in run.text:
                run.text = _DECIMAL_RE.sub(",", run.text)


def _apply_french_decimals(doc: Document) -> None:
    """Convention francaise pour la virgule decimale.

    Le tableau de correspondance SOW est exclu : ses cellules commencent par un
    numero de tache (1.1, 1.2) qui n'est pas un nombre decimal.
    """
    _frenchify_runs(doc.paragraphs)
    for table in doc.tables:
        entetes = " ".join(cell.text for cell in table.rows[0].cells)
        if "Livrable SOW" in entetes:
            continue
        for row in table.rows:
            for cell in row.cells:
                _frenchify_runs(cell.paragraphs)


def _clean_docx_metadata(doc: Document) -> None:
    """Retire les mentions d'outil laissees dans les proprietes du document."""
    props = doc.core_properties
    props.comments = ""
    props.category = ""
    props.keywords = ""
    props.last_modified_by = "Aliou Barry"


def _add_docx_run(paragraph, text: str, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_docx_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def _ordered_summary(summary: pd.DataFrame) -> pd.DataFrame:
    out = summary.copy()
    out["season_order"] = out["season"].map({s: i for i, s in enumerate(SEASON_ORDER)}).fillna(99)
    return out.sort_values(["season_order", "season"]).drop(columns=["season_order"])


def _ordered_confidence(confidence: pd.DataFrame) -> pd.DataFrame:
    out = confidence.copy()
    out["season_order"] = out["season"].map({s: i for i, s in enumerate(SEASON_ORDER)}).fillna(99)
    out["level_order"] = out["reinforced_confidence_level"].map(CONFIDENCE_ORDER).fillna(99)
    return out.sort_values(["season_order", "level_order"]).drop(columns=["season_order", "level_order"])


def _concordance_rows(concordance: pd.DataFrame) -> list[list[str]]:
    rows = [["Expérience", "Scans alignés", "Avec alerte", "Taux"]]
    items = []
    label_order = {SEASON_LABELS[s]: i for i, s in enumerate(SEASON_ORDER)}
    for _, row in concordance.iterrows():
        exp = str(row.get("Experiment", row.get("experience", "")))
        exp = CONCORDANCE_LABELS.get(exp, exp)
        scans = row.get("n_scans", row.get("scans", ""))
        alerted = row.get("scans_avec_alerte", row.get("n_scans_avec_alerte", ""))
        pct = row.get("taux_concurrence_%", row.get("pct_scans_avec_alerte", ""))
        items.append((label_order.get(exp, 99), [exp, str(int(scans)), str(int(alerted)), f"{float(pct):.1f}%"]))
    rows.extend(item for _, item in sorted(items, key=lambda x: x[0]))
    return rows


def _load_current_sls_summary() -> dict[str, object]:
    import json

    summary = json.loads((SLS_CURRENT_DIR / "mcgill_summary.json").read_text(encoding="utf-8"))
    metric = next(
        item for item in summary["primary_metrics"] if item["metric"] == "pre7_hybrid_notifs"
    )
    return {
        "n_evaluable": int(summary["cohort"]["n_evaluable"]),
        "n_sls_ge_2": int(summary["cohort"]["n_sls_ge_2"]),
        "n_sls_lt_2": int(summary["cohort"]["n_sls_lt_2"]),
        "auc": float(metric["auc"]),
        "mann_whitney_p": float(metric["mann_whitney_p"]),
        "spearman_rho": float(metric["spearman_rho"]),
        "spearman_p": float(metric["spearman_p"]),
    }


def _display_n_cows(season: str, value: object) -> str:
    n = int(value)
    if season == "fall_2021":
        return f"{n} traités / 8 complets"
    return str(n)


def _season_interpretation(season: str) -> str:
    notes = {
        "winter_2019": (
            "Saison la plus marquee par un contexte collectif: 89 alertes sur 149 sont classees "
            "comme evenement collectif probable. L'episode du debut fevrier touche une grande "
            "partie du troupeau et ne doit pas etre lu comme une multiplication de cas individuels."
        ),
        "summer_2019": (
            "Volume important d'alertes, avec une part collective plus faible que Winter 2019. "
            "Les 16 alertes A constituent les signaux individuels les plus prioritaires a examiner."
        ),
        "fall_2019": (
            "Saison riche en signaux, mais avec plusieurs jours collectifs. La reclassification "
            "permet de separer les alertes individuelles a verifier des episodes partages par le troupeau."
        ),
        "fall_2021": (
            "Fenetre courte et peu d'alertes. Les resultats sont utiles pour tester la reproductibilite "
            "technique, mais trop limites pour une interpretation scientifique forte."
        ),
    }
    return notes.get(season, "Saison traitee par le pipeline et incluse dans les tableaux de synthese.")


def _build_pdf() -> Path:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = PDF_DIR / "Objectif1_rapport_detaille_lisible.pdf"

    summary = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_season.csv")
    summary = _ordered_summary(summary)
    confidence = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_confidence.csv")
    confidence = _ordered_confidence(confidence)
    cows = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_cow.csv")
    collective_days = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_collective_days.csv")
    concordance = pd.read_csv(SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv")

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleClean",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#000000"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Clean",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#000000"),
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Clean",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#000000"),
            spaceBefore=9,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyClean",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.2,
            leading=14.5,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallClean",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#555555"),
        )
    )

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="Objectif 1 - Rapport detaille lisible",
    )

    story = []
    story.append(Spacer(1, 1.0 * inch))
    story.append(_p("Objectif 1", styles["TitleClean"]))
    story.append(_p("Rapport detaille lisible", styles["TitleClean"]))
    story.append(Spacer(1, 0.15 * inch))
    story.append(_p("Application et evaluation du pipeline de detection sur les donnees IceTag McGill", styles["BodyClean"]))
    story.append(_p(f"Version nettoyee du {date.today().isoformat()}", styles["SmallClean"]))
    story.append(Spacer(1, 0.35 * inch))
    total_alerts = int(summary["initial_notifications"].sum())
    total_a = int(summary["A_individuelle_prioritaire"].sum())
    total_b = int(summary["B_individuelle_a_verifier"].sum())
    total_c = int(summary["C_probable_evenement_collectif"].sum())
    story.append(
        _callout(
            f"<b>Conclusion courte.</b> Les quatre corpus ont ete traites. Le pipeline produit "
            f"{total_alerts} notifications, maintenant requalifiees en {total_a} alertes A, "
            f"{total_b} alertes B et {total_c} alertes C. Ces sorties sont defendables comme "
            "signaux comportementaux a prioriser, pas comme diagnostics cliniques confirmes.",
            styles["BodyClean"],
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        _short_table(
            [
                ["Element", "Statut"],
                ["Corpus IceTag", "4 saisons traitees"],
                ["Pipeline", "Application complete sur intervalles de 15 min"],
                ["Renforcement", "Normalisation troupeau + filtre collectif + niveaux A/B/C/D"],
                ["Validation clinique", "Non revendiquee sans labels synchrones robustes"],
                ["Livrables", "1 PDF lisible + tableaux CSV reproductibles"],
            ],
            [2.0 * inch, 4.9 * inch],
        )
    )

    story.append(PageBreak())
    story.append(_p("1. Perimetre du livrable", styles["H1Clean"]))
    story.append(
        _p(
            "Ce rapport documente ce qui a ete fait pour l'Objectif 1: appliquer le pipeline de "
            "detection sur les donnees accelerometriques McGill, analyser les alertes obtenues, "
            "les comparer aux observations disponibles et renforcer leur interpretation par un "
            "contexte troupeau.",
            styles["BodyClean"],
        )
    )
    story.append(
        _p(
            "Le livrable ne pretend pas mesurer une sensibilite ou une specificite clinique. Pour "
            "cela, il faudrait des labels de boiterie synchrones, dates et suffisamment nets. La "
            "formulation defendable est donc: detection et priorisation de signaux comportementaux "
            "compatibles avec une perturbation locomotrice.",
            styles["BodyClean"],
        )
    )
    story.append(_p("Travaux couverts", styles["H2Clean"]))
    for item in [
        "Conversion et harmonisation des donnees IceTag par saison.",
        "Construction de bins de 15 minutes et verification de la couverture.",
        "Application du pipeline de detection sur chaque vache et chaque saison.",
        "Compilation des notifications et syntheses multi-saisons.",
        "Comparaison temporelle avec les observations comportementales disponibles.",
        "Ajout d'une normalisation troupeau et d'un filtre d'evenements collectifs.",
        "Reclassification des alertes en niveaux de confiance A, B, C et D.",
    ]:
        story.append(_p(f"- {item}", styles["BodyClean"]))

    story.append(_p("2. Donnees traitees", styles["H1Clean"]))
    data_rows = [["Saison", "Vaches", "Bins 15 min", "Periode", "Couverture"]]
    for _, row in summary.iterrows():
        data_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["n_cows"])),
                _fmt_int(row["n_intervals"]),
                SEASON_PERIODS.get(row["season"], ""),
                f"{row['mean_coverage_pct']:.1f}%",
            ]
        )
    story.append(_short_table(data_rows, [1.2 * inch, 0.7 * inch, 1.05 * inch, 1.9 * inch, 0.85 * inch]))
    story.append(
        _p(
            "Les quatre jeux de donnees ont une couverture suffisante pour produire des alertes. "
            "Fall 2021 est le corpus le plus court; il est donc utile surtout pour verifier que la "
            "chaine de traitement s'execute correctement sur ce format.",
            styles["BodyClean"],
        )
    )

    story.append(PageBreak())
    story.append(_p("3. Pipeline applique", styles["H1Clean"]))
    story.append(
        _p(
            "La chaine appliquee reste volontairement simple a auditer: elle part des mesures "
            "IceTag, reconstruit des indicateurs temporels par vache, detecte les anomalies, puis "
            "ajoute un niveau d'interpretation troupeau pour eviter de confondre un evenement collectif "
            "avec plusieurs alertes individuelles independantes.",
            styles["BodyClean"],
        )
    )
    pipeline_rows = [
        ["Etape", "Role", "Sortie"],
        ["1. Lecture", "Importer les fichiers saisonniers harmonises", "Table par vache et timestamp"],
        ["2. Binning", "Regrouper les donnees en intervalles de 15 min", "Series regulieres comparables"],
        ["3. Features", "Calculer activite, variations et contexte temporel", "Variables de detection"],
        ["4. Detection", "Identifier les deviations individuelles", "Notifications brutes"],
        ["5. Contexte troupeau", "Comparer chaque vache au troupeau au meme moment", "Signal propre vs collectif"],
        ["6. Reclassification", "Attribuer un niveau A/B/C/D", "Tableaux interpretables"],
    ]
    story.append(_short_table(pipeline_rows, [1.1 * inch, 3.2 * inch, 2.2 * inch]))
    story.append(_p("Definition des niveaux de confiance", styles["H2Clean"]))
    level_rows = [
        ["Niveau", "Interpretation"],
        ["A", "Alerte individuelle prioritaire: non collective, bonne couverture, deviation propre a la vache."],
        ["B", "Alerte individuelle a verifier: signal non collectif mais support troupeau moins net."],
        ["C", "Probable evenement collectif: plusieurs vaches alertees dans une fenetre courte."],
        ["D", "Qualite ou contexte insuffisant: alerte conservee mais interpretation faible."],
    ]
    story.append(_short_table(level_rows, [0.75 * inch, 5.85 * inch]))

    story.append(PageBreak())
    story.append(_p("4. Resultats globaux", styles["H1Clean"]))
    alert_rows = [["Saison", "Total", "A prioritaire", "B a verifier", "C collectif"]]
    for _, row in summary.iterrows():
        alert_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["initial_notifications"])),
                str(int(row["A_individuelle_prioritaire"])),
                str(int(row["B_individuelle_a_verifier"])),
                str(int(row["C_probable_evenement_collectif"])),
            ]
        )
    story.append(_short_table(alert_rows, [1.25 * inch, 0.7 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch]))
    story.append(
        _p(
            "La reclassification change la lecture de l'Objectif 1. Le nombre brut d'alertes reste "
            "important, mais toutes les alertes n'ont pas la meme priorite. Les alertes C ne sont pas "
            "ignorees; elles sont conservees comme signaux de contexte collectif et sorties de la liste "
            "des suspicions individuelles prioritaires.",
            styles["BodyClean"],
        )
    )
    story.append(_p("Synthese par niveau", styles["H2Clean"]))
    conf_rows = [["Saison", "Niveau", "Alertes", "Vaches", "Alertes collectives"]]
    for _, row in confidence.iterrows():
        conf_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(row["reinforced_confidence_level"]).replace("_", " "),
                str(int(row["notifications"])),
                str(int(row["cows"])),
                str(int(row["collective_alerts"])),
            ]
        )
    story.append(_short_table(conf_rows, [1.0 * inch, 2.45 * inch, 0.75 * inch, 0.7 * inch, 1.15 * inch]))

    story.append(PageBreak())
    story.append(_p("5. Concordance avec les observations disponibles", styles["H1Clean"]))
    story.append(_short_table(_concordance_rows(concordance), [1.35 * inch, 1.0 * inch, 1.0 * inch, 0.75 * inch]))
    story.append(
        _p(
            "La concordance temporelle montre qu'une partie des scans comportementaux se situe a "
            "proximite d'une alerte, surtout pour Fall 2019, Summer 2019 et Winter 2019. Cela confirme "
            "que les alertes captent bien des episodes comportementaux dans certaines fenetres.",
            styles["BodyClean"],
        )
    )
    story.append(
        _p(
            "Cette concordance ne constitue pas une validation clinique. Les scans sont ponctuels, "
            "les alertes sont continues, et les observations disponibles ne forment pas une verite-terrain "
            "clinique complete. Le resultat doit donc etre presente comme une verification de coherence, "
            "pas comme une estimation de performance diagnostique.",
            styles["BodyClean"],
        )
    )

    story.append(PageBreak())
    story.append(_p("6. Lecture par saison", styles["H1Clean"]))
    story.append(
        _p(
            "Chaque saison est presentee avec les memes elements: volume de donnees, nombre "
            "d'alertes, reclassification A/B/C, vaches les plus concernees et principaux jours collectifs.",
            styles["BodyClean"],
        )
    )
    existing_seasons = [season for season in SEASON_ORDER if not summary.loc[summary["season"] == season].empty]
    for index, season in enumerate(existing_seasons):
        row = summary.loc[summary["season"] == season]
        r = row.iloc[0]
        if index > 0:
            story.append(PageBreak())
        story.append(_p(SEASON_LABELS.get(season, season), styles["H1Clean"]))
        story.append(
            _short_table(
                [
                    ["Indicateur", "Valeur"],
                    ["Periode", SEASON_PERIODS.get(season, "")],
                    ["Vaches", str(int(r["n_cows"]))],
                    ["Bins 15 min", _fmt_int(r["n_intervals"])],
                    ["Notifications brutes", str(int(r["initial_notifications"]))],
                    ["Vaches avec alerte", str(int(r["alerted_cows"]))],
                    ["Alertes A", str(int(r["A_individuelle_prioritaire"]))],
                    ["Alertes B", str(int(r["B_individuelle_a_verifier"]))],
                    ["Alertes C", str(int(r["C_probable_evenement_collectif"]))],
                    ["Part collective", _fmt_pct(r["collective_flagged_pct"])],
                ],
                [2.1 * inch, 2.5 * inch],
            )
        )
        story.append(_p(_season_interpretation(season), styles["BodyClean"]))

        season_cows = cows[cows["season"] == season].sort_values("total_notifications", ascending=False).head(6)
        if not season_cows.empty:
            story.append(_p("Vaches avec le plus de notifications", styles["H2Clean"]))
            cow_rows = [["Vache", "Total", "A", "B", "C", "Collectif"]]
            for _, cr in season_cows.iterrows():
                cow_rows.append(
                    [
                        str(cr["Cow"]),
                        str(int(cr["total_notifications"])),
                        str(int(cr["A_individuelle_prioritaire"])),
                        str(int(cr["B_individuelle_a_verifier"])),
                        str(int(cr["C_probable_evenement_collectif"])),
                        str(int(cr["collective_flagged"])),
                    ]
                )
            story.append(_short_table(cow_rows, [0.8 * inch, 0.7 * inch, 0.55 * inch, 0.55 * inch, 0.55 * inch, 0.8 * inch]))

        season_days = collective_days[
            (collective_days["season"] == season) & (collective_days["collective_alerts"] > 0)
        ].sort_values("collective_alerts", ascending=False).head(5)
        if not season_days.empty:
            story.append(_p("Principaux jours collectifs", styles["H2Clean"]))
            day_rows = [["Jour", "Alertes", "Vaches", "Collectives", "Fraction troupeau"]]
            for _, dr in season_days.iterrows():
                day_rows.append(
                    [
                        str(dr["Day"]),
                        str(int(dr["n_alerts"])),
                        str(int(dr["alerted_cows"])),
                        str(int(dr["collective_alerts"])),
                        _fmt_pct(dr["alerted_cow_frac"]),
                    ]
                )
            story.append(_short_table(day_rows, [1.1 * inch, 0.75 * inch, 0.75 * inch, 0.9 * inch, 1.05 * inch]))

    story.append(PageBreak())
    story.append(_p("7. Limites scientifiques", styles["H1Clean"]))
    for item in [
        "Les donnees cliniques disponibles ne permettent pas de calculer une sensibilite ou une specificite robuste.",
        "Les cas de boiterie synchrones disponibles sont insuffisants ou trop legers pour servir de verite-terrain forte.",
        "L'IceTag mesure l'activite et les postures; il ne mesure pas directement l'asymetrie fine de la demarche.",
        "Les evenements collectifs doivent etre distingues des anomalies individuelles pour eviter une surestimation.",
        "Fall 2021 est une fenetre courte: le corpus appuie surtout la verification technique.",
    ]:
        story.append(_p(f"- {item}", styles["BodyClean"]))

    story.append(_p("8. Conclusion defendable", styles["H1Clean"]))
    story.append(
        _p(
            "L'Objectif 1 est atteint comme livrable de transfert et d'evaluation: les quatre corpus ont ete traites, "
            "les alertes ont ete produites, comparees aux observations disponibles, puis requalifiees selon leur contexte. "
            "La formulation defendable est: detection et priorisation de signaux comportementaux compatibles avec "
            "une perturbation locomotrice, sans revendication de diagnostic clinique confirme.",
            styles["BodyClean"],
        )
    )
    story.append(
        _callout(
            "Amelioration principale par rapport a la lecture initiale: le rapport ne se limite plus au "
            "nombre brut d'alertes. Il explique lesquelles sont prioritaires, lesquelles doivent etre "
            "verifiees, et lesquelles correspondent probablement a un contexte collectif de troupeau.",
            styles["BodyClean"],
            "#F4F7EC",
        )
    )

    story.append(PageBreak())
    story.append(_p("9. Fichiers du paquet propre", styles["H1Clean"]))
    files_rows = [
        ["Dossier", "Contenu", "Usage"],
        ["RAPPORTS_PDF", "Objectif1_rapport_detaille_lisible.pdf", "Rapport principal a lire/envoyer"],
        ["DONNEES_TRAITEES_ALERTES", "Predictions, alertes seules et resumes par saison", "Livrable SOW 1.1"],
        ["TABLEAUX_CSV", "Syntheses par saison, confiance, vache et jours collectifs", "Resultats exploitables et auditables"],
        ["A_LIRE", "Rapports Markdown propres + README", "Documentation texte et notes de livraison"],
        ["ARCHIVE_INTERNE", "Anciens exports, code et analyses detaillees", "Interne; ne pas envoyer par defaut"],
    ]
    story.append(_short_table(files_rows, [1.55 * inch, 2.9 * inch, 2.15 * inch]))
    story.append(
        _p(
            "Le paquet visible a ete nettoye pour eviter les doublons, les exports PDF illisibles et les "
            "references internes. Les donnees detaillees restent disponibles en CSV pour verification.",
            styles["BodyClean"],
        )
    )

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return pdf_path


def _build_docx() -> Path:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = WORD_DIR / "Objectif1_rapport_detaille_lisible.docx"

    summary = _ordered_summary(
        pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_season.csv")
    )
    confidence = _ordered_confidence(
        pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_confidence.csv")
    )
    cows = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_cow.csv")
    collective_days = pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_collective_days.csv")
    concordance = pd.read_csv(SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv")

    doc = Document()
    _configure_docx_styles(doc)
    doc.core_properties.title = "Objectif 1 - Rapport detaille lisible"
    doc.core_properties.subject = "Projet McGill / WELL-E"
    doc.core_properties.author = "Aliou Barry"

    total_alerts = int(summary["initial_notifications"].sum())
    total_a = int(summary["A_individuelle_prioritaire"].sum())
    total_b = int(summary["B_individuelle_a_verifier"].sum())
    total_c = int(summary["C_probable_evenement_collectif"].sum())

    doc.add_paragraph("Objectif 1", style="Title")
    doc.add_paragraph("Rapport detaille lisible", style="Title")
    doc.add_paragraph(
        "Application et evaluation du pipeline de detection sur les donnees IceTag McGill",
        style="Subtitle",
    )
    doc.add_paragraph(f"Version nettoyee du {date.today().isoformat()}", style="Subtitle")
    _add_docx_callout(
        doc,
        f"Conclusion courte. Les quatre corpus ont ete traites. Le pipeline produit "
        f"{total_alerts} notifications, maintenant requalifiees en {total_a} alertes A, "
        f"{total_b} alertes B et {total_c} alertes C. Ces sorties sont defendables comme "
        "signaux comportementaux a prioriser, pas comme diagnostics cliniques confirmes.",
    )
    _add_docx_table(
        doc,
        [
            ["Element", "Statut"],
            ["Corpus IceTag", "4 saisons traitees"],
            ["Pipeline", "Application complete sur intervalles de 15 min"],
            ["Renforcement", "Normalisation troupeau + filtre collectif + niveaux A/B/C/D"],
            ["Validation clinique", "Non revendiquee sans labels synchrones robustes"],
            ["Livrables", "PDF + Word + tableaux CSV reproductibles"],
        ],
        [1.8, 4.6],
    )

    doc.add_page_break()
    doc.add_heading("1. Perimetre du livrable", level=1)
    doc.add_paragraph(
        "Ce rapport documente ce qui a ete fait pour l'Objectif 1: appliquer le pipeline de "
        "detection sur les donnees accelerometriques McGill, analyser les alertes obtenues, "
        "les comparer aux observations disponibles et renforcer leur interpretation par un contexte troupeau."
    )
    doc.add_paragraph(
        "Le livrable ne pretend pas mesurer une sensibilite ou une specificite clinique. Pour cela, il faudrait "
        "des labels de boiterie synchrones, dates et suffisamment nets. La formulation defendable est donc: "
        "detection et priorisation de signaux comportementaux compatibles avec une perturbation locomotrice."
    )
    doc.add_heading("Travaux couverts", level=2)
    _add_docx_bullets(
        doc,
        [
            "Conversion et harmonisation des donnees IceTag par saison.",
            "Construction de bins de 15 minutes et verification de la couverture.",
            "Application du pipeline de detection sur chaque vache et chaque saison.",
            "Compilation des notifications et syntheses multi-saisons.",
            "Comparaison temporelle avec les observations comportementales disponibles.",
            "Ajout d'une normalisation troupeau et d'un filtre d'evenements collectifs.",
            "Reclassification des alertes en niveaux de confiance A, B, C et D.",
        ],
    )

    doc.add_heading("2. Donnees traitees", level=1)
    data_rows = [["Saison", "Vaches", "Bins 15 min", "Periode", "Couverture"]]
    for _, row in summary.iterrows():
        data_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["n_cows"])),
                _fmt_int(row["n_intervals"]),
                SEASON_PERIODS.get(row["season"], ""),
                f"{row['mean_coverage_pct']:.1f}%",
            ]
        )
    _add_docx_table(doc, data_rows, [1.2, 0.7, 1.05, 2.2, 0.9])
    doc.add_paragraph(
        "Les quatre jeux de donnees ont une couverture suffisante pour produire des alertes. Fall 2021 est "
        "le corpus le plus court; il est donc utile surtout pour verifier que la chaine de traitement s'execute "
        "correctement sur ce format."
    )

    doc.add_page_break()
    doc.add_heading("3. Pipeline applique", level=1)
    doc.add_paragraph(
        "La chaine appliquee reste volontairement simple a auditer: elle part des mesures IceTag, reconstruit "
        "des indicateurs temporels par vache, detecte les anomalies, puis ajoute un niveau d'interpretation "
        "troupeau pour eviter de confondre un evenement collectif avec plusieurs alertes individuelles independantes."
    )
    _add_docx_table(
        doc,
        [
            ["Etape", "Role", "Sortie"],
            ["1. Lecture", "Importer les fichiers saisonniers harmonises", "Table par vache et timestamp"],
            ["2. Binning", "Regrouper les donnees en intervalles de 15 min", "Series regulieres comparables"],
            ["3. Features", "Calculer activite, variations et contexte temporel", "Variables de detection"],
            ["4. Detection", "Identifier les deviations individuelles", "Notifications brutes"],
            ["5. Contexte troupeau", "Comparer chaque vache au troupeau au meme moment", "Signal propre vs collectif"],
            ["6. Reclassification", "Attribuer un niveau A/B/C/D", "Tableaux interpretables"],
        ],
        [1.2, 3.1, 2.0],
    )
    doc.add_heading("Definition des niveaux de confiance", level=2)
    _add_docx_table(
        doc,
        [
            ["Niveau", "Interpretation"],
            ["A", "Alerte individuelle prioritaire: non collective, bonne couverture, deviation propre a la vache."],
            ["B", "Alerte individuelle a verifier: signal non collectif mais support troupeau moins net."],
            ["C", "Probable evenement collectif: plusieurs vaches alertees dans une fenetre courte."],
            ["D", "Qualite ou contexte insuffisant: alerte conservee mais interpretation faible."],
        ],
        [0.8, 5.6],
    )

    doc.add_page_break()
    doc.add_heading("4. Resultats globaux", level=1)
    alert_rows = [["Saison", "Total", "A prioritaire", "B a verifier", "C collectif"]]
    for _, row in summary.iterrows():
        alert_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["initial_notifications"])),
                str(int(row["A_individuelle_prioritaire"])),
                str(int(row["B_individuelle_a_verifier"])),
                str(int(row["C_probable_evenement_collectif"])),
            ]
        )
    _add_docx_table(doc, alert_rows, [1.2, 0.8, 1.2, 1.2, 1.2])
    doc.add_paragraph(
        "La reclassification change la lecture de l'Objectif 1. Le nombre brut d'alertes reste important, "
        "mais toutes les alertes n'ont pas la meme priorite. Les alertes C ne sont pas ignorees; elles sont "
        "conservees comme signaux de contexte collectif et sorties de la liste des suspicions individuelles prioritaires."
    )
    doc.add_heading("Synthese par niveau", level=2)
    conf_rows = [["Saison", "Niveau", "Alertes", "Vaches", "Alertes collectives"]]
    level_short = {
        "A_individuelle_prioritaire": "A",
        "B_individuelle_a_verifier": "B",
        "C_probable_evenement_collectif": "C",
        "D_qualite_ou_contexte_insuffisant": "D",
    }
    for _, row in confidence.iterrows():
        conf_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                level_short.get(str(row["reinforced_confidence_level"]), str(row["reinforced_confidence_level"])),
                str(int(row["notifications"])),
                str(int(row["cows"])),
                str(int(row["collective_alerts"])),
            ]
        )
    _add_docx_table(doc, conf_rows, [1.3, 0.8, 0.9, 0.8, 1.3])

    doc.add_page_break()
    doc.add_heading("5. Concordance avec les observations disponibles", level=1)
    _add_docx_table(doc, _concordance_rows(concordance), [1.4, 1.0, 1.0, 0.8])
    doc.add_paragraph(
        "La concordance temporelle montre qu'une partie des scans comportementaux se situe a proximite d'une "
        "alerte, surtout pour Fall 2019, Summer 2019 et Winter 2019. Cela confirme que les alertes captent bien "
        "des episodes comportementaux dans certaines fenetres."
    )
    doc.add_paragraph(
        "Cette concordance ne constitue pas une validation clinique. Les scans sont ponctuels, les alertes sont "
        "continues, et les observations disponibles ne forment pas une verite-terrain clinique complete. Le resultat "
        "doit donc etre presente comme une verification de coherence, pas comme une estimation de performance diagnostique."
    )

    doc.add_page_break()
    doc.add_heading("6. Lecture par saison", level=1)
    doc.add_paragraph(
        "Chaque saison est presentee avec les memes elements: volume de donnees, nombre d'alertes, "
        "reclassification A/B/C, vaches les plus concernees et principaux jours collectifs."
    )
    existing_seasons = [season for season in SEASON_ORDER if not summary.loc[summary["season"] == season].empty]
    for index, season in enumerate(existing_seasons):
        if index > 0:
            doc.add_page_break()
        row = summary.loc[summary["season"] == season].iloc[0]
        doc.add_heading(SEASON_LABELS.get(season, season), level=1)
        _add_docx_table(
            doc,
            [
                ["Indicateur", "Valeur"],
                ["Periode", SEASON_PERIODS.get(season, "")],
                ["Vaches", str(int(row["n_cows"]))],
                ["Bins 15 min", _fmt_int(row["n_intervals"])],
                ["Notifications brutes", str(int(row["initial_notifications"]))],
                ["Vaches avec alerte", str(int(row["alerted_cows"]))],
                ["Alertes A", str(int(row["A_individuelle_prioritaire"]))],
                ["Alertes B", str(int(row["B_individuelle_a_verifier"]))],
                ["Alertes C", str(int(row["C_probable_evenement_collectif"]))],
                ["Part collective", _fmt_pct(row["collective_flagged_pct"])],
            ],
            [2.0, 2.6],
        )
        doc.add_paragraph(_season_interpretation(season))

        season_cows = cows[cows["season"] == season].sort_values("total_notifications", ascending=False).head(6)
        if not season_cows.empty:
            doc.add_heading("Vaches avec le plus de notifications", level=2)
            cow_rows = [["Vache", "Total", "A", "B", "C", "Collectif"]]
            for _, cr in season_cows.iterrows():
                cow_rows.append(
                    [
                        str(cr["Cow"]),
                        str(int(cr["total_notifications"])),
                        str(int(cr["A_individuelle_prioritaire"])),
                        str(int(cr["B_individuelle_a_verifier"])),
                        str(int(cr["C_probable_evenement_collectif"])),
                        str(int(cr["collective_flagged"])),
                    ]
                )
            _add_docx_table(doc, cow_rows, [0.85, 0.75, 0.55, 0.55, 0.55, 0.85])

        season_days = collective_days[
            (collective_days["season"] == season) & (collective_days["collective_alerts"] > 0)
        ].sort_values("collective_alerts", ascending=False).head(5)
        if not season_days.empty:
            doc.add_heading("Principaux jours collectifs", level=2)
            day_rows = [["Jour", "Alertes", "Vaches", "Collectives", "Fraction troupeau"]]
            for _, dr in season_days.iterrows():
                day_rows.append(
                    [
                        str(dr["Day"]),
                        str(int(dr["n_alerts"])),
                        str(int(dr["alerted_cows"])),
                        str(int(dr["collective_alerts"])),
                        _fmt_pct(dr["alerted_cow_frac"]),
                    ]
                )
            _add_docx_table(doc, day_rows, [1.15, 0.75, 0.75, 0.9, 1.05])

    doc.add_page_break()
    doc.add_heading("7. Limites scientifiques", level=1)
    _add_docx_bullets(
        doc,
        [
            "Les donnees cliniques disponibles ne permettent pas de calculer une sensibilite ou une specificite robuste.",
            "Les cas de boiterie synchrones disponibles sont insuffisants ou trop legers pour servir de verite-terrain forte.",
            "L'IceTag mesure l'activite et les postures; il ne mesure pas directement l'asymetrie fine de la demarche.",
            "Les evenements collectifs doivent etre distingues des anomalies individuelles pour eviter une surestimation.",
            "Fall 2021 est une fenetre courte: le corpus appuie surtout la verification technique.",
        ],
    )
    doc.add_heading("8. Conclusion defendable", level=1)
    doc.add_paragraph(
        "L'Objectif 1 est atteint comme livrable de transfert et d'evaluation: les quatre corpus ont ete traites, "
        "les alertes ont ete produites, comparees aux observations disponibles, puis requalifiees selon leur contexte. "
        "La formulation defendable est: detection et priorisation de signaux comportementaux compatibles avec une "
        "perturbation locomotrice, sans revendication de diagnostic clinique confirme."
    )
    _add_docx_callout(
        doc,
        "Amelioration principale par rapport a la lecture initiale: le rapport ne se limite plus au nombre brut "
        "d'alertes. Il explique lesquelles sont prioritaires, lesquelles doivent etre verifiees, et lesquelles "
        "correspondent probablement a un contexte collectif de troupeau.",
    )

    doc.add_page_break()
    doc.add_heading("9. Fichiers du paquet propre", level=1)
    _add_docx_table(
        doc,
        [
            ["Dossier", "Contenu", "Usage"],
            ["RAPPORTS_PDF", "Objectif1_rapport_detaille_lisible.pdf", "Rapport principal a lire/envoyer"],
            ["RAPPORTS_WORD", "Objectif1_rapport_detaille_lisible.docx", "Version Word ouvrable dans Pages"],
            ["DONNEES_TRAITEES_ALERTES", "Predictions, alertes seules et resumes par saison", "Livrable SOW 1.1"],
            ["TABLEAUX_CSV", "Syntheses par saison, confiance, vache et jours collectifs", "Resultats exploitables"],
            ["A_LIRE", "Rapports Markdown propres + README", "Documentation texte et notes de livraison"],
            ["ARCHIVE_INTERNE", "Anciens exports, code et analyses detaillees", "Interne; ne pas envoyer par defaut"],
        ],
        [1.45, 3.05, 1.9],
    )
    doc.add_paragraph(
        "Le paquet visible a ete nettoye pour eviter les doublons, les exports illisibles et les references internes. "
        "Les donnees detaillees restent disponibles en CSV pour verification."
    )

    doc.save(docx_path)
    return docx_path


def _build_pdf() -> Path:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = PDF_DIR / "Objectif1_rapport_livraison.pdf"

    summary = _ordered_summary(
        pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_season.csv")
    )
    coverage_by_season = (
        pd.read_csv(SOURCE / "objective1_multi_season_summary.csv").set_index("season")["mean_coverage"].to_dict()
    )
    concordance = pd.read_csv(SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv")
    sls_current = _load_current_sls_summary()

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="DeliveryTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#000000"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DeliveryH1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#000000"),
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DeliveryBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14.5,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DeliverySmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=10.5,
            textColor=colors.HexColor("#555555"),
        )
    )

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="Objectif 1 - Rapport de livraison",
    )

    total_alerts = int(summary["initial_notifications"].sum())
    total_a = int(summary["A_individuelle_prioritaire"].sum())
    total_b = int(summary["B_individuelle_a_verifier"].sum())
    total_c = int(summary["C_probable_evenement_collectif"].sum())

    story = [
        Spacer(1, 0.65 * inch),
        _p("Objectif 1", styles["DeliveryTitle"]),
        _p("Rapport de livraison", styles["DeliveryTitle"]),
        _p("Projet McGill / WELL-E - Données IceTag", styles["DeliveryBody"]),
        Spacer(1, 0.25 * inch),
        _callout(
            f"<b>Conclusion.</b> Les quatre saisons ont été traitées. Le pipeline a produit "
            f"{total_alerts} alertes, reclassées en {total_a} alertes A, {total_b} alertes B "
            f"et {total_c} alertes C. Le résultat est défendable comme détection de signaux "
            "comportementaux à vérifier, sans revendication de diagnostic clinique confirmé.",
            styles["DeliveryBody"],
        ),
    ]

    story.append(_p("1. Correspondance avec le SOW", styles["DeliveryH1"]))
    story.append(
        _short_table(
            [
                ["Livrable SOW", "Où le trouver"],
                ["1.1 Données traitées avec alertes", "DONNEES_TRAITEES_ALERTES/"],
                ["1.1 Note technique de reproductibilité", "NOTES_SOW/note_technique_reproductibilite.md"],
                ["1.2 Table de concordance", "TABLEAUX_CSV/table_concordance.csv"],
                ["1.2 Rapport court de validation", "NOTES_SOW/rapport_validation_concordance.md"],
            ],
            [2.8 * inch, 4.2 * inch],
        )
    )

    story.append(_p("2. Données traitées", styles["DeliveryH1"]))
    data_rows = [["Saison", "Profils", "Intervalles 15 min", "Période", "Couverture"]]
    for _, row in summary.iterrows():
        data_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                _display_n_cows(row["season"], row["n_cows"]),
                _fmt_int(row["n_intervals"]),
                SEASON_PERIODS_SHORT.get(row["season"], ""),
                f"{float(coverage_by_season.get(row['season'], row['mean_coverage_pct'])):.1f}%",
            ]
        )
    story.append(
        _short_table(
            data_rows,
            [1.15 * inch, 1.35 * inch, 1.25 * inch, 2.35 * inch, 0.9 * inch],
        )
    )
    story.append(
        _p(
            "Fall 2021 contient 10 profils traités pour la traçabilité, dont 8 couvrent la fenêtre complète annoncée au SOW. "
            "Les deux profils partiels servent uniquement à la vérification technique.",
            styles["DeliverySmall"],
        )
    )

    story.append(_p("3. Alertes produites", styles["DeliveryH1"]))
    alert_rows = [["Saison", "Total", "A", "B", "C"]]
    for _, row in summary.iterrows():
        alert_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["initial_notifications"])),
                str(int(row["A_individuelle_prioritaire"])),
                str(int(row["B_individuelle_a_verifier"])),
                str(int(row["C_probable_evenement_collectif"])),
            ]
        )
    story.append(_short_table(alert_rows, [2.2 * inch, 1.2 * inch, 1.2 * inch, 1.2 * inch, 1.2 * inch]))
    story.append(
        _p(
            "A = prioritaire à vérifier individuellement; B = individuelle à vérifier; "
            "C = probablement liée à un contexte collectif du troupeau.",
            styles["DeliveryBody"],
        )
    )

    story.append(_p("4. Concordance avec les observations", styles["DeliveryH1"]))
    story.append(_short_table(_concordance_rows(concordance), [2.2 * inch, 1.6 * inch, 1.6 * inch, 1.6 * inch]))
    story.append(
        _p(
            "La concordance temporelle est présente pour Winter 2019, Summer 2019 et Fall 2019. "
            "Elle confirme une cohérence entre certaines alertes et les observations disponibles, "
            "mais elle ne remplace pas une validation clinique.",
            styles["DeliveryBody"],
        )
    )

    story.append(PageBreak())
    story.append(_p("5. Concordance exploratoire avec les scores SLS", styles["DeliveryH1"]))
    story.append(
        _short_table(
            [
                ["Évaluation", "Cohorte", "Résultat", "Lecture"],
                [
                    "Pipeline initiale IF + règles",
                    f'{SLS_INITIAL["n_evaluable"]} vaches; {SLS_INITIAL["n_sls_ge_2"]} avec SLS >= 2',
                    f'p = {SLS_INITIAL["mann_whitney_p"]:.3f}; rho = {SLS_INITIAL["spearman_rho"]:.3f}',
                    "Pas de concordance observée avec les SLS.",
                ],
                [
                    "Pipeline HYPO + instabilité + hybride",
                    f'{sls_current["n_evaluable"]} évaluables; {sls_current["n_sls_ge_2"]} avec SLS >= 2',
                    f'AUC = {sls_current["auc"]:.3f}; p = {sls_current["mann_whitney_p"]:.3f}',
                    "Signal encourageant, strictement exploratoire.",
                ],
            ],
            [1.55 * inch, 1.65 * inch, 1.55 * inch, 2.25 * inch],
        )
    )
    story.append(
        _p(
            "Des scores SLS synchronisés existent pour Winter 2019. Les deux évaluations utilisent toutefois des cohortes et "
            "des protocoles différents; elles ne constituent pas une comparaison directe de performance. Avec seulement trois "
            "vaches SLS >= 2 dans l'évaluation actuelle et un effet du traitement Exercise, aucune sensibilité ou spécificité "
            "robuste ne peut être revendiquée.",
            styles["DeliveryBody"],
        )
    )

    story.append(_p("6. Limites et interprétation", styles["DeliveryH1"]))
    for item in [
        "Les alertes sont des signaux comportementaux, pas des diagnostics cliniques confirmés.",
        "Les scores SLS disponibles soutiennent une concordance exploratoire limitée, pas une validation clinique complète.",
        "La petite cohorte et le faible nombre de cas SLS >= 2 empêchent une estimation robuste de sensibilité et de spécificité.",
        "Les épisodes collectifs sont séparés des alertes individuelles pour éviter une surestimation.",
        "Fall 2021 comporte 8 profils complets et 2 profils partiels; son rôle principal est la compatibilité technique.",
    ]:
        story.append(_p(f"- {item}", styles["DeliveryBody"]))

    story.append(_p("7. Contenu du dossier livré", styles["DeliveryH1"]))
    story.append(
        _short_table(
            [
                ["Dossier", "Contenu"],
                ["RAPPORTS", "Rapport Word et présentation PowerPoint"],
                ["DONNEES_TRAITEES_ALERTES", "Prédictions, alertes seules et résumés par saison"],
                ["TABLEAUX_CSV", "Tables de synthèse et de concordance"],
                ["NOTES_SOW", "Deux notes courtes correspondant aux livrables SOW"],
                ["ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride", "Comparaison de la pipeline actuelle et validation SLS exploratoire"],
            ],
            [2.5 * inch, 4.5 * inch],
        )
    )

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return pdf_path


def _build_docx() -> Path:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = WORD_DIR / "Objectif1_rapport_livraison.docx"

    summary = _ordered_summary(
        pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_season.csv")
    )
    coverage_by_season = (
        pd.read_csv(SOURCE / "objective1_multi_season_summary.csv").set_index("season")["mean_coverage"].to_dict()
    )
    concordance = pd.read_csv(SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv")
    sls_current = _load_current_sls_summary()

    doc = Document()
    _configure_docx_styles(doc)
    doc.core_properties.title = "Objectif 1 - Rapport de livraison"
    doc.core_properties.subject = "Projet McGill / WELL-E"
    doc.core_properties.author = "Aliou Barry"

    total_alerts = int(summary["initial_notifications"].sum())
    total_a = int(summary["A_individuelle_prioritaire"].sum())
    total_b = int(summary["B_individuelle_a_verifier"].sum())
    total_c = int(summary["C_probable_evenement_collectif"].sum())

    doc.add_paragraph("Objectif 1", style="Title")
    doc.add_paragraph("Rapport de livraison", style="Title")
    doc.add_paragraph("Projet McGill / WELL-E - Données IceTag", style="Subtitle")
    _add_docx_callout(
        doc,
        f"Conclusion. Les quatre saisons ont été traitées. Le pipeline a produit {total_alerts} alertes, "
        f"reclassées en {total_a} alertes A, {total_b} alertes B et {total_c} alertes C. Le résultat "
        "est défendable comme détection de signaux comportementaux à vérifier, sans revendication de "
        "diagnostic clinique confirmé.",
    )

    cadrage = doc.add_paragraph()
    _add_docx_run(cadrage, "Cadrage des pipelines. ", bold=True)
    _add_docx_run(
        cadrage,
        f"La pipeline IF + règles est la baseline initiale livrée (gelée) au titre du SOW, "
        f"qui a produit les {total_alerts} alertes. L'approche HYPO + instabilité + hybride du "
        "mémoire, fournie en annexe, est l'évolution montrée supérieure (localisation, "
        "concordance SLS exploratoire) et représente la direction actuelle du travail.",
    )

    doc.add_heading("1. Correspondance avec le SOW", level=1)
    _add_docx_table(
        doc,
        [
            ["Livrable SOW", "Où le trouver"],
            ["1.1 Données traitées avec alertes", "DONNEES_TRAITEES_ALERTES/"],
            ["1.1 Note technique de reproductibilité", "NOTES_SOW/note_technique_reproductibilite.md"],
            ["1.2 Table de concordance", "TABLEAUX_CSV/table_concordance.csv"],
            ["1.2 Rapport court de validation", "NOTES_SOW/rapport_validation_concordance.md"],
        ],
        [2.6, 3.9],
    )

    doc.add_heading("2. Données traitées", level=1)
    data_rows = [["Saison", "Profils", "Intervalles 15 min", "Période", "Couverture"]]
    for _, row in summary.iterrows():
        data_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                _display_n_cows(row["season"], row["n_cows"]),
                _fmt_int(row["n_intervals"]),
                SEASON_PERIODS_SHORT.get(row["season"], ""),
                f"{float(coverage_by_season.get(row['season'], row['mean_coverage_pct'])):.1f}%",
            ]
        )
    _add_docx_table(doc, data_rows, [1.1, 1.25, 1.25, 2.0, 0.9])
    doc.add_paragraph(
        "Fall 2021 contient 10 profils traités pour la traçabilité, dont 8 couvrent la fenêtre complète annoncée au SOW. "
        "Les deux profils partiels servent uniquement à la vérification technique."
    )

    doc.add_heading("3. Alertes produites", level=1)
    alert_rows = [["Saison", "Total", "A", "B", "C"]]
    for _, row in summary.iterrows():
        alert_rows.append(
            [
                SEASON_LABELS.get(row["season"], row["season"]),
                str(int(row["initial_notifications"])),
                str(int(row["A_individuelle_prioritaire"])),
                str(int(row["B_individuelle_a_verifier"])),
                str(int(row["C_probable_evenement_collectif"])),
            ]
        )
    _add_docx_table(doc, alert_rows, [2.1, 1.1, 1.1, 1.1, 1.1])
    doc.add_paragraph(
        "A = prioritaire à vérifier individuellement; B = individuelle à vérifier; "
        "C = probablement liée à un contexte collectif du troupeau."
    )

    doc.add_heading("4. Concordance avec les observations", level=1)
    _add_docx_table(doc, _concordance_rows(concordance), [2.0, 1.5, 1.5, 1.5])
    doc.add_paragraph(
        "La concordance temporelle est présente pour Winter 2019, Summer 2019 et Fall 2019. "
        "Elle confirme une cohérence entre certaines alertes et les observations disponibles, mais elle ne "
        "remplace pas une validation clinique."
    )

    doc.add_page_break()
    doc.add_heading("5. Concordance exploratoire avec les scores SLS", level=1)
    _add_docx_table(
        doc,
        [
            ["Évaluation", "Cohorte", "Résultat", "Lecture"],
            [
                "Pipeline initiale IF + règles",
                f'{SLS_INITIAL["n_evaluable"]} vaches; {SLS_INITIAL["n_sls_ge_2"]} avec SLS >= 2',
                f'p = {SLS_INITIAL["mann_whitney_p"]:.3f}; rho = {SLS_INITIAL["spearman_rho"]:.3f}',
                "Pas de concordance observée avec les SLS.",
            ],
            [
                "Pipeline HYPO + instabilité + hybride",
                f'{sls_current["n_evaluable"]} évaluables; {sls_current["n_sls_ge_2"]} avec SLS >= 2',
                f'AUC = {sls_current["auc"]:.3f}; p = {sls_current["mann_whitney_p"]:.3f}',
                "Signal encourageant, strictement exploratoire.",
            ],
        ],
        [1.35, 1.45, 1.35, 2.35],
    )
    doc.add_paragraph(
        "Des scores SLS synchronisés existent pour Winter 2019. Les deux évaluations utilisent toutefois des cohortes et des "
        "protocoles différents; elles ne constituent pas une comparaison directe de performance. Avec seulement trois vaches "
        "SLS >= 2 dans l'évaluation actuelle et un effet du traitement Exercise, aucune sensibilité ou spécificité robuste ne peut "
        "être revendiquée."
    )

    doc.add_heading("6. Limites et interprétation", level=1)
    _add_docx_bullets(
        doc,
        [
            "Les alertes sont des signaux comportementaux, pas des diagnostics cliniques confirmés.",
            "Les scores SLS disponibles soutiennent une concordance exploratoire limitée, pas une validation clinique complète.",
            "La petite cohorte et le faible nombre de cas SLS >= 2 empêchent une estimation robuste de sensibilité et de spécificité.",
            "Les épisodes collectifs sont séparés des alertes individuelles pour éviter une surestimation.",
            "Fall 2021 comporte 8 profils complets et 2 profils partiels; son rôle principal est la compatibilité technique.",
        ],
    )

    doc.add_heading("7. Contenu du dossier livré", level=1)
    _add_docx_table(
        doc,
        [
            ["Dossier", "Contenu"],
            ["RAPPORTS", "Rapport Word et présentation PowerPoint"],
            ["DONNEES_TRAITEES_ALERTES", "Prédictions, alertes seules et résumés par saison"],
            ["TABLEAUX_CSV", "Tables de synthèse et de concordance"],
            ["NOTES_SOW", "Deux notes courtes correspondant aux livrables SOW"],
            ["ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride", "Comparaison de la pipeline actuelle et validation SLS exploratoire"],
        ],
        [2.2, 4.3],
    )

    _apply_french_decimals(doc)
    _clean_docx_metadata(doc)
    doc.save(docx_path)
    return docx_path


def _write_sow_notes() -> None:
    READ_ME.mkdir(parents=True, exist_ok=True)
    summary = _ordered_summary(
        pd.read_csv(SOURCE / "renforcement_scientifique" / "objective1_reinforced_summary_by_season.csv")
    )
    concordance = pd.read_csv(SOURCE / "tache1_2_concordance" / "concordance_par_experience.csv")
    sls_current = _load_current_sls_summary()

    total_intervals = int(summary["n_intervals"].sum())
    total_alerts = int(summary["initial_notifications"].sum())
    seasons = ", ".join(SEASON_LABELS.get(s, s) for s in summary["season"].tolist())

    technical_note = f"""# Note technique de reproductibilité - Objectif 1

Cette note confirme que le pipeline de détection a été appliqué aux quatre corpus IceTag fournis pour l'Objectif 1.

## Données traitées
- Saisons: {seasons}
- Total analysé: {_fmt_int(total_intervals)} intervalles de 15 minutes
- Alertes produites: {total_alerts}
- Fall 2021: 10 profils traités, dont 8 couvrent la fenêtre complète du SOW et 2 sont partiels.

## Traitement appliqué
- Harmonisation des timestamps et des identifiants de vaches.
- Construction de séries régulières par vache en intervalles de 15 minutes.
- Application du pipeline de détection sur chaque saison.
- Export des prédictions, alertes seules et résumés par saison.
- Ajout d'une lecture troupeau pour séparer les alertes individuelles des épisodes collectifs.

## Sorties
- DONNEES_TRAITEES_ALERTES/: prédictions, alertes et résumés par saison.
- TABLEAUX_CSV/: synthèses et tables de concordance.

## Limite
Les sorties sont reproductibles comme alertes comportementales. Des scores SLS synchronisés sont disponibles pour une sous-cohorte Winter 2019, mais le faible nombre de cas SLS >= 2 ne permet pas une validation clinique complète.
"""

    rows = _concordance_rows(concordance)
    table_lines = ["| Expérience | Scans alignés | Avec alerte | Taux |", "|---|---:|---:|---:|"]
    for row in rows[1:]:
        table_lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")

    validation_note = f"""# Rapport court de validation - Objectif 1.2

Objet: comparer les alertes produites avec les observations comportementales disponibles.

## Concordance temporelle
{chr(10).join(table_lines)}

## Interprétation
Une concordance temporelle est observée pour Winter 2019, Summer 2019 et Fall 2019. Fall 2021 est trop court pour soutenir une interprétation forte.

## Concordance exploratoire avec les scores SLS
- Pipeline initiale IF + règles: {SLS_INITIAL["n_evaluable"]} vaches, dont {SLS_INITIAL["n_sls_ge_2"]} avec SLS >= 2; Mann-Whitney p = {SLS_INITIAL["mann_whitney_p"]:.3f}; Spearman rho = {SLS_INITIAL["spearman_rho"]:.3f}. Aucune concordance n'est observée.
- Pipeline HYPO + instabilité + hybride: {sls_current["n_evaluable"]} vaches évaluables, dont {sls_current["n_sls_ge_2"]} avec SLS >= 2; AUC = {sls_current["auc"]:.3f}; Mann-Whitney p = {sls_current["mann_whitney_p"]:.3f}. Le signal est encourageant mais strictement exploratoire.

Les deux évaluations utilisent des cohortes et des protocoles différents. Elles ne constituent donc pas une comparaison directe de performance, ni une estimation clinique de sensibilité ou de spécificité.

## Conclusion
La comparaison soutient une cohérence temporelle entre certaines alertes et les observations disponibles. Les SLS ajoutent une validation observationnelle exploratoire limitée, mais ne permettent pas de calculer une sensibilité ou une spécificité robuste.
"""

    (READ_ME / "note_technique_reproductibilite.md").write_text(technical_note, encoding="utf-8")
    (READ_ME / "rapport_validation_concordance.md").write_text(validation_note, encoding="utf-8")


def _write_readme(docx_path: Path) -> None:
    _write_sow_notes()
    text = f"""Objectif 1 - livrables McGill / WELL-E

Contenu à envoyer:
- RAPPORTS/
- DONNEES_TRAITEES_ALERTES/
- TABLEAUX_CSV/
- NOTES_SOW/
- ANNEXE_pipeline_actuelle_HYPO_instabilite_hybride/ (analyse complémentaire)

Rapports principaux:
- RAPPORTS/{docx_path.name}
- RAPPORTS/Objectif1_presentation_detaillee.pptx

Correspondance SOW:
- 1.1 Données traitées avec alertes: DONNEES_TRAITEES_ALERTES/
- 1.1 Note technique de reproductibilité: NOTES_SOW/note_technique_reproductibilite.md
- 1.2 Table de concordance: TABLEAUX_CSV/table_concordance.csv
- 1.2 Rapport court de validation: NOTES_SOW/rapport_validation_concordance.md

Résumé:
Les quatre saisons IceTag ont été traitées. Les alertes produites sont livrées avec des tableaux de synthèse, une interprétation par niveaux de confiance et une concordance SLS exploratoire sur Winter 2019. Les résultats doivent être lus comme des signaux comportementaux à vérifier, non comme des diagnostics cliniques confirmés.

Note sur les dates des fichiers:
Les CSV de pipeline par saison conservent leur date de génération originale du 29 mai 2026. Ils sont inclus tels quels pour préserver la traçabilité. Le fichier objective1_reinforced_alerts.csv et les tableaux de synthèse renforcés correspondent à l'analyse finale du 13 juillet 2026.
"""
    ROOT_README.write_text(text, encoding="utf-8")


def main() -> None:
    DEST.mkdir(parents=True, exist_ok=True)
    _archive_existing_visible_noise()
    _copy_csv_outputs()
    docx_path = _build_docx()
    _write_readme(docx_path)
    print(f"Version Word creee: {docx_path}")
    print(f"README: {ROOT_README}")
    print(f"CSV: {CSV_DIR}")
    print(f"Notes SOW: {READ_ME}")
    print(f"Archive interne hors livrable: {ARCHIVE}")


if __name__ == "__main__":
    main()
